# ╔══════════════════════════════════════════════════════════════════╗
# ║  ui/tab5_report.py — ITM Pave Pro                               ║
# ║  Project Save / Load — JSON                                     ║
# ║  พัฒนาโดย รศ.ดร.อิทธิพล มีผล | ภาควิชาครุศาสตร์โยธา มจพ.    ║
# ╚══════════════════════════════════════════════════════════════════╝

import json
import io
from datetime import datetime
import streamlit as st
import pandas as pd
import numpy as np


# ─────────────────────────────────────────────
#  Combined Report Helpers
# ─────────────────────────────────────────────

def _make_combined_settings(chapter: int) -> dict:
    """
    สร้าง report_settings ทุกส่วนจาก chapter number เดียว
    ตรงกับโครงสร้างรายงานจริง:
      N.2   ESAL  (Flex: N.2.2 / ตาราง N-1→N-4,  Rigid: N.2.3 / ตาราง N-5→N-7)
      N.3   CBR   (ตาราง N-7,  รูป N-7)
      N.4   Flexible (ตาราง N-8, N-9, N-10,  รูป N-8)
      N.5   Rigid   (รูป N-4→N-8)
    """
    c = str(chapter)
    return {
        'esal': {
            'flex_section_number':  f'{c}.2.2',
            'rigid_section_number': f'{c}.2.3',
            'flex_table_start':     f'{c}-1',
            'rigid_table_start':    f'{c}-5',
        },
        'cbr': {
            'section_number': f'{c}.3',
            'table_number':   f'{c}-7',
            'figure_number':  f'{c}-7',
        },
        'flex': {
            'section_number':  f'{c}.4',
            'table_inputs':    f'{c}-8',
            'table_materials': f'{c}-9',
            'table_sn':        f'{c}-10',
            'figure_number':   f'{c}-8',
        },
        'rigid': {
            'section_number': f'{c}.5',
            'figure_prefix':  f'{c}-',
            'figure_start':   4,
            'inc_summary':    True,
        },
    }


def _merge_docx_bytes(composer, sub_bytes: bytes):
    """
    Append เนื้อหาจาก sub_bytes เข้า master ผ่าน docxcompose.Composer

    ใช้ Composer.append() แทนการ copy XML element ดิบ
    (master_doc.element.body.append) เพราะวิธีดิบจะไม่ copy
    relationship ของรูปภาพ (rId ใน document.xml.rels) ตามไปด้วย
    ทำให้รูปในรายงานรวมหาย ("The picture can't be displayed")
    docxcompose จะ remap relationship + copy ไฟล์รูปให้ครบถูกต้อง
    """
    from docx import Document as DocxDoc
    sub_doc = DocxDoc(io.BytesIO(sub_bytes))
    composer.append(sub_doc)


def _generate_combined_report(ss, chapter: int):
    """
    เรียก builder ทุกส่วนตามข้อมูลที่มี แล้ว merge เป็นไฟล์เดียว
    ผลลัพธ์เก็บใน ss['_combined_report_bytes']
    """
    import copy
    from docx import Document as DocxDoc
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from datetime import datetime

    settings = _make_combined_settings(chapter)
    proj     = ss.get('project_name', '') or 'โครงการ'
    errors   = []
    sections_built = []

    # ── สร้าง master document ──
    master = DocxDoc()
    style  = master.styles['Normal']
    style.font.name = 'TH SarabunPSK'
    style.font.size = Pt(15)
    try:
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')
    except Exception:
        pass

    def _run(para, text, bold=False, size=15, color=None):
        r = para.add_run(text)
        r.font.name = 'TH SarabunPSK'
        r.font.size = Pt(size)
        r.font.bold = bold
        if color:
            r.font.color.rgb = color
        try:
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')
        except Exception:
            pass

    # ── Cover page ──
    p1 = master.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p1, "รายการคำนวณออกแบบโครงสร้างชั้นทาง", bold=True, size=20)

    p2 = master.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p2, "ตามวิธี AASHTO 1993 Guide for Design of Pavement Structures", size=15)

    p3 = master.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p3, f"โครงการ: {proj}", bold=True, size=15,
         color=RGBColor(0x00, 0x47, 0xAB))

    p4 = master.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p4,
         f"บทที่ {chapter}  |  "
         f"วันที่: {datetime.now().strftime('%d/%m/%Y %H:%M')}  |  "
         "พัฒนาโดย รศ.ดร.อิทธิพล มีผล  |  ภาควิชาครุศาสตร์โยธา มจพ.",
         size=12, color=RGBColor(80, 80, 80))

    master.add_page_break()

    # ── สร้าง Composer สำหรับ merge ทุก section (รูปไม่หาย) ──
    # lazy import — กันพังตอน startup
    from docxcompose.composer import Composer
    composer = Composer(master)

    # ── 1. ESAL ──
    if ss.get('traffic_df') is not None:
        try:
            from engine.report_esal import build_esal_report
            ss_esal = dict(ss)
            esal_settings = dict(settings['esal'])
            esal_settings['survey_map_img'] = ss.get('_survey_map_img')  # bytes หรือ None
            ss_esal['report_settings'] = esal_settings
            b = build_esal_report(ss_esal)
            if b:
                _merge_docx_bytes(composer, b)
                master.add_page_break()
                sections_built.append('🚛 ESAL')
        except Exception as e:
            errors.append(f'ESAL: {e}')

    # ── 2. CBR ──
    if ss.get('cbr_values'):
        try:
            from engine.report_cbr import build_cbr_report
            ss_cbr = dict(ss)
            ss_cbr.update(settings['cbr'])
            b = build_cbr_report(ss_cbr)
            if b:
                _merge_docx_bytes(composer, b)
                master.add_page_break()
                sections_built.append('📊 CBR')
        except Exception as e:
            errors.append(f'CBR: {e}')

    # ── 3. Flexible ──
    if ss.get('flex_results'):
        try:
            from engine.report_flexible import build_flexible_report
            ss_flex = dict(ss)
            ss_flex['report_settings'] = settings['flex']
            b = build_flexible_report(ss_flex)
            if b:
                _merge_docx_bytes(composer, b)
                master.add_page_break()
                sections_built.append('🔧 Flexible')
        except Exception as e:
            errors.append(f'Flexible: {e}')

    # ── 4. Rigid ──
    if ss.get('rigid_results'):
        try:
            from engine.report_rigid import build_rigid_report
            ss_rigid = dict(ss)
            ss_rigid['report_settings'] = settings['rigid']
            b = build_rigid_report(ss_rigid)
            if b:
                _merge_docx_bytes(composer, b)
                sections_built.append('🏗️ Rigid')
        except Exception as e:
            errors.append(f'Rigid: {e}')

    # ── บันทึก bytes ──
    buf = io.BytesIO()
    composer.save(buf)
    buf.seek(0)
    ss['_combined_report_bytes'] = buf.read()

    # ── แสดงผล ──
    if sections_built:
        st.success(f"✅ สร้างรายงานรวมสำเร็จ — {len(sections_built)} ส่วน: {' | '.join(sections_built)}")
    if errors:
        for err in errors:
            st.warning(f"⚠️ ข้ามส่วน {err}")


def _has_data(val) -> bool:
    """เช็คว่า session state value มีข้อมูลจริง รองรับ DataFrame ด้วย"""
    if val is None:
        return False
    try:
        return len(val) > 0
    except TypeError:
        return bool(val)


# ─────────────────────────────────────────────
#  Save ทั้ง session (ไม่ใช้ SAVE_KEYS allowlist อีกต่อไป)
#  เพื่อไม่ให้ข้อมูล widget ของหน้าโครงสร้างชั้นทาง (fmat_i, jpcp_name_i,
#  jpcp_n, flex_n_layers, fsub_i ฯลฯ) หลุดหายเหมือนที่ผ่านมา —
#  การ apply ค่ากลับตอน Load ทำที่ app.py ก่อน st.tabs() render จึงไม่ชนกับ
#  กฎ "ห้าม set widget key หลัง widget instantiate แล้ว" ของ Streamlit
# ─────────────────────────────────────────────
_SKIP = object()


def _to_jsonable(v):
    """แปลงค่าให้ JSON-safe แบบ recursive — คืน _SKIP ถ้า serialize ไม่ได้/ไม่ควรเก็บ
    (bytes เช่นรูปกราฟ/ไฟล์ report จะถูกข้าม เพราะ generate ใหม่ได้จากปุ่มในแอป)"""
    if v is None or isinstance(v, (str, int, float, bool)):
        return v
    if isinstance(v, bytes):
        return _SKIP
    if isinstance(v, np.integer):
        return int(v)
    if isinstance(v, np.floating):
        return float(v)
    if isinstance(v, np.ndarray):
        return v.tolist()
    if isinstance(v, pd.DataFrame):
        return {'__df__': True, 'records': v.to_dict('records')}
    if isinstance(v, dict):
        out = {}
        for k, vv in v.items():
            c = _to_jsonable(vv)
            if c is not _SKIP:
                out[str(k)] = c
        return out
    if isinstance(v, (list, tuple)):
        out = []
        for x in v:
            c = _to_jsonable(x)
            if c is not _SKIP:
                out.append(c)
        return out
    try:
        json.dumps(v)
        return v
    except (TypeError, ValueError):
        return _SKIP


def render():
    ss = st.session_state

    _just_loaded = ss.pop('_just_loaded_count', None)
    if _just_loaded is not None:
        st.success(f"✅ โหลดข้อมูลโปรเจกต์สำเร็จ — {_just_loaded} รายการ "
                   f"(ข้อมูลอัปเดตในทุกแท็บแล้ว)")

    st.markdown("### 💾 Project Save / Load")
    st.markdown("บันทึกและโหลดข้อมูลโปรเจกต์ทั้งหมดเป็นไฟล์ JSON")

    st.markdown("---")

    col_save, col_load = st.columns(2)

    # ════════════════════════════════
    #  SAVE
    # ════════════════════════════════
    with col_save:
        st.markdown("#### 📥 บันทึกโปรเจกต์")

        # แสดงสรุปข้อมูลที่มี
        with st.container(border=True):
            st.markdown("**ข้อมูลที่จะบันทึก:**")
            items = [
                ('esal_rigid',    '🚛 ESAL Rigid'),
                ('esal_flex',     '🚛 ESAL Flexible'),
                ('cbr_values',    '📊 CBR Analysis'),
                ('flex_results',  '🔧 Flexible Design'),
                ('rigid_results', '🏗️ Rigid Design'),
            ]
            for key, label in items:
                val = ss.get(key)
                has = _has_data(val)
                icon = '✅' if has else '—'
                color = '#2E7D32' if has else '#9E9E9E'
                st.markdown(
                    f'<div style="font-size:0.9rem;color:{color};padding:2px 0">'
                    f'{icon} {label}</div>',
                    unsafe_allow_html=True)

        if st.button("📥 สร้างไฟล์ Save", type="primary",
                      use_container_width=True, key="btn_save"):
            save_data = {}
            _skip_keys = {'json_uploader'}  # widget key ของ file_uploader เอง ไม่ต้อง save
            for key in list(ss.keys()):
                if key.startswith('_') or key in _skip_keys:
                    continue
                val = _to_jsonable(ss[key])
                if val is not _SKIP:
                    save_data[key] = val

            save_data['_saved_at']      = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            save_data['_version']       = 'ITM_Pave_Pro_v3'
            save_data['_project_name']  = ss.get('project_name', '')

            json_str = json.dumps(save_data, ensure_ascii=False, indent=2)
            proj     = ss.get('project_name', 'Project') or 'Project'
            fname    = f"{proj}_{datetime.now().strftime('%Y%m%d_%H%M')}.json"

            st.download_button(
                "💾 Download JSON",
                data=json_str.encode('utf-8'),
                file_name=fname,
                mime="application/json",
                use_container_width=True,
                key="dl_json")

            st.success(f"✅ พร้อม Download — {len(save_data)} รายการ")

    # ════════════════════════════════
    #  LOAD
    # ════════════════════════════════
    with col_load:
        st.markdown("#### 📤 โหลดโปรเจกต์")

        uploaded_json = st.file_uploader(
            "เลือกไฟล์ JSON", type=['json'], key="json_uploader")

        if uploaded_json:
            try:
                raw      = uploaded_json.read().decode('utf-8')
                data     = json.loads(raw)
                saved_at = data.get('_saved_at', 'ไม่ทราบ')
                proj_name = data.get('_project_name', '')
                version  = data.get('_version', '')

                with st.container(border=True):
                    st.markdown(f"**โปรเจกต์:** {proj_name or '—'}")
                    st.markdown(f"**บันทึกเมื่อ:** {saved_at}")
                    st.markdown(f"**Version:** {version}")

                    # แสดงสรุปข้อมูลในไฟล์
                    has_items = []
                    if data.get('esal_rigid') or data.get('esal_flex'):
                        has_items.append('🚛 ESAL')
                    if data.get('cbr_values'):
                        has_items.append('📊 CBR')
                    if data.get('flex_results'):
                        has_items.append('🔧 Flexible')
                    if data.get('rigid_results'):
                        has_items.append('🏗️ Rigid')
                    if has_items:
                        st.markdown("**มีข้อมูล:** " + " | ".join(has_items))

                if st.button("📤 โหลดข้อมูลนี้", type="primary",
                              use_container_width=True, key="btn_load"):
                    # เก็บข้อมูลไว้ใน session_state ชั่วคราวก่อน แล้ว rerun
                    # การ set ค่าจริงให้ widget keys (fmat_i, jpcp_name_i ฯลฯ)
                    # ต้องทำที่ app.py "ก่อน" st.tabs() render ในรอบถัดไป
                    # ไม่งั้นจะชนกฎ "ห้าม set widget key หลัง widget instantiate แล้ว"
                    # ของ Streamlit เพราะแท็บ 1-4 render ไปก่อนแท็บนี้ในทุกรอบ
                    ss['_pending_load_data'] = data
                    st.rerun()

            except Exception as e:
                st.error(f"❌ ไม่สามารถอ่านไฟล์ได้: {e}")

    # ════════════════════════════════
    #  Reset
    # ════════════════════════════════
    st.markdown("---")
    with st.expander("🗑️ ล้างข้อมูลทั้งหมด", expanded=False):
        st.warning("⚠️ จะลบข้อมูลทั้งหมดในเซสชันนี้ — ไม่สามารถย้อนกลับได้")
        if st.button("🗑️ ล้างข้อมูลทั้งหมด", type="primary", key="btn_reset"):
            # เดิมลบ key ตรงนี้เลย แต่แท็บ 1-4 render (สร้าง widget) ไปก่อนแท็บนี้
            # ในทุกรอบอยู่แล้ว (st.tabs() render ทุกแท็บพร้อมกัน) การ del widget key
            # ที่เพิ่ง instantiate ไปในรอบเดียวกันเสี่ยงชนกฎของ Streamlit เหมือนตอน Load
            # จึงพัก flag ไว้แล้วให้ app.py ล้างจริงก่อน render แท็บใดๆ ในรอบถัดไป
            ss['_pending_reset'] = True
            st.rerun()

    # ════════════════════════════════
    #  COMBINED REPORT
    # ════════════════════════════════
    st.markdown("---")
    st.markdown("#### 📄 สร้างรายงานรวม (Combined Report)")
    st.markdown("รวมทุกส่วนที่คำนวณแล้วเป็นไฟล์ Word เดียว โดยอัตโนมัติ")

    with st.container(border=True):

        # ── แสดงสถานะข้อมูลที่มี ──
        st.markdown("**ข้อมูลที่พร้อมใช้:**")
        report_items = [
            ('traffic_df',    '🚛 ESAL'),
            ('cbr_values',    '📊 CBR'),
            ('flex_results',  '🔧 Flexible Pavement'),
            ('rigid_results', '🏗️ Rigid Pavement'),
        ]
        has_any = False
        cols_status = st.columns(4)
        for ci, (key, label) in enumerate(report_items):
            val = ss.get(key)
            has = _has_data(val)
            if has:
                has_any = True
            with cols_status[ci]:
                if has:
                    st.success(label, icon="✅")
                else:
                    st.caption(f"— {label}")

        st.markdown("")

        # ── Upload รูปแผนที่สำรวจจราจร ──
        st.markdown("**🗺️ รูปแผนที่ตำแหน่งสำรวจจราจร** *(สำหรับ section N.2.1)*")
        col_up, col_prev = st.columns([2, 1])
        with col_up:
            survey_img_file = st.file_uploader(
                f"Upload รูปที่ N-1 (PNG / JPG)",
                type=['png', 'jpg', 'jpeg'],
                key="survey_map_uploader",
                help="รูปแผนที่แสดงตำแหน่งสำรวจปริมาณจราจร — แต่ละโครงการต่างกัน"
            )
            if survey_img_file is not None:
                ss['_survey_map_img'] = survey_img_file.read()
                st.success("✅ โหลดรูปแล้ว — จะแทรกในรายงาน section N.2.1")
            elif ss.get('_survey_map_img'):
                st.info("📌 ใช้รูปที่ upload ไว้ก่อนหน้า")
            else:
                st.caption("⚠️ ไม่มีรูป — จะแทรก placeholder แดงแทน (แก้ใน Word ภายหลัง)")
        with col_prev:
            if ss.get('_survey_map_img'):
                st.image(ss['_survey_map_img'], caption="Preview", use_container_width=True)

        st.markdown("")

        # ── Chapter input ──
        col_ch, col_preview = st.columns([1, 2])
        with col_ch:
            chapter = st.number_input(
                "บทที่ (Chapter)",
                min_value=1, max_value=9,
                value=int(ss.get('_combined_chapter', 3)),
                step=1,
                key="combined_chapter_input",
                help="เลขบทในรายงานจริง เช่น บทที่ 3 → หัวข้อ 3.2, 3.3, 3.4, 3.5"
            )
            ss['_combined_chapter'] = chapter

        with col_preview:
            c = str(chapter)
            st.markdown(f"""
**Preview หัวข้อ/ตาราง/รูป (บทที่ {chapter}):**

| ส่วน | Section | ตาราง | รูป |
|---|---|---|---|
| ESAL Flex | {c}.2.2 | {c}-1 → {c}-4 | {c}-1 |
| ESAL Rigid | {c}.2.3 | {c}-5 → {c}-7 | — |
| CBR | {c}.3 | {c}-7 | {c}-7 |
| Flexible | {c}.4 | {c}-8, {c}-9, {c}-10 | {c}-8 |
| Rigid | {c}.5 | — | {c}-4 → {c}-8 |
""")

        # ── ปุ่ม Generate ──
        if not has_any:
            st.warning("⚠️ ยังไม่มีข้อมูลการคำนวณ — กรุณาคำนวณอย่างน้อย 1 ส่วนก่อน")
        else:
            if st.button("📄 สร้างรายงานรวม", type="primary",
                          use_container_width=True, key="btn_combined_report"):
                _generate_combined_report(ss, chapter)

        if ss.get('_combined_report_bytes'):
            proj  = ss.get('project_name', '') or 'Report'
            fname = f"Combined_Report_Ch{chapter}_{proj}.docx"
            st.download_button(
                "📥 Download รายงานรวม (.docx)",
                data=ss['_combined_report_bytes'],
                file_name=fname,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="dl_combined_report")

    # Footer
    st.markdown("---")
    st.markdown("""
    <div style='text-align:center;color:#4A5568;font-size:0.85rem;padding:0.5rem;'>
        🛣️ <b>ITM Pave Pro v3.0</b> — AASHTO 1993 Pavement Design System<br>
        พัฒนาโดย รศ.ดร.อิทธิพล มีผล | ภาควิชาครุศาสตร์โยธา | คณะครุศาสตร์อุตสาหกรรม | มจพ.
    </div>
    """, unsafe_allow_html=True)
