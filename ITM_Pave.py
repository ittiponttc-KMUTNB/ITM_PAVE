# ╔══════════════════════════════════════════════════════════════════╗
# ║  ITM Pave Pro — AASHTO 1993 Pavement Design System             ║
# ║  Single-file Streamlit App | Engineering Green Theme            ║
# ║  พัฒนาโดย รศ.ดร.อิทธิพล มีผล | ภาควิชาครุศาสตร์โยธา มจพ.    ║
# ╠══════════════════════════════════════════════════════════════════╣
# ║  TAB 1 │ ESAL Calculator (Pt Global | Upload / Manual)        ║
# ║  TAB 2 │ CBR Analysis (Percentile)                             ║
# ║  TAB 3 │ Flexible Pavement Design                              ║
# ║  TAB 4 │ Rigid Design (K-Nomograph + JPCP/JRCP + CRCP)        ║
# ║  TAB 5 │ Report & Save (Word + JSON)       ปรับปรุง ครั้งที่ 2.1   ║
# ╚══════════════════════════════════════════════════════════════════╝

# ─────────────────────────────────────────────
#  SEC 1: IMPORTS
# ─────────────────────────────────────────────
import streamlit as st
import pandas as pd
import numpy as np
import math
import json
import io
import base64
from datetime import datetime

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as patches
MPL_OK = True

import plotly.graph_objects as go
PLOTLY_OK = True

try:
    from scipy.optimize import brentq as _brentq
except ImportError:
    def _brentq(f, a, b, xtol=1e-6, maxiter=500):
        fa, fb = f(a), f(b)
        if fa * fb > 0:
            raise ValueError("No sign change")
        for _ in range(maxiter):
            mid = (a + b) / 2.0
            fm = f(mid)
            if abs(fm) < xtol or (b - a) / 2.0 < xtol:
                return mid
            if fa * fm < 0:
                b, fb = mid, fm
            else:
                a, fa = mid, fm
        return (a + b) / 2.0

# openpyxl — required for pd.read_excel
try:
    import openpyxl
    from openpyxl import load_workbook
    OPENPYXL_OK = True
except ModuleNotFoundError:
    OPENPYXL_OK = False
    openpyxl = None
    load_workbook = None

from docx import Document as DocxDoc
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
DOCX_OK = True

try:
    from PIL import Image as PILImage, ImageDraw as PILDraw
    PIL_OK = True
except ImportError:
    PILImage = None
    PILDraw  = None
    PIL_OK   = False

# ─────────────────────────────────────────────
#  SEC 2: PAGE CONFIG
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="ITM Pave Pro – AASHTO 1993",
    page_icon="🛣️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────
#  SEC 3: CSS — Engineering Green Theme
# ─────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Sarabun:wght@300;400;600;700&family=IBM+Plex+Mono:wght@400;600&display=swap');

html, body, [class*="css"] {
    font-family: 'Sarabun', sans-serif;
}

/* ── Header ── */
.main-header {
    background: linear-gradient(135deg, #1B5E20 0%, #388E3C 60%, #43A047 100%);
    color: white;
    padding: 1.4rem 2rem;
    border-radius: 14px;
    margin-bottom: 1.5rem;
    box-shadow: 0 6px 20px rgba(27,94,32,0.35);
    border-left: 6px solid #A5D6A7;
}
.main-header h1 { margin:0; font-size:1.75rem; font-weight:700; letter-spacing:-0.5px; }
.main-header p  { margin:0.3rem 0 0; font-size:0.9rem; opacity:0.88; }

/* ── Section Cards ── */
.card {
    background: #fff;
    border: 1px solid #C8E6C9;
    border-left: 5px solid #2E7D32;
    border-radius: 10px;
    padding: 1rem 1.3rem;
    margin-bottom: 1rem;
    box-shadow: 0 2px 8px rgba(46,125,50,0.08);
}
.card h4 { color:#1B5E20; margin:0 0 0.8rem; font-size:1rem; font-weight:700; }

/* ── Status Badges ── */
.badge-ready  { background:#E8F5E9; color:#2E7D32; border:1px solid #A5D6A7;
                border-radius:20px; padding:0.25rem 0.8rem; font-size:0.82rem; font-weight:600; display:inline-block; }
.badge-wait   { background:#FFF8E1; color:#E65100; border:1px solid #FFE082;
                border-radius:20px; padding:0.25rem 0.8rem; font-size:0.82rem; font-weight:600; display:inline-block; }
.badge-na     { background:#F5F5F5; color:#757575; border:1px solid #E0E0E0;
                border-radius:20px; padding:0.25rem 0.8rem; font-size:0.82rem; font-weight:600; display:inline-block; }

/* ── Result Boxes ── */
.result-pass { background:#E8F5E9; border:1px solid #A5D6A7; border-radius:8px;
               padding:0.8rem 1rem; color:#1B5E20; font-weight:600; margin:0.3rem 0; }
.result-fail { background:#FFEBEE; border:1px solid #EF9A9A; border-radius:8px;
               padding:0.8rem 1rem; color:#B71C1C; font-weight:600; margin:0.3rem 0; }
.result-info { background:#E3F2FD; border:1px solid #90CAF9; border-radius:8px;
               padding:0.8rem 1rem; color:#0D47A1; font-weight:600; margin:0.3rem 0; }
.result-warn { background:#FFF8E1; border:1px solid #FFE082; border-radius:8px;
               padding:0.8rem 1rem; color:#E65100; font-weight:600; margin:0.3rem 0; }

/* ── Metric Box ── */
.metric-box {
    background:#fff; border:1px solid #C8E6C9; border-radius:12px;
    padding:1rem; text-align:center;
    box-shadow:0 2px 8px rgba(46,125,50,0.10);
}
.metric-box .val { font-size:1.5rem; font-weight:700; color:#1B5E20;
                   font-family:'IBM Plex Mono', monospace; }
.metric-box .lbl { font-size:0.78rem; color:#558B2F; margin-top:0.2rem; }

/* ── Tabs ── */
[data-baseweb="tab-list"] { gap:3px; }
[data-baseweb="tab"] {
    background:#E8F5E9 !important; border-radius:8px 8px 0 0 !important;
    font-weight:600 !important; color:#1B5E20 !important;
    padding:0.45rem 0.9rem !important;
}
[aria-selected="true"][data-baseweb="tab"] {
    background:#2E7D32 !important; color:white !important;
}

/* ── Sidebar ── */
[data-testid="stSidebar"] { background:#1B5E20; }
[data-testid="stSidebar"] * { color:#E8F5E9 !important; }
[data-testid="stSidebar"] hr { border-color:#2E7D32; }

/* ── Buttons ── */
button[kind="primary"] {
    background:#2E7D32 !important; border-radius:8px !important;
    font-weight:700 !important;
}

/* ── Number inputs ── */
.stNumberInput > div > div > input {
    font-family:'IBM Plex Mono', monospace; font-weight:600;
}

/* ── DataFrames ── */
.stDataFrame { border-radius:8px; overflow:hidden; }

/* ── Flow Arrow ── */
.flow-arrow {
    text-align:center; font-size:1.5rem; color:#43A047;
    margin:0.3rem 0; line-height:1;
}
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
#  SEC 4: CONSTANTS & LOOKUP TABLES
# ─────────────────────────────────────────────
TON_TO_KIP = 2.2046

VEHICLE_AXLES = {
    "MB":  [(4,  1, 1), (11, 1, 1)],
    "HB":  [(5,  1, 1), (20, 2, 1)],
    "MT":  [(4,  1, 1), (11, 1, 1)],
    "HT":  [(5,  1, 1), (20, 2, 1)],
    "TR":  [(5,  1, 1), (20, 2, 1), (11, 1, 2)],
    "STR": [(5,  1, 1), (20, 2, 2)],
}
VEHICLE_LABELS = {
    "MB": "Medium Bus",  "HB": "Heavy Bus",
    "MT": "Medium Truck","HT": "Heavy Truck",
    "TR": "Trailer",     "STR": "Semi Trailer",
}
VEHICLE_COLS = ["MB", "HB", "MT", "HT", "TR", "STR"]

# SLAB_THICKNESSES: cm ตามมาตรฐานไทย (กรมทางหลวง)
# แสดงผลพร้อมค่าอ้างอิงนิ้วกลม สำหรับเทียบตาราง AASHTO
SLAB_THICKNESSES = [25, 28, 30, 32, 35]          # cm (ใช้คำนวณจริง)
SLAB_LABELS      = ['25 cm (10 in)', '28 cm (11 in)', '30 cm (12 in)',
                     '32 cm (13 in)', '35 cm (14 in)']  # แสดงใน UI/รายงาน
SLAB_INCH        = [10, 11, 12, 13, 14]          # reference เท่านั้น
SN_DEFAULTS      = [6.5, 7.1, 7.5, 8.0]

ZR_MAP = {
    50:0.000,60:-0.253,70:-0.524,75:-0.674,80:-0.841,85:-1.037,
    90:-1.282,91:-1.340,92:-1.405,93:-1.476,94:-1.555,95:-1.645,
    96:-1.751,97:-1.881,98:-2.054,99:-2.327,
}

J_VALUES = {"JPCP/JRCP": 2.8, "CRCP": 2.6}

# วัสดุที่ lock mi=1.0 (AC และ PMA)
AC_MATERIALS_LOCK_MI = {
    "ผิวทางลาดยาง PMA",
    "ผิวทางแอสฟัลต์คอนกรีต (AC)",
}

# วัสดุ AC/PMA สำหรับ sub-layer checkbox
AC_SURFACE_MATERIALS = {
    "ผิวทางลาดยาง PMA",
    "ผิวทางแอสฟัลต์คอนกรีต (AC)",
}

FLEX_LAYER_MATERIALS = {
    "ไม่เลือก":                                              (None, None),
    "ผิวทางลาดยาง PMA":                                      (0.44, 1.0),
    "ผิวทางแอสฟัลต์คอนกรีต (AC)":                           (0.40, 1.0),
    "(CTB) หินคลุกปรับปรุงด้วยปูนซีเมนต์ UCS 40 ksc ": (0.18, 1.0),
    "หินคลุกผสมซีเมนต์ UCS 24.5 ksc":               (0.15, 1.0),
    "ดินซีเมนต์ UCS 17.5 ksc":                       (0.13, 1.0),
    "หินคลุก CBR 80%":                               (0.13, 1.0),
    "วัสดุหมุนเวียน (Recycling)":                    (0.15, 1.0),
    "วัสดุมวลรวม CBR 25%":                        (0.10, 1.0),
    "วัสดุคัดเลือก ก":                                       (0.08, 1.0),
    "ดินถมคันทาง CBR 10%":                                   (0.08, 1.0),
}

RIGID_LAYER_MATERIALS = {
    "ไม่เลือก":                                               None,
    "AC รองใต้ผิวคอนกรีต":                                   2500,
    "วัสดุหมุนเวียน (Recycling)":                     850,
    "(CTB) หินคลุกปรับปรุงด้วยปูนซีเมนต์ UCS 40 ksc ": 1200,
    "หินคลุกผสมซีเมนต์ UCS 24.5 ksc":                850,
    "ดินซีเมนต์ UCS 17.5 ksc":                        350,
    "หินคลุก CBR 80%":                                350,
    "รองพื้นทางวัสดุมวลรวม CBR 25%":                         150,
    "วัสดุคัดเลือก ก":                                        100,
    "ดินถมคันทาง CBR 10%":                                    100,
}

# E_MPa default ต่อวัสดุ (สำหรับ auto-fill Layer Editor)
# NOTE: key ต้องตรงกับ RIGID_LAYER_MATERIALS ทุกตัว
RIGID_LAYER_E_DEFAULT = {
    "ไม่เลือก":                                               0,
    "AC รองใต้ผิวคอนกรีต":                                   2500,
    "วัสดุหมุนเวียน (Recycling)":                             850,
    "(CTB) หินคลุกปรับปรุงด้วยปูนซีเมนต์ UCS 40 ksc ":      1200,
    "หินคลุกผสมซีเมนต์ UCS 24.5 ksc":                        850,
    "ดินซีเมนต์ UCS 17.5 ksc":                                350,
    "หินคลุก CBR 80%":                                        350,
    "รองพื้นทางวัสดุมวลรวม CBR 25%":                         150,
    "วัสดุคัดเลือก ก":                                        100,
    "ดินถมคันทาง CBR 10%":                                    100,
}

SAMPLE_CBR = [14.8,14.37,5.31,17.37,5.48,18.46,4.85,6.23,
              5.02,10.78,10.52,14,15.5,8.7,12.93,8.19,
              8.1,15.56,16.88,20.75,20.3,8,7.84,7.48,
              23.55,8.92,13.3,13.5,13.86,7.18,6.95,5.8,
              6,11.18,9.69,7.48]

# ─────────────────────────────────────────────
#  SEC 5: ENGINE FUNCTIONS
# ─────────────────────────────────────────────

def ealf_flex(L1_ton, L2, SN, Pt):
    L1  = L1_ton * TON_TO_KIP
    Gt  = math.log10((4.2 - Pt) / (4.2 - 1.5))
    Bx  = 0.40 + 0.081*(L1+L2)**3.23 / ((SN+1)**5.19 * L2**3.23)
    B18 = 0.40 + 0.081*(18+1)**3.23  / ((SN+1)**5.19 * 1.0**3.23)
    return 10**(4.79*math.log10(L1+L2) - 4.33*math.log10(L2)
                - 4.79*math.log10(19) + Gt*(1/B18 - 1/Bx))

def ealf_rigid(L1_ton, L2, D_cm, Pt):
    L1  = L1_ton * TON_TO_KIP
    D   = round(D_cm / 2.54)   # AASHTO 1993 ใช้ความหนาเป็นจำนวนเต็มนิ้ว
    Gt  = math.log10((4.5 - Pt) / (4.5 - 1.5))
    Bx  = 1.0 + 3.63*(L1+L2)**5.20 / ((D+1)**8.46 * L2**3.52)
    B18 = 1.0 + 3.63*(18+1)**5.20  / ((D+1)**8.46 * 1.0**3.52)
    return 10**(4.62*math.log10(L1+L2) - 3.28*math.log10(L2)
                - 4.62*math.log10(19) + Gt*(1/B18 - 1/Bx))

def truck_factor_flex(vtype, SN, Pt):
    return sum(ealf_flex(L1,L2,SN,Pt)*cnt for L1,L2,cnt in VEHICLE_AXLES[vtype])

def truck_factor_rigid(vtype, D_cm, Pt):
    return sum(ealf_rigid(L1,L2,D_cm,Pt)*cnt for L1,L2,cnt in VEHICLE_AXLES[vtype])

def compute_esal_from_df(traffic_df, ldf, ddf, Pt, mode="rigid", sn_list=None):
    """
    คำนวณ ESAL จาก traffic DataFrame (AASHTO 1993)

    สูตร: ESAL = Σ_ปี [ AADT(คัน/วัน) × 365 × DDF × LDF × TF ]

    traffic_df : DataFrame คอลัมน์ Year, MB, HB, MT, HT, TR, STR
                 ค่า = AADT 2 ทิศทาง (คัน/วัน) ต่อปี
    ldf        : Lane Distribution Factor
    ddf        : Directional Distribution Factor
    Pt         : Terminal Serviceability
    Returns    : {D_cm: esal} for rigid | {SN: esal} for flexible
    """
    DAYS_PER_YEAR = 365
    if mode == "rigid":
        keys    = SLAB_THICKNESSES
        results = {k: 0.0 for k in keys}
        for _, row in traffic_df.iterrows():
            for vtype in VEHICLE_COLS:
                cnt = float(row.get(vtype, 0) or 0)
                if cnt <= 0:
                    continue
                for D in keys:
                    tf_val = truck_factor_rigid(vtype, D, Pt)
                    results[D] += cnt * DAYS_PER_YEAR * ddf * ldf * tf_val
        return results
    else:
        keys    = sn_list or SN_DEFAULTS
        results = {k: 0.0 for k in keys}
        for _, row in traffic_df.iterrows():
            for vtype in VEHICLE_COLS:
                cnt = float(row.get(vtype, 0) or 0)
                if cnt <= 0:
                    continue
                for SN in keys:
                    tf_val = truck_factor_flex(vtype, SN, Pt)
                    results[SN] += cnt * DAYS_PER_YEAR * ddf * ldf * tf_val
        return results

def aashto_sn_required(esal, zr, so, pi, pt, mr_psi):
    delta_psi = pi - pt
    logW18 = math.log10(max(esal, 1))
    def eq(SN):
        if SN <= 0: return -1e10
        t1 = zr * so
        t2 = 9.36*math.log10(SN+1) - 0.20
        t3 = math.log10(delta_psi/2.7) / (0.40 + 1094/(SN+1)**5.19)
        t4 = 2.32*math.log10(mr_psi) - 8.07
        return t1+t2+t3+t4 - logW18
    try:
        return _brentq(eq, 0.1, 30, xtol=1e-4)
    except:
        return None

def aashto_rigid_w18(d_cm, pi, pt, zr, so, sc_psi, cd, j, ec_psi, k_pci):
    d_in = round(d_cm / 2.54)   # AASHTO 1993 ใช้ความหนาเป็นจำนวนเต็มนิ้ว
    delta_psi = pi - pt
    t1 = zr * so
    t2 = 7.35*math.log10(d_in+1) - 0.06
    t3 = math.log10(delta_psi/3.0) / (1 + 1.624e7/(d_in+1)**8.46)
    num4 = sc_psi * cd * (d_in**0.75 - 1.132)
    den4 = 215.63 * j * (d_in**0.75 - 18.42/(ec_psi/k_pci)**0.25)
    if num4 <= 0 or den4 <= 0:
        return None
    inner = num4 / den4
    if inner <= 0:
        return None
    t4 = (4.22 - 0.32*pt) * math.log10(inner)
    return 10**(t1+t2+t3+t4)

def cbr_to_mr(cbr):
    return 1500.0 * cbr

def mr_to_k(mr_psi):
    return mr_psi / 19.4

def calc_percentile_cbr(cbr_values):
    arr = np.sort(np.array(cbr_values, dtype=float))
    n   = len(arr)
    unique_cbr = np.unique(arr)
    unique_pct = np.array([np.sum(arr >= v)/n*100 for v in unique_cbr])
    return arr, n, unique_cbr, unique_pct

def grow_traffic(base_row, growth_rate_pct, years):
    """Generate DataFrame with yearly traffic from base year + growth rate."""
    r = growth_rate_pct / 100.0
    rows = []
    for y in range(1, years+1):
        factor = (1 + r)**(y - 1)
        row = {"Year": y}
        for v in VEHICLE_COLS:
            row[v] = int(round(base_row.get(v, 0) * factor))
        rows.append(row)
    return pd.DataFrame(rows)

# ─────────────────────────────────────────────
#  SEC 6: NOMOGRAPH DRAW FUNCTIONS (Matplotlib)
# ─────────────────────────────────────────────

def draw_k_infinity_nomograph(esb_psi, dsb_in, k_sub_pci):
    """
    Composite k∞ Nomograph (AASHTO 1993 Fig. 3.3 approximation)
    3 axes: Esb (left), DSB (center), k∞ (right)
    """
    fig, ax = plt.subplots(figsize=(8, 9))
    ax.set_xlim(0, 10); ax.set_ylim(0, 10)
    ax.axis('off')
    ax.set_facecolor('#F1F8E9')
    fig.patch.set_facecolor('#F1F8E9')

    # Axis x-positions
    x_esb, x_dsb, x_kinf = 1.5, 5.0, 8.5

    # ── Axis lines ──
    for x in [x_esb, x_dsb, x_kinf]:
        ax.plot([x, x], [0.5, 9.5], color='#1B5E20', lw=2.5)

    # ── Esb axis (left): 5,000 – 100,000 psi ──
    esb_range = [5000, 10000, 20000, 30000, 50000, 100000]
    esb_log_min = math.log10(5000)
    esb_log_max = math.log10(100000)
    def esb_to_y(v):
        return 0.5 + 9.0*(math.log10(v)-esb_log_min)/(esb_log_max-esb_log_min)
    for v in esb_range:
        y = esb_to_y(v)
        ax.plot([x_esb-0.15, x_esb+0.15], [y, y], color='#1B5E20', lw=1.5)
        ax.text(x_esb-0.25, y, f"{v:,}", ha='right', va='center', fontsize=8, color='#1B5E20')
    ax.text(x_esb, 9.8, "Esb (psi)", ha='center', va='bottom', fontsize=9, fontweight='bold', color='#1B5E20')

    # ── DSB axis (center): 0 – 36 in ──
    dsb_range = [0, 4, 8, 12, 16, 20, 24, 28, 32, 36]
    def dsb_to_y(v):
        return 0.5 + 9.0*(v/36.0)
    for v in dsb_range:
        y = dsb_to_y(v)
        ax.plot([x_dsb-0.15, x_dsb+0.15], [y, y], color='#2E7D32', lw=1.5)
        ax.text(x_dsb+0.25, y, f"{v}", ha='left', va='center', fontsize=8, color='#2E7D32')
    ax.text(x_dsb, 9.8, "DSB (in)", ha='center', va='bottom', fontsize=9, fontweight='bold', color='#2E7D32')

    # ── k∞ axis (right): 50 – 1000 pci ──
    kinf_range = [50, 100, 150, 200, 300, 500, 700, 1000]
    kinf_log_min = math.log10(50)
    kinf_log_max = math.log10(1000)
    def kinf_to_y(v):
        return 0.5 + 9.0*(math.log10(v)-kinf_log_min)/(kinf_log_max-kinf_log_min)
    for v in kinf_range:
        y = kinf_to_y(v)
        ax.plot([x_kinf-0.15, x_kinf+0.15], [y, y], color='#388E3C', lw=1.5)
        ax.text(x_kinf+0.25, y, f"{v}", ha='left', va='center', fontsize=8, color='#388E3C')
    ax.text(x_kinf, 9.8, "k∞ (pci)", ha='center', va='bottom', fontsize=9, fontweight='bold', color='#388E3C')

    # ── Compute k∞ from inputs via AASHTO approximation ──
    # Odemark: k∞ = k_sub * (1 + Esb*DSB/(k_sub*19.4*DSB^0.5 + 1))^0.4  (simplified)
    if esb_psi > 0 and dsb_in >= 0 and k_sub_pci > 0:
        if dsb_in == 0:
            k_inf_calc = k_sub_pci
        else:
            h_eq = dsb_in * (esb_psi / (k_sub_pci * 19.4))**(1/3)
            k_inf_calc = min(k_sub_pci * (1 + 0.55*h_eq**0.45), 1000)
        k_inf_calc = max(50, min(1000, k_inf_calc))
    else:
        k_inf_calc = k_sub_pci

    # ── Draw reading lines ──
    y_esb  = esb_to_y(max(5000, min(100000, esb_psi)))
    y_dsb  = dsb_to_y(max(0, min(36, dsb_in)))
    y_kinf = kinf_to_y(max(50, min(1000, k_inf_calc)))

    ax.annotate("", xy=(x_dsb, y_dsb), xytext=(x_esb, y_esb),
                arrowprops=dict(arrowstyle="-", color='red', lw=2, linestyle='dashed'))
    ax.annotate("", xy=(x_kinf, y_kinf), xytext=(x_dsb, y_dsb),
                arrowprops=dict(arrowstyle="->", color='red', lw=2, linestyle='dashed'))

    # Dots at reading points
    for (xp, yp) in [(x_esb, y_esb), (x_dsb, y_dsb), (x_kinf, y_kinf)]:
        ax.plot(xp, yp, 'o', color='red', markersize=8, zorder=5)

    # k∞ result label
    ax.text(x_kinf+1.0, y_kinf, f"k∞ = {k_inf_calc:.0f} pci",
            ha='left', va='center', fontsize=11, fontweight='bold',
            color='red', bbox=dict(boxstyle='round,pad=0.3', facecolor='white', edgecolor='red', alpha=0.9))

    ax.set_title("Composite k∞ Nomograph (AASHTO 1993 Fig.3.3)", fontsize=11,
                 fontweight='bold', color='#1B5E20', pad=15)
    plt.tight_layout()
    return fig, k_inf_calc


def draw_loss_of_support_nomograph(k_inf_pci, ls_value):
    """
    Loss of Support Nomograph (AASHTO 1993 Fig. 3.7)
    k_corrected = k_inf / 10^(LS*0.5) approximation
    """
    fig, ax = plt.subplots(figsize=(7, 8))
    ax.set_xlim(0, 10); ax.set_ylim(0, 10)
    ax.axis('off')
    ax.set_facecolor('#F1F8E9')
    fig.patch.set_facecolor('#F1F8E9')

    x_kinf, x_kcorr = 2.5, 7.5
    ls_colors = {0.0:'#1B5E20', 0.5:'#2E7D32', 1.0:'#43A047',
                 1.5:'#66BB6A', 2.0:'#EF6C00', 3.0:'#B71C1C'}

    k_range = [10, 20, 50, 100, 200, 300, 500, 700, 1000, 1500, 2000, 3000]
    k_log_min = math.log10(10)
    k_log_max = math.log10(3000)
    def k_to_y(v):
        return 0.5 + 9.0*(math.log10(max(v,10))-k_log_min)/(k_log_max-k_log_min)

    for x in [x_kinf, x_kcorr]:
        ax.plot([x,x],[0.5,9.5], color='#1B5E20', lw=2.5)

    for v in k_range:
        for x in [x_kinf, x_kcorr]:
            y = k_to_y(v)
            ax.plot([x-0.15, x+0.15],[y,y], color='#1B5E20', lw=1.5)
        ax.text(x_kinf-0.25, k_to_y(v), f"{v}", ha='right', va='center', fontsize=8, color='#1B5E20')
        ax.text(x_kcorr+0.25, k_to_y(v), f"{v}", ha='left', va='center', fontsize=8, color='#388E3C')

    ax.text(x_kinf, 9.8, "k∞ (pci)", ha='center', va='bottom', fontsize=9, fontweight='bold', color='#1B5E20')
    ax.text(x_kcorr, 9.8, "k_eff (pci)", ha='center', va='bottom', fontsize=9, fontweight='bold', color='#388E3C')

    # Draw LS lines for all LS values
    for ls, lc in ls_colors.items():
        for k_val in [20, 50, 100, 200, 500, 1000, 2000]:
            k_corr_ls = k_val / (10**(ls * 0.5))
            k_corr_ls = max(10, min(3000, k_corr_ls))
            y1 = k_to_y(k_val)
            y2 = k_to_y(k_corr_ls)
            ax.plot([x_kinf, x_kcorr],[y1, y2], color=lc, lw=0.8, alpha=0.4)
        ax.text(5.0, k_to_y(50/(10**(ls*0.5))) + ls*0.3,
                f"LS={ls}", ha='center', va='center', fontsize=7, color=lc, alpha=0.8)

    # ── Compute k_corrected ──
    k_corr_calc = k_inf_pci / (10**(ls_value * 0.5))
    k_corr_calc = max(10, min(3000, k_corr_calc))

    # Draw user's reading line
    y1 = k_to_y(max(10, min(3000, k_inf_pci)))
    y2 = k_to_y(k_corr_calc)
    ax.annotate("", xy=(x_kcorr, y2), xytext=(x_kinf, y1),
                arrowprops=dict(arrowstyle="->", color='red', lw=2.5))
    for (xp, yp) in [(x_kinf, y1), (x_kcorr, y2)]:
        ax.plot(xp, yp, 'o', color='red', markersize=9, zorder=5)

    ax.text(x_kcorr+1.2, y2,
            f"k_eff =\n{k_corr_calc:.0f} pci",
            ha='left', va='center', fontsize=10, fontweight='bold',
            color='red', bbox=dict(boxstyle='round,pad=0.3', facecolor='white', edgecolor='red', alpha=0.9))

    ax.set_title(f"Loss of Support Nomograph  (LS = {ls_value})", fontsize=11,
                 fontweight='bold', color='#1B5E20', pad=15)

    # Legend
    legend_x, legend_y = 3.5, 1.5
    for i, (ls, lc) in enumerate(ls_colors.items()):
        ax.plot(legend_x, legend_y - i*0.3, 's', color=lc, markersize=7)
        ax.text(legend_x+0.2, legend_y - i*0.3, f"LS = {ls}", va='center', fontsize=7, color=lc)

    plt.tight_layout()
    return fig, k_corr_calc

def fig_to_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor=fig.get_facecolor())
    buf.seek(0)
    return buf.read()

# ─────────────────────────────────────────────
#  SEC 6b: PAVEMENT STRUCTURE FIGURE
# ─────────────────────────────────────────────

# แปลงชื่อวัสดุภาษาไทย → อังกฤษสำหรับแสดงในรูป
_LAYER_NAME_EN = {
    "ผิวทางลาดยาง PMA":                                    "PMA Surface",
    "ผิวทางแอสฟัลต์คอนกรีต (AC)":                         "AC Surface",
    "(CTB) หินคลุกปรับปรุงด้วยปูนซีเมนต์ UCS 40 ksc ":   "Cement Treated Base",
    "หินคลุกผสมซีเมนต์ UCS 24.5 ksc":                     "MOD. Crushed Rock",
    "ดินซีเมนต์ UCS 17.5 ksc":                             "Soil Cement",
    "หินคลุก CBR 80%":                                     "Crushed Rock (CBR 80%)",
    "วัสดุหมุนเวียน (Recycling)":                          "Recycled Material",
    "วัสดุมวลรวม CBR 25%":                                 "Aggregate Subbase",
    "รองพื้นทางวัสดุมวลรวม CBR 25%":                      "Aggregate Subbase",
    "วัสดุคัดเลือก ก":                                     "Selected A",
    "AC รองใต้ผิวคอนกรีต":                                 "AC Interlayer",
    "ดินถมคันทาง CBR 10%":                                 "Embankment",
    "ดินถมคันทาง / ดินเดิม":                               "Embankment",
    "Concrete Slab":                                        "Concrete Slab",
}

# สีแต่ละวัสดุ
_LAYER_COLORS = {
    "ผิวทางลาดยาง PMA":                                    "#1A252F",
    "ผิวทางแอสฟัลต์คอนกรีต (AC)":                         "#2C3E50",
    "Concrete Slab":                                        "#78909C",
    "(CTB) หินคลุกปรับปรุงด้วยปูนซีเมนต์ UCS 40 ksc ":   "#7F8C8D",
    "หินคลุกผสมซีเมนต์ UCS 24.5 ksc":                     "#95A5A6",
    "ดินซีเมนต์ UCS 17.5 ksc":                             "#AAB7B8",
    "หินคลุก CBR 80%":                                     "#BDC3C7",
    "วัสดุหมุนเวียน (Recycling)":                          "#85929E",
    "วัสดุมวลรวม CBR 25%":                                 "#FFCC99",
    "รองพื้นทางวัสดุมวลรวม CBR 25%":                      "#FFCC99",
    "วัสดุคัดเลือก ก":                                     "#E8DAEF",
    "AC รองใต้ผิวคอนกรีต":                                 "#34495E",
    "ดินถมคันทาง CBR 10%":                                 "#F5CBA7",
    "ดินถมคันทาง / ดินเดิม":                               "#F5CBA7",
}

# วัสดุที่ใช้ text สีขาว (background เข้ม)
_DARK_LAYERS = {
    "ผิวทางลาดยาง PMA", "ผิวทางแอสฟัลต์คอนกรีต (AC)",
    "Concrete Slab", "AC รองใต้ผิวคอนกรีต",
    "(CTB) หินคลุกปรับปรุงด้วยปูนซีเมนต์ UCS 40 ksc ",
    "หินคลุกผสมซีเมนต์ UCS 24.5 ksc", "วัสดุหมุนเวียน (Recycling)",
}

def draw_pavement_structure(layers, mode="flex",
                             cbr_subgrade=3.0,
                             d_concrete_cm=None,
                             ptype="JPCP"):
    """
    วาดรูปโครงสร้างชั้นทาง (matplotlib) — ชื่อวัสดุภาษาอังกฤษ

    Parameters
    ----------
    layers        : list of dict
    mode          : "flex" | "rigid"
    cbr_subgrade  : CBR ดินเดิม (%)
    d_concrete_cm : ความหนา Slab (rigid)
    ptype         : "JPCP" | "JRCP" | "CRCP"
    """
    MIN_H   = 5      # ความสูงขั้นต่ำ (display units)
    W       = 3.0    # ความกว้าง block
    X_CTR   = 5.0    # กึ่งกลาง X
    X_START = X_CTR - W / 2

    # ── เตรียม layer list ──
    all_layers = []
    if mode == "rigid" and d_concrete_cm:
        all_layers.append({
            "name":        "Concrete Slab",
            "thickness_cm": d_concrete_cm,
            "label":       f"Concrete Slab ({ptype})",
            "side_info":   None,
        })

    valid = [l for l in layers if l.get("thickness_cm", 0) > 0]
    for l in valid:
        name  = l.get("name", "")
        h     = l.get("thickness_cm", 0)
        en    = _LAYER_NAME_EN.get(name, name)
        if mode == "flex":
            ai  = l.get("ai",  None)
            sni = l.get("sni", None)
            side = f"ai={ai:.2f} | SNi={sni:.3f}" if ai and sni else None
        else:
            e_mpa = l.get("E_MPa", None)
            side  = f"E={e_mpa:,} MPa" if e_mpa else None
        all_layers.append({"name": name, "thickness_cm": h,
                            "label": en, "side_info": side})

    # Subgrade
    all_layers.append({
        "name":        "ดินถมคันทาง / ดินเดิม",
        "thickness_cm": 0,
        "label":       f"Subgrade (CBR≥{cbr_subgrade:.0f}%)",
        "side_info":   None,
    })

    if len(all_layers) <= 1:
        return None

    # ── display height — normalize ──
    n_layers  = len(all_layers)
    real_h    = [l["thickness_cm"] for l in all_layers]
    max_real  = max((h for h in real_h if h > 0), default=30)
    # scale ให้ layer ใหญ่สุด = 40 display units, ขั้นต่ำ MIN_H
    SCALE     = 40.0 / max_real
    display_h = []
    for h in real_h:
        display_h.append(MIN_H if h == 0 else max(round(h * SCALE, 1), MIN_H))

    total_disp  = sum(display_h)
    total_thick = sum(h for h in real_h if h > 0)

    # ── figsize: กว้างคงที่ สูงตาม layer count ──
    fig_h = max(4.0, min(n_layers * 1.4, 8.0))
    fig_w = 9.0
    fs_lbl = max(10.0, 12.0 - n_layers * 0.3)   # fontsize ชื่อวัสดุ
    fs_h   = max(10.0, 12.0 - n_layers * 0.3)   # fontsize ความหนา

    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    ax.set_xlim(0, 11)
    ax.set_ylim(-4, total_disp + 5)
    ax.axis('off')
    fig.patch.set_facecolor('white')

    y = total_disp

    for i, layer in enumerate(all_layers):
        dh    = display_h[i]
        name  = layer["name"]
        label = layer["label"]
        h_cm  = layer["thickness_cm"]
        side  = layer["side_info"]

        color   = _LAYER_COLORS.get(name, "#D5D8DC")
        hatch   = '///' if "หมุนเวียน" in name else None
        is_dark = name in _DARK_LAYERS
        txt_col = 'white' if is_dark else '#1A1A1A'

        y_bot = y - dh
        y_ctr = y_bot + dh / 2

        # rectangle
        rect = patches.Rectangle(
            (X_START, y_bot), W, dh,
            linewidth=1.5, edgecolor='#2C3E50',
            facecolor=color, hatch=hatch, zorder=2
        )
        ax.add_patch(rect)

        # ความหนาในกล่อง
        h_text = f"{h_cm} cm" if h_cm > 0 else "∞"
        ax.text(X_CTR, y_ctr, h_text,
                ha='center', va='center',
                fontsize=fs_h, fontweight='bold',
                color=txt_col, zorder=3)

        # ชื่อวัสดุซ้าย
        ax.text(X_START - 0.2, y_ctr, label,
                ha='right', va='center',
                fontsize=fs_lbl, fontweight='bold',
                color='#1B2631', zorder=3)

        # ข้อมูลขวา
        if side:
            ax.text(X_START + W + 0.2, y_ctr, side,
                    ha='left', va='center',
                    fontsize=max(fs_lbl - 1, 7.0),
                    color='#154360', zorder=3)

        # เส้นคั่น
        if i > 0:
            ax.plot([X_START, X_START + W], [y, y],
                    color='#2C3E50', lw=1.0, zorder=3)

        y = y_bot

    # ── ลูกศร Total ──
    x_arr       = X_START + W + 2.8
    y_solid     = sum(display_h[:-1])   # ไม่รวม subgrade
    ax.annotate('', xy=(x_arr, total_disp),
                xytext=(x_arr, total_disp - y_solid),
                arrowprops=dict(arrowstyle='<->', color='#C0392B', lw=1.8))
    ax.text(x_arr + 0.2, total_disp - y_solid / 2,
            f"Total\n{total_thick} cm",
            ha='left', va='center',
            fontsize=9, color='#C0392B', fontweight='bold')

    plt.tight_layout()
    return fig


# ─────────────────────────────────────────────
#  SEC 7: WORD REPORT FUNCTIONS
# ─────────────────────────────────────────────

def _new_doc():
    if not DOCX_OK:
        return None
    doc = DocxDoc()
    style = doc.styles['Normal']
    style.font.name = 'TH SarabunPSK'
    style.font.size = Pt(15)
    try:
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')
    except:
        pass
    return doc

def _th_run(para, text, bold=False, size=15, color=None):
    run = para.add_run(text)
    run.font.name = 'TH SarabunPSK'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    try:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')
    except:
        pass
    return run

def _set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _doc_to_bytes(doc):
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def _add_section_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = _th_run(p, text, bold=True, size=16 if level==1 else 14)
    return p

def _add_simple_table(doc, headers, rows, hdr_bg='C8E6C9'):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        p = cell.paragraphs[0]
        _th_run(p, h, bold=True, size=13)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell, hdr_bg)
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = tbl.rows[i+1].cells[j]
            p = cell.paragraphs[0]
            _th_run(p, str(val), size=13)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return tbl

def build_report_esal(ss):
    doc = _new_doc()
    if not doc: return None
    _add_section_heading(doc, "1. ผลการคำนวณ ESAL (Equivalent Single Axle Load)")
    p = doc.add_paragraph()
    _th_run(p, f"วิธีคำนวณ AASHTO 1993  |  LDF = {ss.get('ldf',0.9)}  |  DDF = {ss.get('ddf',0.5)}", size=13)
    doc.add_paragraph()

    if ss.get('esal_rigid'):
        _add_section_heading(doc, "1.1 ESAL – ผิวทางคอนกรีต", level=2)
        hdr = ["Slab Thickness", "ESAL (Design Lane)"]
        rows = [[lbl, f"{ss['esal_rigid'].get(D,0):,.0f}"]
                for D, lbl in zip(SLAB_THICKNESSES, SLAB_LABELS)]
        _add_simple_table(doc, hdr, rows)
        doc.add_paragraph()

    if ss.get('esal_flex'):
        _add_section_heading(doc, "1.2 ESAL – ผิวทางลาดยาง", level=2)
        hdr = ["Structure Number (SN)", "ESAL (Design Lane)"]
        rows = [[f"SN = {sn}", f"{v:,.0f}"] for sn,v in ss['esal_flex'].items()]
        _add_simple_table(doc, hdr, rows)

    _add_footer(doc)
    return _doc_to_bytes(doc)

def build_report_cbr(ss):
    doc = _new_doc()
    if not doc: return None
    _add_section_heading(doc, "2. ผลการวิเคราะห์ค่า CBR")
    cbr_vals = ss.get('cbr_values', [])
    if cbr_vals:
        arr, n, u_cbr, u_pct = calc_percentile_cbr(cbr_vals)
        pct = ss.get('cbr_percentile', 90)
        cbr_d = float(np.interp(pct, u_pct[::-1], u_cbr[::-1]))
        p = doc.add_paragraph()
        _th_run(p, f"จำนวนตัวอย่าง: {n}  |  Percentile: {pct}%  |  CBR ออกแบบ: {cbr_d:.2f}%", size=13)
        doc.add_paragraph()
        hdr = ["ลำดับ","CBR (%)","จำนวน ≥","Percentile (%)"]
        rows = []
        for i,(v,p2) in enumerate(zip(u_cbr, u_pct)):
            cnt = int(np.sum(arr >= v))
            rows.append([str(i+1), f"{v:.2f}", str(cnt), f"{p2:.1f}"])
        _add_simple_table(doc, hdr, rows)
    _add_footer(doc)
    return _doc_to_bytes(doc)

def build_report_flexible(ss):
    doc = _new_doc()
    if not doc: return None
    _add_section_heading(doc, "3. ผลการออกแบบโครงสร้างชั้นทางลาดยาง (AASHTO 1993)")
    res = ss.get('flex_results', {})
    if res:
        p = doc.add_paragraph()
        _th_run(p, f"Design ESAL = {res.get('esal',0):,.0f}  |  SN Required = {res.get('sn_req',0):.3f}  |  SN Provided = {res.get('sn_prov',0):.3f}", size=13)
        p2 = doc.add_paragraph()
        status = "✅ PASS" if res.get('pass') else "❌ FAIL"
        _th_run(p2, f"ผลการตรวจสอบ: {status}", bold=True, size=14)
        doc.add_paragraph()
        layers = res.get('layers', [])
        if layers:
            hdr = ["ชั้นที่","วัสดุ","หนา (cm)","ai","mi","SNi","ΣSNi"]
            rows = [[str(l['layer']),l['material'],str(l['h_cm']),
                     f"{l['ai']:.2f}",f"{l['mi']:.1f}",
                     f"{l['sni']:.3f}",f"{l['cum_sn']:.3f}"] for l in layers]
            _add_simple_table(doc, hdr, rows)
    _add_footer(doc)
    return _doc_to_bytes(doc)

def build_report_kvalue(ss):
    doc = _new_doc()
    if not doc: return None
    _add_section_heading(doc, "4. ค่า k_eff (Effective Modulus of Subgrade Reaction)")
    p = doc.add_paragraph()
    _th_run(p, f"k∞ = {ss.get('k_inf',0):.1f} pci  |  LS = {ss.get('ls_value',0)}  |  k_eff = {ss.get('k_corrected',0):.1f} pci", size=13)
    imgs = ss.get('nomograph_img_k')
    if imgs:
        doc.add_paragraph()
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_img.add_run()
        r.add_picture(io.BytesIO(imgs), width=Cm(12))
    imgs2 = ss.get('nomograph_img_ls')
    if imgs2:
        doc.add_paragraph()
        p_img2 = doc.add_paragraph()
        p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p_img2.add_run()
        r2.add_picture(io.BytesIO(imgs2), width=Cm(12))
    _add_footer(doc)
    return _doc_to_bytes(doc)

def build_report_rigid(ss):
    doc = _new_doc()
    if not doc: return None
    _add_section_heading(doc, "5. ผลการออกแบบความหนาถนนคอนกรีต (AASHTO 1993)")
    rigid_res = ss.get('rigid_results', {})
    for ptype, res in rigid_res.items():
        if not res: continue
        _add_section_heading(doc, f"5.{list(rigid_res.keys()).index(ptype)+1} {ptype}", level=2)
        p = doc.add_paragraph()
        _th_run(p, f"Slab = {res.get('d_cm',0)} cm  |  k_eff = {res.get('k_eff',0):.1f} pci  |  f'c = {res.get('fc',0)} ksc", size=13)
        p2 = doc.add_paragraph()
        status = "✅ PASS" if res.get('pass') else "❌ FAIL"
        _th_run(p2, f"W18 Capacity = {res.get('w18_cap',0):,.0f}  |  W18 Required = {res.get('w18_req',0):,.0f}  |  {status}", bold=True, size=13)
        doc.add_paragraph()
    _add_footer(doc)
    return _doc_to_bytes(doc)

def build_report_full(ss):
    doc = _new_doc()
    if not doc: return None
    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _th_run(p, "รายการคำนวณออกแบบโครงสร้างชั้นทาง", bold=True, size=20)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _th_run(p2, "ตามวิธี AASHTO 1993", size=16)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _th_run(p3, f"วันที่: {datetime.now().strftime('%d/%m/%Y %H:%M')}", size=13)
    doc.add_page_break()

    sections = [
        ('esal_rigid', build_report_esal),
        ('cbr_values', build_report_cbr),
        ('flex_results', build_report_flexible),
        ('k_corrected', build_report_kvalue),
        ('rigid_results', build_report_rigid),
    ]
    for key, fn in sections:
        if ss.get(key):
            sub_bytes = fn(ss)
            if sub_bytes:
                sub_doc = DocxDoc(io.BytesIO(sub_bytes))
                for elem in sub_doc.element.body:
                    doc.element.body.append(elem)
    _add_footer(doc)
    return _doc_to_bytes(doc)

def _add_footer(doc):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _th_run(p, "พัฒนาโดย รศ.ดร.อิทธิพล มีผล  |  ภาควิชาครุศาสตร์โยธา  |  มจพ.", italic=False, size=12, color=(80,80,80))

# patch _th_run for italic param
_th_run_orig = _th_run
def _th_run(para, text, bold=False, size=15, color=None, italic=False):
    run = para.add_run(text)
    run.font.name = 'TH SarabunPSK'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    try:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')
    except:
        pass
    return run

# ─────────────────────────────────────────────
#  SEC 8: SESSION STATE INIT
# ─────────────────────────────────────────────
def ss_init():
    defaults = {
        # Traffic & ESAL
        'traffic_df':      None,
        'esal_rigid':      {},
        'esal_flex':       {},
        'ldf':             0.9,
        'ddf':             0.5,
        'pt_global':       2.5,
        'pt_rigid':        2.5,
        'pt_flex':         2.5,
        '_pt_sync':        2.5,
        'sn_list':         [6.5, 7.1, 7.5, 8.0],
        # CBR
        'cbr_values':      [],
        'cbr_percentile':  90.0,
        'cbr_design':      3.0,
        'mr_subgrade_psi': 4500.0,
        'k_subgrade_pci':  231.9,
        # Flexible
        'flex_results':    {},
        'cbr_fl_val':      3.0,
        'mr_fl_val':       4500.0,
        # K-Value
        'k_inf':           0.0,
        'k_corrected':     0.0,
        'ls_value':        1.0,
        'nomograph_img_k': None,
        'nomograph_img_ls':None,
        'layer_esb_psi':   50000,
        'layer_dsb_in':    6.0,
        # Rigid
        'rigid_results':   {},
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

ss_init()
ss = st.session_state

# ─────────────────────────────────────────────
#  HELPER: Status badge
# ─────────────────────────────────────────────
def status_badge(key, label=None):
    val = ss.get(key)
    has = (val is not None and val != {} and val != [] and val != 0.0)
    cls = "badge-ready" if has else "badge-wait"
    icon = "✅" if has else "⚠️"
    lbl = label or key
    return f'<span class="{cls}">{icon} {lbl}</span>'

# ─────────────────────────────────────────────
#  SIDEBAR
# ─────────────────────────────────────────────
with st.sidebar:
    # ── Dependency warnings ──
    if not OPENPYXL_OK:
        st.warning("⚠️ openpyxl ไม่พร้อม — Upload Excel จะไม่ทำงาน")
    if not DOCX_OK:
        st.warning("⚠️ python-docx ไม่พร้อม — Word report จะไม่ทำงาน")

    st.markdown("""
    <div style='text-align:center;padding:1rem 0 0.5rem;'>
        <div style='font-size:2.2rem;'>🛣️</div>
        <div style='font-weight:700;font-size:1.1rem;color:#A5D6A7;'>ITM Pave Pro</div>
        <div style='font-size:0.78rem;color:#81C784;'>AASHTO 1993</div>
    </div>
    """, unsafe_allow_html=True)
    st.divider()

    st.markdown("**📊 สถานะข้อมูล**")
    st.markdown(status_badge('esal_rigid','ESAL Rigid'), unsafe_allow_html=True)
    st.markdown(status_badge('esal_flex','ESAL Flexible'), unsafe_allow_html=True)
    st.markdown(status_badge('cbr_values','CBR Data'), unsafe_allow_html=True)
    st.markdown(status_badge('flex_results','Flex Design'), unsafe_allow_html=True)
    st.markdown(status_badge('k_corrected','K-Value'), unsafe_allow_html=True)
    st.markdown(status_badge('rigid_results','Rigid Design'), unsafe_allow_html=True)
    st.divider()

    st.markdown("**💾 Save / Load Project**")
    if st.button("💾 Save JSON", use_container_width=True):
        save_data = {
            'esal_rigid':     ss.esal_rigid,
            'esal_flex':      {str(k):v for k,v in ss.esal_flex.items()},
            'ldf':            ss.ldf, 'ddf': ss.ddf,
            'pt_global':      ss.pt_global,
            'pt_rigid':       ss.pt_rigid, 'pt_flex': ss.pt_flex,
            'sn_list':        ss.sn_list,
            'cbr_values':     ss.cbr_values,
            'cbr_percentile': ss.cbr_percentile,
            'cbr_design':     ss.cbr_design,
            'mr_subgrade_psi':ss.mr_subgrade_psi,
            'k_subgrade_pci': ss.k_subgrade_pci,
            'flex_results':   ss.flex_results,
            'k_inf':          ss.k_inf,
            'k_corrected':    ss.k_corrected,
            'ls_value':       ss.ls_value,
            'rigid_results':  ss.rigid_results,
            'traffic_df':     ss.traffic_df.to_dict('records') if ss.traffic_df is not None else None,
        }
        json_bytes = json.dumps(save_data, ensure_ascii=False, indent=2).encode('utf-8')
        st.download_button("📥 Download JSON", json_bytes,
                           file_name=f"itm_pave_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
                           mime="application/json", use_container_width=True)

    uploaded_json = st.file_uploader("📂 Load JSON", type=['json'])
    if uploaded_json:
        try:
            data = json.loads(uploaded_json.read().decode('utf-8'))
            for k in ['esal_rigid','ldf','ddf','pt_global','pt_rigid','pt_flex','sn_list',
                      'cbr_values','cbr_percentile','cbr_design','mr_subgrade_psi',
                      'k_subgrade_pci','flex_results','k_inf','k_corrected',
                      'ls_value','rigid_results']:
                if k in data:
                    ss[k] = data[k]
            if 'esal_flex' in data:
                ss.esal_flex = {float(k):v for k,v in data['esal_flex'].items()}
            if data.get('traffic_df'):
                ss.traffic_df = pd.DataFrame(data['traffic_df'])
            st.success("✅ โหลดข้อมูลสำเร็จ!")
            st.rerun()
        except Exception as e:
            st.error(f"❌ โหลดไม่สำเร็จ: {e}")

    st.divider()
    st.markdown("""
    <div style='font-size:0.72rem;color:#81C784;text-align:center;line-height:1.8;'>
    รศ.ดร.อิทธิพล มีผล<br>
    ภาควิชาครุศาสตร์โยธา<br>
    คณะครุศาสตร์อุตสาหกรรม มจพ.<br>
    ITM Pave Pro v2.1
    </div>
    """, unsafe_allow_html=True)

# ─────────────────────────────────────────────
#  HEADER
# ─────────────────────────────────────────────
st.markdown("""
<div class="main-header">
    <h1>🛣️ ITM Pave Pro — ระบบออกแบบโครงสร้างชั้นทาง AASHTO 1993</h1>
    <p>ESAL Calculator · CBR Analysis · Flexible Design · Rigid Design · Report</p>
</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
#  MAIN TABS
# ─────────────────────────────────────────────
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "🚛 ESAL Calculator",
    "📊 CBR Analysis",
    "🔧 Flexible Design",
    "🏗️ Rigid Design",
    "📄 Report & Save",
])

# ══════════════════════════════════════════════
#  TAB 1: ESAL CALCULATOR
# ══════════════════════════════════════════════
with tab1:
    st.markdown("### 🚛 ESAL Calculator — AASHTO 1993")

    # ── Pt Global ──
    st.markdown('<div class="card"><h4>🎯 Terminal Serviceability (Pt) — ค่าร่วม</h4>', unsafe_allow_html=True)
    pt_g_col1, pt_g_col2 = st.columns([1, 3])
    with pt_g_col1:
        pt_global = st.number_input(
            "Pt (Global Default)",
            value=float(ss.pt_global), step=0.1,
            min_value=2.0, max_value=3.0,
            key="pt_global_input",
            help="ค่านี้จะเป็น default ใน TAB Flexible และ TAB Rigid — แก้ได้อิสระในแต่ละ TAB"
        )
        if pt_global != ss.pt_global:
            ss.pt_global  = pt_global
            ss['_pt_sync'] = pt_global
            st.rerun()
    with pt_g_col2:
        st.markdown(f"""
        <div class="result-info" style="margin-top:1.6rem;">
            Pt = <b>{pt_global}</b> &nbsp;→&nbsp;
            จะส่งเป็น default ไปยัง <b>TAB Flexible Design</b> และ <b>TAB Rigid Design</b>
            &nbsp;(แก้ไขได้อิสระในแต่ละ TAB)
        </div>""", unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    sub_esal_rigid, sub_esal_flex = st.tabs(["🔴 Rigid Pavement", "🟢 Flexible Pavement"])

    # ── Shared Traffic Input (above sub-tabs) ──
    with st.expander("📋 ข้อมูลปริมาณจราจร (ใช้ร่วมกันทั้ง Rigid & Flexible)", expanded=True):
        col_inp1, col_inp2 = st.columns([1, 1])

        with col_inp1:
            st.markdown('<div class="card"><h4>📁 Upload Excel / กรอกมือ</h4>', unsafe_allow_html=True)
            input_mode = st.radio("วิธีกรอกข้อมูล", ["📁 Upload Excel", "✏️ กรอกมือ + Growth Rate"], horizontal=True)

            if input_mode == "📁 Upload Excel":
                uploaded_xl = st.file_uploader("เลือกไฟล์ Excel (.xlsx)", type=['xlsx'])
                st.caption("รูปแบบ: คอลัมน์ Year, MB, HB, MT, HT, TR, STR")
                if uploaded_xl:
                    if not OPENPYXL_OK:
                        st.error("❌ openpyxl ไม่ได้ติดตั้ง — กรุณาใช้วิธี 'กรอกมือ + Growth Rate' แทน")
                    else:
                        try:
                            df_up = pd.read_excel(uploaded_xl, engine='openpyxl')
                            df_up.columns = [c.strip() for c in df_up.columns]
                            col_map = {}
                            for c in df_up.columns:
                                for vc in ['Year']+VEHICLE_COLS:
                                    if c.upper() == vc.upper():
                                        col_map[c] = vc
                            df_up = df_up.rename(columns=col_map)
                            for vc in VEHICLE_COLS:
                                if vc not in df_up.columns:
                                    df_up[vc] = 0
                            ss.traffic_df = df_up[['Year']+VEHICLE_COLS].fillna(0)
                            st.success(f"✅ อ่านข้อมูล {len(df_up)} ปีสำเร็จ")
                        except Exception as e:
                            st.error(f"❌ {e}")
            else:
                st.markdown("**ปริมาณจราจรปีแรก (คัน/วัน)**")
                base_cols = st.columns(6)
                base_row = {}
                defaults_base = {"MB":120,"HB":60,"MT":250,"HT":180,"TR":100,"STR":120}
                for i, vc in enumerate(VEHICLE_COLS):
                    with base_cols[i]:
                        base_row[vc] = st.number_input(vc, value=defaults_base[vc],
                                                       min_value=0, step=10, key=f"base_{vc}")
                gc1, gc2 = st.columns(2)
                with gc1:
                    growth_rate = st.number_input("Growth Rate (%/ปี)", value=4.5, step=0.5, min_value=0.0, max_value=20.0)
                with gc2:
                    design_years = st.number_input("Design Period (ปี)", value=20, min_value=1, max_value=40, step=1)
                if st.button("🔄 สร้างตารางจราจร", type="primary"):
                    ss.traffic_df = grow_traffic(base_row, growth_rate, int(design_years))
                    st.success(f"✅ สร้างตาราง {int(design_years)} ปีสำเร็จ")

            st.markdown('</div>', unsafe_allow_html=True)

        with col_inp2:
            if ss.traffic_df is not None:
                st.markdown('<div class="card"><h4>📊 ตารางปริมาณจราจร</h4>', unsafe_allow_html=True)
                st.dataframe(ss.traffic_df.style.format({c: "{:,.0f}" for c in VEHICLE_COLS}),
                             use_container_width=True, height=280)
                total_row = {vc: ss.traffic_df[vc].sum() for vc in VEHICLE_COLS}
                st.markdown(f'<div class="result-info">📊 รวมตลอดอายุออกแบบ: '
                            + " | ".join(f"<b>{vc}</b>: {total_row[vc]:,.0f}" for vc in VEHICLE_COLS)
                            + '</div>', unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)
            else:
                st.info("⬅️ กรอกหรือ Upload ข้อมูลจราจรก่อน")

    # ─── Sub-tab: Rigid ───
    with sub_esal_rigid:
        st.markdown('<div class="card"><h4>⚙️ พารามิเตอร์ – Rigid Pavement</h4>', unsafe_allow_html=True)
        c1,c2,c3,c4 = st.columns(4)
        with c1: ldf_r = st.number_input("Lane Distribution Factor", value=0.9, step=0.05, min_value=0.1, max_value=1.0, key="ldf_r")
        with c2: ddf_r = st.number_input("Directional Dist. Factor", value=0.5, step=0.05, min_value=0.1, max_value=1.0, key="ddf_r")
        with c3: pt_r  = st.number_input("Terminal Serviceability Pt", value=2.5, step=0.1, min_value=1.5, max_value=3.5, key="pt_r")
        with c4: st.markdown(f"<br><small>Pi (Rigid) = 4.5</small>", unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

        # EALF Table
        st.markdown('<div class="card"><h4>📋 Truck Factor (EALF/คัน) ตาม Slab Thickness</h4>', unsafe_allow_html=True)
        tf_rows = []
        for vt in VEHICLE_COLS:
            row = {"ประเภทรถ": f"{VEHICLE_LABELS[vt]} ({vt})"}
            for D, lbl in zip(SLAB_THICKNESSES, SLAB_LABELS):
                row[lbl] = f"{truck_factor_rigid(vt,D,pt_r):.3f}"
            tf_rows.append(row)
        st.dataframe(pd.DataFrame(tf_rows), use_container_width=True, hide_index=True)
        st.markdown('</div>', unsafe_allow_html=True)

        if st.button("🔄 คำนวณ ESAL Rigid", type="primary", key="calc_r"):
            if ss.traffic_df is None:
                st.warning("⚠️ กรุณากรอกข้อมูลจราจรก่อน")
            else:
                esal_r = compute_esal_from_df(ss.traffic_df, ldf_r, ddf_r, pt_r, mode="rigid")
                ss.esal_rigid = esal_r
                ss.ldf = ldf_r; ss.ddf = ddf_r; ss.pt_rigid = pt_r

                st.markdown("---")
                st.markdown("#### 📊 ผลการคำนวณ ESAL – Rigid Pavement")
                cols_m = st.columns(len(SLAB_THICKNESSES))
                for i, (D, lbl) in enumerate(zip(SLAB_THICKNESSES, SLAB_LABELS)):
                    with cols_m[i]:
                        st.markdown(f"""
                        <div class="metric-box">
                            <div class="val">{esal_r[D]:,.0f}</div>
                            <div class="lbl">ESAL – {lbl}</div>
                        </div>""", unsafe_allow_html=True)
                st.markdown('<div class="result-info">✅ ค่า ESAL บันทึกแล้ว → ใช้ได้ใน Tab K-Value และ Rigid Design</div>', unsafe_allow_html=True)

        if ss.esal_rigid:
            st.markdown("**ค่า ESAL Rigid ปัจจุบัน:**")
            df_er = pd.DataFrame({
                "Slab": SLAB_LABELS,
                "ESAL": [f"{ss.esal_rigid.get(D,0):,.0f}" for D in SLAB_THICKNESSES]
            })
            st.dataframe(df_er, use_container_width=True, hide_index=True)

    # ─── Sub-tab: Flexible ───
    with sub_esal_flex:
        st.markdown('<div class="card"><h4>⚙️ พารามิเตอร์ – Flexible Pavement</h4>', unsafe_allow_html=True)
        c1,c2,c3,c4 = st.columns(4)
        with c1: ldf_f = st.number_input("Lane Distribution Factor", value=0.9, step=0.05, min_value=0.1, max_value=1.0, key="ldf_f")
        with c2: ddf_f = st.number_input("Directional Dist. Factor", value=0.5, step=0.05, min_value=0.1, max_value=1.0, key="ddf_f")
        with c3: pt_f  = st.number_input("Terminal Serviceability Pt", value=2.5, step=0.1, min_value=1.5, max_value=3.5, key="pt_f")
        with c4: st.markdown(f"<br><small>Pi (Flexible) = 4.2</small>", unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="card"><h4>📐 กำหนด Structure Number (SN)</h4>', unsafe_allow_html=True)
        sn_cols = st.columns(4)
        user_sn = []
        sn_defs = [6.5, 7.1, 7.5, 8.0]
        for i, col in enumerate(sn_cols):
            with col:
                user_sn.append(round(st.number_input(f"SN {i+1}", value=sn_defs[i],
                                min_value=1.0, max_value=20.0, step=0.1,
                                key=f"sn_{i}", format="%.1f"), 2))
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="card"><h4>📋 Truck Factor (EALF/คัน) ตาม SN</h4>', unsafe_allow_html=True)
        tf_rows_f = []
        for vt in VEHICLE_COLS:
            row = {"ประเภทรถ": f"{VEHICLE_LABELS[vt]} ({vt})"}
            for sn in user_sn:
                row[f"SN={sn}"] = f"{truck_factor_flex(vt,sn,pt_f):.3f}"
            tf_rows_f.append(row)
        st.dataframe(pd.DataFrame(tf_rows_f), use_container_width=True, hide_index=True)
        st.markdown('</div>', unsafe_allow_html=True)

        if st.button("🔄 คำนวณ ESAL Flexible", type="primary", key="calc_f"):
            if ss.traffic_df is None:
                st.warning("⚠️ กรุณากรอกข้อมูลจราจรก่อน")
            else:
                esal_fv = compute_esal_from_df(ss.traffic_df, ldf_f, ddf_f, pt_f, mode="flex", sn_list=user_sn)
                ss.esal_flex = esal_fv
                ss.sn_list = user_sn
                ss.pt_flex = pt_f

                st.markdown("---")
                st.markdown("#### 📊 ผลการคำนวณ ESAL – Flexible Pavement")
                cols_m2 = st.columns(len(user_sn))
                for i, sn in enumerate(user_sn):
                    with cols_m2[i]:
                        st.markdown(f"""
                        <div class="metric-box">
                            <div class="val">{esal_fv[sn]:,.0f}</div>
                            <div class="lbl">ESAL – SN {sn}</div>
                        </div>""", unsafe_allow_html=True)
                st.markdown('<div class="result-info">✅ ค่า ESAL บันทึกแล้ว → ใช้ได้ใน Tab Flexible Design</div>', unsafe_allow_html=True)

        if ss.esal_flex:
            st.markdown("**ค่า ESAL Flexible ปัจจุบัน:**")
            df_ef = pd.DataFrame({"SN": [f"SN {k}" for k in ss.esal_flex.keys()],
                                   "ESAL": [f"{v:,.0f}" for v in ss.esal_flex.values()]})
            st.dataframe(df_ef, use_container_width=True, hide_index=True)

# ══════════════════════════════════════════════
#  TAB 2: CBR ANALYSIS
# ══════════════════════════════════════════════
with tab2:
    st.markdown("### 📊 CBR Analysis — Percentile Method")

    col_cbr_l, col_cbr_r = st.columns([1, 1])

    with col_cbr_l:
        st.markdown('<div class="card"><h4>📁 ข้อมูล CBR</h4>', unsafe_allow_html=True)
        cbr_mode = st.radio("แหล่งข้อมูล", ["📁 Upload Excel", "✏️ กรอกค่า", "📌 ใช้ข้อมูลตัวอย่าง"], horizontal=True)

        cbr_vals_input = None
        if cbr_mode == "📁 Upload Excel":
            cbr_xl = st.file_uploader("ไฟล์ Excel (คอลัมน์ CBR)", type=['xlsx'], key="cbr_xl")
            if cbr_xl:
                try:
                    df_cbr = pd.read_excel(cbr_xl, engine='openpyxl')
                    col_cbr = next((c for c in df_cbr.columns if 'cbr' in c.lower()), df_cbr.columns[0])
                    cbr_vals_input = pd.to_numeric(df_cbr[col_cbr], errors='coerce').dropna().tolist()
                    st.success(f"✅ {len(cbr_vals_input)} ตัวอย่าง")
                except Exception as e:
                    st.error(str(e))
        elif cbr_mode == "✏️ กรอกค่า":
            cbr_txt = st.text_area("กรอกค่า CBR (%) คั่นด้วย , หรือ Enter",
                                   placeholder="6.5, 7.2, 8.1, 5.3, ...",
                                   height=120)
            if cbr_txt.strip():
                import re
                parts = re.split(r'[,\n\r\s]+', cbr_txt.strip())
                try:
                    cbr_vals_input = [float(x) for x in parts if x]
                    st.success(f"✅ {len(cbr_vals_input)} ค่า")
                except:
                    st.error("กรุณากรอกตัวเลขเท่านั้น")
        else:
            cbr_vals_input = SAMPLE_CBR
            st.info(f"📌 ใช้ข้อมูลตัวอย่าง {len(SAMPLE_CBR)} ค่า")

        if cbr_vals_input:
            ss.cbr_values = cbr_vals_input

        target_pct = st.slider("Percentile ที่ต้องการ (%)", 50, 99, int(ss.cbr_percentile), step=1, key="pct_slider")
        ss.cbr_percentile = float(target_pct)
        st.markdown('</div>', unsafe_allow_html=True)

        if ss.cbr_values:
            arr, n, u_cbr, u_pct = calc_percentile_cbr(ss.cbr_values)
            cbr_at_pct = float(np.interp(target_pct, u_pct[::-1], u_cbr[::-1]))

            st.markdown('<div class="card"><h4>🎯 ค่า CBR ที่ใช้ออกแบบ</h4>', unsafe_allow_html=True)
            mr_auto = cbr_to_mr(cbr_at_pct)
            k_auto  = mr_to_k(mr_auto)

            c_m1, c_m2, c_m3 = st.columns(3)
            with c_m1:
                st.markdown(f"""<div class="metric-box">
                    <div class="val">{cbr_at_pct:.2f}</div>
                    <div class="lbl">CBR @ P{target_pct:.0f} (%)</div>
                </div>""", unsafe_allow_html=True)
            with c_m2:
                st.markdown(f"""<div class="metric-box">
                    <div class="val">{mr_auto:,.0f}</div>
                    <div class="lbl">Mr (psi) = 1500×CBR</div>
                </div>""", unsafe_allow_html=True)
            with c_m3:
                st.markdown(f"""<div class="metric-box">
                    <div class="val">{k_auto:.1f}</div>
                    <div class="lbl">k subgrade (pci)</div>
                </div>""", unsafe_allow_html=True)

            design_cbr_input = st.number_input(
                "CBR ที่ใช้ออกแบบจริง (ปรับได้)",
                value=float(round(cbr_at_pct, 1)),
                min_value=0.5, max_value=100.0, step=0.5,
                key="design_cbr_input"
            )
            mr_design = cbr_to_mr(design_cbr_input)
            k_design  = mr_to_k(mr_design)
            st.markdown(f"""
            <div class="result-info">
                CBR ออกแบบ = <b>{design_cbr_input:.1f}%</b> →
                Mr = <b>{mr_design:,.0f} psi</b> →
                k_subgrade = <b>{k_design:.1f} pci</b>
            </div>""", unsafe_allow_html=True)

            if st.button("✅ ใช้ค่านี้", type="primary", key="use_cbr"):
                ss.cbr_design      = design_cbr_input
                ss.mr_subgrade_psi = mr_design
                ss.k_subgrade_pci  = k_design
                st.success("✅ บันทึกค่า CBR/Mr/k แล้ว → ใช้ได้ใน Tab Flexible, K-Value, Rigid Design")
            st.markdown('</div>', unsafe_allow_html=True)

    with col_cbr_r:
        if ss.cbr_values:
            arr, n, u_cbr, u_pct = calc_percentile_cbr(ss.cbr_values)
            cbr_at_pct = float(np.interp(target_pct, u_pct[::-1], u_cbr[::-1]))

            st.markdown('<div class="card"><h4>📈 กราฟ Percentile vs CBR</h4>', unsafe_allow_html=True)
            if PLOTLY_OK:
                fig_cbr = go.Figure()
                fig_cbr.add_trace(go.Scatter(
                    x=u_cbr, y=u_pct, mode='lines+markers', name='CBR Distribution',
                    line=dict(color='#2E7D32', width=2.5),
                    marker=dict(size=7, symbol='x', color='#1B5E20')
                ))
                fig_cbr.add_trace(go.Scatter(
                    x=[0, cbr_at_pct], y=[target_pct, target_pct],
                    mode='lines', name=f'P{target_pct:.0f}%',
                    line=dict(color='red', width=2, dash='dash')
                ))
                fig_cbr.add_trace(go.Scatter(
                    x=[cbr_at_pct, cbr_at_pct], y=[0, target_pct],
                    mode='lines', name=f'CBR={cbr_at_pct:.2f}%',
                    line=dict(color='red', width=2, dash='dash')
                ))
                fig_cbr.add_annotation(
                    x=cbr_at_pct, y=0,
                    text=f"<b>{cbr_at_pct:.2f}%</b>",
                    showarrow=True, arrowhead=2, arrowcolor='red',
                    font=dict(size=14, color='red'), ay=40
                )
                fig_cbr.update_layout(
                    xaxis_title="CBR (%)", yaxis_title="Percentile (%)",
                    plot_bgcolor='white', height=380,
                    xaxis=dict(range=[0, max(u_cbr)*1.1], gridcolor='#E8F5E9'),
                    yaxis=dict(range=[0, 100], gridcolor='#E8F5E9'),
                    legend=dict(bgcolor='rgba(255,255,255,0.8)', bordercolor='#C8E6C9'),
                    margin=dict(l=50, r=30, t=30, b=50)
                )
                st.plotly_chart(fig_cbr, use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)

            # Stats
            st.markdown('<div class="card"><h4>📋 สถิติ CBR</h4>', unsafe_allow_html=True)
            c1,c2,c3,c4 = st.columns(4)
            with c1: st.metric("n", n)
            with c2: st.metric("Min", f"{np.min(ss.cbr_values):.2f}%")
            with c3: st.metric("Max", f"{np.max(ss.cbr_values):.2f}%")
            with c4: st.metric("Mean", f"{np.mean(ss.cbr_values):.2f}%")
            st.markdown('</div>', unsafe_allow_html=True)

# ─────────────────────────────────────────────
#  AASHTO 1993 D_min HELPER
# ─────────────────────────────────────────────
def _aashto_badge(sn_req, cum_sn_before, ai, mi, h_cm):
    """
    คำนวณ D_min ตาม AASHTO 1993:
        Di_min = SNi_remaining / (ai × mi)   [inch → cm]
    โดย SNi_remaining = SN_required - ΣSN ชั้นบน

    Returns: HTML badge string
    """
    if sn_req <= 0 or ai <= 0:
        return ""
    sn_remaining = sn_req - cum_sn_before          # SN ที่ชั้นนี้ต้องรับ
    if sn_remaining <= 0:
        # ชั้นบนๆ รับครบแล้ว — ชั้นนี้ไม่จำเป็น
        d_min_cm = 0.0
    else:
        d_min_in = sn_remaining / (ai * mi)
        d_min_cm = d_min_in * 2.54

    passed = h_cm >= d_min_cm - 0.05              # tolerance 0.05 cm

    if d_min_cm <= 0:
        return ('<div style="font-size:0.78rem;color:#2E7D32;font-weight:600;'
                'margin-top:0.25rem;">✅ ผ่าน — ชั้นบนรับ SN ครบแล้ว</div>')
    elif passed:
        return (f'<div style="font-size:0.78rem;color:#2E7D32;font-weight:600;'
                f'background:#E8F5E9;border-radius:5px;padding:0.2rem 0.5rem;'
                f'margin-top:0.25rem;border:1px solid #A5D6A7;">'
                f'✅ ผ่าน &nbsp;—&nbsp; D<sub>min</sub> = '
                f'<b>{d_min_cm:.1f} cm</b> '
                f'({d_min_in:.2f} in) &nbsp;|&nbsp; '
                f'SN<sub>i</sub> ต้องการ = {sn_remaining:.3f}</div>')
    else:
        return (f'<div style="font-size:0.78rem;color:#B71C1C;font-weight:600;'
                f'background:#FFF8E1;border-radius:5px;padding:0.2rem 0.5rem;'
                f'margin-top:0.25rem;border:1px solid #FFE082;">'
                f'💡 D<sub>min</sub> = <b>{d_min_cm:.1f} cm</b> '
                f'({d_min_in:.2f} in) &nbsp;|&nbsp; '
                f'SN<sub>i</sub> ต้องการ = {sn_remaining:.3f} &nbsp;'
                f'<span style="color:#E65100;">(กรอกอยู่ {h_cm} cm)</span></div>')

# ─────────────────────────────────────────────
#  CBR ↔ Mr SYNC CALLBACKS
# ─────────────────────────────────────────────
def _on_cbr_fl_change():
    """CBR เปลี่ยน → คำนวณ Mr อัตโนมัติ"""
    cbr_val = max(0.5, float(st.session_state.get('cbr_fl_input', 3.0)))
    st.session_state['cbr_fl_val'] = cbr_val
    st.session_state['mr_fl_val']  = max(500.0, cbr_to_mr(cbr_val))

def _on_mr_fl_change():
    """Mr เปลี่ยน → คำนวณ CBR ย้อนกลับ"""
    mr_val = max(500.0, float(st.session_state.get('mr_fl_input', 4500.0)))
    st.session_state['mr_fl_val']  = mr_val
    st.session_state['cbr_fl_val'] = max(0.5, mr_val / 1500.0)

# ══════════════════════════════════════════════
#  TAB 3: FLEXIBLE DESIGN
# ══════════════════════════════════════════════
with tab3:
    st.markdown("### 🔧 Flexible Pavement Design — AASHTO 1993")

    col_fl, col_fr = st.columns([1, 1])

    with col_fl:
        # Auto-fill from session
        st.markdown('<div class="card"><h4>📥 Design ESAL</h4>', unsafe_allow_html=True)
        if ss.esal_flex:
            sn_keys = list(ss.esal_flex.keys())
            sel_idx = st.selectbox("เลือก SN", range(len(sn_keys)),
                                   format_func=lambda i: f"SN {sn_keys[i]}  →  ESAL = {ss.esal_flex[sn_keys[i]]:,.0f}",
                                   key="flex_sn_sel")
            design_esal_f = ss.esal_flex[sn_keys[sel_idx]]
            st.markdown(f'<div class="result-info">📊 Design ESAL = <b>{design_esal_f:,.0f}</b></div>', unsafe_allow_html=True)
        else:
            st.warning("⚠️ ยังไม่มีค่า ESAL — คำนวณใน Tab 1 ก่อน หรือกรอกเอง")
            design_esal_f = st.number_input("Design ESAL (กรอกเอง)", value=0, step=100000, key="flex_esal_manual")
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="card"><h4>🌍 Subgrade</h4>', unsafe_allow_html=True)
        # ── Sync init: ถ้า cbr_design เพิ่งโหลดมาจาก TAB2 ให้อัพเดต ss ด้วย ──
        _cbr_from_tab2 = float(ss.cbr_design) if ss.cbr_design else 3.0
        if abs(ss.get('cbr_fl_val', 0) - _cbr_from_tab2) > 0.001 and ss.get('_cbr_fl_user_edited') != True:
            ss['cbr_fl_val'] = _cbr_from_tab2
            ss['mr_fl_val']  = cbr_to_mr(_cbr_from_tab2)
        # ── clamp ค่าให้อยู่ในช่วง min ก่อน render ──
        ss['cbr_fl_val'] = max(0.5, float(ss.get('cbr_fl_val', 3.0)))
        ss['mr_fl_val']  = max(500.0, float(ss.get('mr_fl_val', 4500.0)))
        c1, c2 = st.columns(2)
        with c1:
            st.number_input(
                "CBR (%)", value=ss['cbr_fl_val'],
                step=0.5, min_value=0.5,
                key="cbr_fl_input",
                on_change=_on_cbr_fl_change
            )
            cbr_fl = ss['cbr_fl_val']
        with c2:
            st.number_input(
                "Mr (psi) [กรอกหรือคำนวณอัตโนมัติ]",
                value=ss['mr_fl_val'],
                step=500.0, min_value=500.0,
                key="mr_fl_input",
                on_change=_on_mr_fl_change
            )
            mr_fl = ss['mr_fl_val']
        ss['_cbr_fl_user_edited'] = True
        st.markdown(f"Mr = **{mr_fl:,.0f} psi**  ({mr_fl/145.038:.1f} MPa)")
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="card"><h4>⚙️ Design Parameters</h4>', unsafe_allow_html=True)
        c1,c2,c3,c4 = st.columns(4)
        with c1:
            r0_fl = st.selectbox("Reliability R0 (%)", list(ZR_MAP.keys()), index=10, key="r0_fl")
            st.caption(f"ZR = {ZR_MAP[r0_fl]}")
        with c2: so_fl = st.number_input("So", value=0.45, step=0.01, min_value=0.3, max_value=0.6, key="so_fl")
        with c3: pi_fl = st.number_input("Pi", value=4.2, step=0.1, key="pi_fl")
        with c4:
            use_pt_global_fl = st.checkbox("ใช้ Pt Global", value=ss.get('use_pt_global_fl', True), key="use_pt_global_fl")
            if use_pt_global_fl:
                pt_fl2 = float(ss.get('pt_global', 2.5))
                st.caption(f"Pt = {pt_fl2} (Global)")
            else:
                pt_fl2 = st.number_input("Pt (Override)", value=float(ss.get('pt_fl2_override', ss.get('pt_global', 2.5))),
                                          step=0.1, min_value=2.0, max_value=3.0, key="pt_fl2_override")
        st.markdown('</div>', unsafe_allow_html=True)

    with col_fr:
        st.markdown('<div class="card"><h4>🔩 Layer Design</h4>', unsafe_allow_html=True)

        # Header
        hc0, hc1, hc2, hc3 = st.columns([3, 1, 1, 4])
        hc0.markdown("**วัสดุ**")
        hc1.markdown("**หนา (cm)**")
        hc2.markdown("**mi**")
        hc3.markdown("**ผลคำนวณ**")

        mat_options   = list(FLEX_LAYER_MATERIALS.keys())
        layer_results = []
        cum_sn        = 0.0

        # ── คำนวณ SN Required ล่วงหน้าสำหรับ Smart Recommendation ──
        _zr_fl   = ZR_MAP.get(ss.get('r0_fl', 90), -1.282)
        _so_fl   = ss.get('so_fl', 0.45)
        _pi_fl   = ss.get('pi_fl', 4.2)
        _pt_fl2  = ss.get('pt_fl2', float(ss.get('pt_global', 2.5)))
        _mr_fl   = float(ss.get('mr_fl_val', 4500.0))
        _esal_f  = ss.get('esal_flex', {})
        _esal_f_val = list(_esal_f.values())[ss.get('flex_sn_sel', 0)] if _esal_f else ss.get('flex_esal_manual', 0)
        try:
            _sn_req = aashto_sn_required(_esal_f_val, _zr_fl, _so_fl, _pi_fl, _pt_fl2, _mr_fl) or 0.0
        except:
            _sn_req = 0.0

        for li in range(6):
            lc0, lc1, lc2, lc3 = st.columns([3, 1, 1, 4])

            # ── ตรวจ do_sub ล่วงหน้า (ก่อน render lc1) ──
            _is_ac_mat = ss.get(f"fmat_{li}", mat_options[0]) in AC_SURFACE_MATERIALS
            _do_sub_now = ss.get(f"fsub_{li}", False) and _is_ac_mat

            with lc0:
                mat_f = st.selectbox(f"L{li+1}", mat_options,
                                     key=f"fmat_{li}", label_visibility="collapsed")
            with lc1:
                if _do_sub_now:
                    # ── #3: read-only แสดงผลรวมชั้นย่อย ──
                    _h_wear_now = ss.get(f"fwear_{li}", 5)
                    _h_bind_now = ss.get(f"fbind_{li}", 5)
                    _h_base_now = ss.get(f"fbase_{li}", 7)
                    _h_total_now = _h_wear_now + _h_bind_now + _h_base_now
                    st.markdown(
                        f'<div style="padding:0.45rem 0.3rem;font-size:0.88rem;'
                        f'font-weight:700;text-align:center;color:#1B5E20;'
                        f'background:#E8F5E9;border-radius:6px;border:1px solid #A5D6A7;">'
                        f'{_h_total_now} 🔒</div>',
                        unsafe_allow_html=True
                    )
                    h_f = _h_total_now
                else:
                    h_f = st.number_input("cm", value=0, step=1, min_value=0,
                                          key=f"fh_{li}", label_visibility="collapsed")
            with lc2:
                is_ac = mat_f in AC_MATERIALS_LOCK_MI
                if mat_f != "ไม่เลือก":
                    if is_ac:
                        st.markdown(
                            '<div style="padding:0.4rem 0.3rem;font-size:0.82rem;'
                            'text-align:center;color:var(--color-text-secondary);">'
                            '1.0 🔒</div>',
                            unsafe_allow_html=True
                        )
                        mi_f = 1.0
                    else:
                        mi_f = st.number_input(
                            "mi", value=1.0, step=0.1,
                            min_value=0.6, max_value=1.4,
                            key=f"fmi_{li}", label_visibility="collapsed",
                            format="%.1f"
                        )
                else:
                    mi_f = 1.0
                    st.markdown("")
            with lc3:
                if mat_f != "ไม่เลือก" and h_f > 0:
                    ai, _ = FLEX_LAYER_MATERIALS[mat_f]
                    # AC sub-layers
                    if mat_f in AC_SURFACE_MATERIALS:
                        do_sub = st.checkbox(
                            "แบ่งชั้นย่อย", key=f"fsub_{li}",
                            help="แบ่งเป็น Wearing / Binder / Base Course"
                        )
                        if do_sub:
                            sc1, sc2, sc3 = st.columns(3)
                            with sc1:
                                h_wear = st.number_input(
                                    "Wearing (cm)", value=5, step=1,
                                    min_value=0, key=f"fwear_{li}",
                                    help="มาตรฐานกรมทางหลวง: 4–7 cm (40–70 mm)"
                                )
                            with sc2:
                                h_bind = st.number_input(
                                    "Binder (cm)", value=5, step=1,
                                    min_value=0, key=f"fbind_{li}",
                                    help="มาตรฐานกรมทางหลวง: 4–8 cm (40–80 mm)"
                                )
                            with sc3:
                                h_base = st.number_input(
                                    "Base (cm)", value=7, step=1,
                                    min_value=0, key=f"fbase_{li}",
                                    help="มาตรฐานกรมทางหลวง: 7–10 cm (70–100 mm)"
                                )
                            # ── แจ้งเตือนเมื่อเกินช่วงมาตรฐาน ──
                            warn_msgs = []
                            if h_wear > 0 and not (4 <= h_wear <= 7):
                                warn_msgs.append(f"⚠️ Wearing {h_wear} cm เกินช่วงมาตรฐาน (4–7 cm)")
                            if h_bind > 0 and not (4 <= h_bind <= 8):
                                warn_msgs.append(f"⚠️ Binder {h_bind} cm เกินช่วงมาตรฐาน (4–8 cm)")
                            if h_base > 0 and not (7 <= h_base <= 10):
                                warn_msgs.append(f"⚠️ Base {h_base} cm เกินช่วงมาตรฐาน (7–10 cm)")
                            if warn_msgs:
                                st.markdown(
                                    f'<div class="result-warn" style="font-size:0.82rem;margin-top:0.3rem;">'
                                    + "<br>".join(warn_msgs) +
                                    "</div>",
                                    unsafe_allow_html=True
                                )
                            h_total = h_wear + h_bind + h_base
                            h_in_total = h_total / 2.54
                            sn_i = ai * h_in_total * mi_f
                            _cum_sn_before = cum_sn          # บันทึก ΣSN ก่อนบวกชั้นนี้
                            cum_sn += sn_i
                            layer_results.append({
                                'layer': li+1, 'material': mat_f,
                                'h_cm': h_total, 'ai': ai, 'mi': mi_f,
                                'sni': round(sn_i,3), 'cum_sn': round(cum_sn,3),
                                'sub': {'wear': h_wear, 'bind': h_bind, 'base': h_base}
                            })
                            st.markdown(
                                f'<div style="padding:0.35rem 0.5rem;font-size:0.80rem;'
                                f'font-family:monospace;background:var(--color-background-secondary);'
                                f'border-radius:6px;margin-top:0.2rem;">'
                                f'Wearing={h_wear}+Binder={h_bind}+Base={h_base} = <b>{h_total} cm</b>'
                                f' | ai={ai:.2f} | SNi={sn_i:.3f} | <b>ΣSNi={cum_sn:.3f}</b>'
                                f'</div>',
                                unsafe_allow_html=True
                            )
                            # ── #4 Smart Recommendation (AASHTO D_min) ──
                            _badge = _aashto_badge(_sn_req, _cum_sn_before, ai, mi_f, h_total)
                            if _badge:
                                st.markdown(_badge, unsafe_allow_html=True)
                        else:
                            h_in = h_f / 2.54
                            sn_i = ai * h_in * mi_f
                            _cum_sn_before = cum_sn
                            cum_sn += sn_i
                            layer_results.append({
                                'layer': li+1, 'material': mat_f,
                                'h_cm': h_f, 'ai': ai, 'mi': mi_f,
                                'sni': round(sn_i,3), 'cum_sn': round(cum_sn,3)
                            })
                            st.markdown(
                                f'<div style="padding:0.35rem 0.5rem;font-size:0.82rem;'
                                f'font-family:monospace;background:var(--color-background-secondary);'
                                f'border-radius:6px;margin-top:0.25rem;">'
                                f'<b>{h_f} cm</b> | ai={ai:.2f} | mi={mi_f:.1f} | '
                                f'SNi={sn_i:.3f} | <b>ΣSNi={cum_sn:.3f}</b>'
                                f'</div>',
                                unsafe_allow_html=True
                            )
                            # ── #4 Smart Recommendation (AASHTO D_min) ──
                            _badge = _aashto_badge(_sn_req, _cum_sn_before, ai, mi_f, h_f)
                            if _badge:
                                st.markdown(_badge, unsafe_allow_html=True)
                    else:
                        h_in = h_f / 2.54
                        sn_i = ai * h_in * mi_f
                        _cum_sn_before = cum_sn
                        cum_sn += sn_i
                        layer_results.append({
                            'layer': li+1, 'material': mat_f,
                            'h_cm': h_f, 'ai': ai, 'mi': mi_f,
                            'sni': round(sn_i,3), 'cum_sn': round(cum_sn,3)
                        })
                        st.markdown(
                            f'<div style="padding:0.35rem 0.5rem;font-size:0.82rem;'
                            f'font-family:monospace;background:var(--color-background-secondary);'
                            f'border-radius:6px;margin-top:0.25rem;">'
                            f'<b>{h_f} cm</b> | ai={ai:.2f} | mi={mi_f:.1f} | '
                            f'SNi={sn_i:.3f} | <b>ΣSNi={cum_sn:.3f}</b>'
                            f'</div>',
                            unsafe_allow_html=True
                        )
                        # ── #4 Smart Recommendation (AASHTO D_min) ──
                        _badge = _aashto_badge(_sn_req, _cum_sn_before, ai, mi_f, h_f)
                        if _badge:
                            st.markdown(_badge, unsafe_allow_html=True)
                else:
                    st.markdown("")

        st.markdown(f"""<div class="result-info" style="margin-top:0.5rem;">
            ΣSN Provided = <b>{cum_sn:.3f}</b>
        </div>""", unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

        if st.button("✅ Design Check", type="primary", key="flex_check"):
            if design_esal_f <= 0:
                st.warning("⚠️ กรุณาใส่ Design ESAL")
            else:
                sn_req = aashto_sn_required(design_esal_f, ZR_MAP[r0_fl], so_fl, pi_fl, pt_fl2, mr_fl)
                if sn_req:
                    passed = cum_sn >= sn_req
                    margin = cum_sn - sn_req
                    css    = "result-pass" if passed else "result-fail"
                    chk    = "✅ PASS" if passed else "❌ FAIL"

                    sn_ratio = cum_sn / sn_req if sn_req > 0 else 0.0
                    c1,c2,c3,c4 = st.columns(4)
                    with c1: st.markdown(f"""<div class="metric-box"><div class="val">{cum_sn:.3f}</div><div class="lbl">SN Provided</div></div>""", unsafe_allow_html=True)
                    with c2: st.markdown(f"""<div class="metric-box"><div class="val">{sn_req:.3f}</div><div class="lbl">SN Required</div></div>""", unsafe_allow_html=True)
                    with c3: st.markdown(f"""<div class="metric-box"><div class="val" style="color:{'#1B5E20' if passed else '#B71C1C'}">{margin:+.3f}</div><div class="lbl">Safety Margin</div></div>""", unsafe_allow_html=True)
                    with c4: st.markdown(f"""<div class="metric-box"><div class="val" style="color:{'#1B5E20' if sn_ratio>=1.0 else '#B71C1C'}">{sn_ratio:.3f}</div><div class="lbl">SN Ratio (≥1.0 = ผ่าน)</div></div>""", unsafe_allow_html=True)

                    st.markdown(f'<div class="{css}" style="margin-top:0.8rem;font-size:1.05rem">{chk} — SN Required = {sn_req:.3f} | SN Provided = {cum_sn:.3f}</div>', unsafe_allow_html=True)

                    ss.flex_results = {
                        'esal': design_esal_f, 'sn_req': sn_req,
                        'sn_prov': cum_sn, 'pass': passed,
                        'layers': layer_results,
                        'mr_psi': mr_fl, 'cbr': cbr_fl,
                    }
                    if layer_results:
                        st.dataframe(pd.DataFrame(layer_results), use_container_width=True, hide_index=True)

                    # ── รูปโครงสร้างชั้นทาง ──
                    if layer_results:
                        flex_layers_fig = [
                            {
                                'name':         l.get('material', ''),
                                'thickness_cm': l.get('h_cm', 0),
                                'ai':           l.get('ai', None),
                                'sni':          l.get('sni', None),
                            }
                            for l in layer_results
                        ]
                        fig_flex = draw_pavement_structure(
                            flex_layers_fig, mode="flex",
                            cbr_subgrade=cbr_fl,
                        )
                        if fig_flex:
                            st.markdown("#### 🖼️ รูปโครงสร้างชั้นทางลาดยาง")
                            st.pyplot(fig_flex, use_container_width=True)
                            ss['flex_structure_img'] = fig_to_bytes(fig_flex)
                            plt.close(fig_flex)
                else:
                    st.error("ไม่สามารถคำนวณ SN Required ได้ — ตรวจสอบ ESAL และ Mr")

# ══════════════════════════════════════════════
#  TAB 4: RIGID DESIGN (รวม K-Nomograph ไว้ด้านใน)
# ══════════════════════════════════════════════
with tab4:
    st.markdown("### 🏗️ Rigid Pavement Design — AASHTO 1993")
    st.markdown("""
    <div class="result-info">
        💡 <b>ขั้นตอน:</b> (1) กำหนด Layer Editor → ได้ DSB, ESB
        → (2) K-Value Nomograph → ได้ k_eff
        → (3) คำนวณ Slab Thickness D
    </div>
    """, unsafe_allow_html=True)

    # ── Status bar ──
    col_s1, col_s2, col_s3 = st.columns(3)
    with col_s1:
        st.markdown(status_badge('esal_rigid', 'ESAL Rigid'), unsafe_allow_html=True)
    with col_s2:
        st.markdown(status_badge('k_subgrade_pci', 'k_subgrade'), unsafe_allow_html=True)
    with col_s3:
        st.markdown(status_badge('cbr_design', 'CBR/Mr'), unsafe_allow_html=True)

    st.markdown("---")

    # ════════════════════════════════════════
    #  SECTION A: K-Value Nomograph (embedded)
    # ════════════════════════════════════════
    with st.expander("📐 SECTION A — K-Value Nomograph (Fig.3.3 & Fig.3.4)", expanded=True):
        st.markdown("""
        <div class="result-info">
            💡 กำหนด Layer ก่อน (Section B) เพื่อให้ได้ DSB และ ESB → แล้วกลับมาลาก Nomograph ที่นี่
        </div>
        """, unsafe_allow_html=True)

        sub_kinf, sub_ls = st.tabs(["📊 Composite k∞ (Fig.3.3)", "📉 Loss of Support (Fig.3.4)"])

        # ─── Sub-tab A: Composite k∞ ───
        with sub_kinf:
            if not PIL_OK:
                st.error("⚠️ ไม่พบ Pillow — ติดตั้งด้วย `pip install Pillow`")
            uploaded_k = st.file_uploader(
                "📂 Upload ภาพ Figure 3.3 (Composite k∞)",
                type=['png','jpg','jpeg'], key='uploader_kinf'
            )
            if uploaded_k is not None:
                raw_k = uploaded_k.read()
                ss['img1_original'] = raw_k
            elif ss.get('img1_original'):
                raw_k = ss['img1_original']
            else:
                raw_k = None

            if raw_k:
                img1 = PILImage.open(io.BytesIO(raw_k)).convert("RGB")
                w1, h1 = img1.size
                img1_draw = img1.copy()
                draw1 = PILDraw.Draw(img1_draw)

                col_ctrl1, col_img1 = st.columns([1, 2])

                with col_ctrl1:
                    st.markdown('<div class="card"><h4>⚙️ ปรับเส้นอ่านค่า</h4>', unsafe_allow_html=True)

                    with st.expander("1. เส้น Turning Line (เขียว)", expanded=True):
                        gx1 = st.slider("X เริ่ม", 0, w1, ss.get('gx1', 699), key="gx1")
                        gy1 = st.slider("Y เริ่ม", 0, h1, ss.get('gy1', 625), key="gy1")
                        gx2 = st.slider("X จบ",   0, w1, ss.get('gx2', 421), key="gx2")
                        gy2 = st.slider("Y จบ",   0, h1, ss.get('gy2', 348), key="gy2")
                    draw1.line([(gx1,gy1),(gx2,gy2)], fill="green", width=5)
                    slope_g = (gy2-gy1)/(gx2-gx1) if (gx2-gx1) != 0 else 0

                    with st.expander("2. พารามิเตอร์ (ส้ม/แดง/น้ำเงิน)", expanded=True):
                        start_x   = st.slider("ตำแหน่งแกน D_SB (ซ้าย)", 0, w1, ss.get('s1_sx',   int(w1*0.15)), key="s1_sx")
                        stop_y_esb= st.slider("ระดับค่า ESB (บน)",       0, h1, ss.get('s1_sy_esb',int(h1*0.10)), key="s1_sy_esb")
                        stop_y_mr = st.slider("ระดับค่า MR (ล่าง)",      0, h1, ss.get('s1_sy_mr', int(h1*0.55)), key="s1_sy_mr")
                        constrained_x = int(gx1 + (stop_y_mr - gy1) / slope_g) if slope_g != 0 else gx1

                    # วาดลูกศร
                    def _arrow(drw, start, end, color, lw=4, arrow=14):
                        drw.line([start, end], fill=color, width=lw)
                        dx = end[0]-start[0]; dy = end[1]-start[1]
                        L = math.sqrt(dx*dx+dy*dy)
                        if L > 0:
                            dx/=L; dy/=L; px=-dy; py=dx
                            bx=end[0]-arrow*dx; by=end[1]-arrow*dy
                            drw.polygon([(end[0],end[1]),(int(bx+arrow*0.5*px),int(by+arrow*0.5*py)),
                                         (int(bx-arrow*0.5*px),int(by-arrow*0.5*py))], fill=color)

                    _arrow(draw1, (start_x, stop_y_esb), (constrained_x, stop_y_esb), "orange")
                    _arrow(draw1, (start_x, stop_y_esb), (start_x, stop_y_mr),         "red")
                    _arrow(draw1, (start_x, stop_y_mr),  (constrained_x, stop_y_mr),   "darkblue")
                    _arrow(draw1, (constrained_x, stop_y_mr), (constrained_x, stop_y_esb), "blue")
                    r=8
                    draw1.ellipse([(constrained_x-r,stop_y_mr-r),(constrained_x+r,stop_y_mr+r)],
                                   fill="black", outline="white")

                    st.markdown('</div>', unsafe_allow_html=True)
                    st.markdown('<div class="card"><h4>📝 บันทึกค่าที่อ่านได้</h4>', unsafe_allow_html=True)

                    # ── auto-fill จาก session state ──
                    mr_auto_k  = int(ss.mr_subgrade_psi) if ss.mr_subgrade_psi else 7000
                    esb_auto_k = int(ss.get('layer_esb_psi', 50000))
                    dsb_auto_k = float(ss.get('layer_dsb_in', 6.0))

                    if ss.mr_subgrade_psi:
                        st.markdown(f'<div class="badge-ready">📊 MR จาก Tab CBR = {mr_auto_k:,} psi</div>', unsafe_allow_html=True)
                    if ss.get('layer_esb_psi'):
                        st.markdown(f'<div class="badge-ready">📐 ESB จาก Rigid Layers = {esb_auto_k:,} psi</div>', unsafe_allow_html=True)
                    if ss.get('layer_dsb_in'):
                        st.markdown(f'<div class="badge-ready">📏 DSB จาก Rigid Layers = {dsb_auto_k:.1f} in</div>', unsafe_allow_html=True)

                    mr_val  = st.number_input("MR (psi) — แก้ไขได้",     value=mr_auto_k,  step=500,  key="nomo_mr")
                    esb_val = st.number_input("ESB (psi) — แก้ไขได้",    value=esb_auto_k, step=1000, key="nomo_esb")
                    dsb_val = st.number_input("DSB (inches) — แก้ไขได้", value=dsb_auto_k, step=0.5,  key="nomo_dsb")
                    k_inf_read = st.number_input("k∞ ที่อ่านได้ (pci)",  value=int(ss.get('nomo_k_inf', 400)), step=10, key="nomo_k_inf")

                    if st.button("✅ บันทึก k∞", type="primary", key="save_kinf"):
                        ss.k_inf = float(k_inf_read)
                        buf1 = io.BytesIO()
                        img1_draw.save(buf1, format='PNG')
                        ss['nomograph_img_k'] = buf1.getvalue()
                        ss['img1_original']   = raw_k
                        st.markdown(f'<div class="result-pass">k∞ = <b>{k_inf_read} pci</b> บันทึกแล้ว → ใช้ใน Tab Loss of Support และ Rigid Design</div>', unsafe_allow_html=True)

                    st.markdown('</div>', unsafe_allow_html=True)

                with col_img1:
                    if ss.get('k_inf'):
                        st.markdown(f'<div class="result-pass" style="margin-bottom:0.5rem;">k∞ ปัจจุบัน = <b>{ss.k_inf:.0f} pci</b></div>', unsafe_allow_html=True)
                    st.image(img1_draw, caption="Composite k∞ Nomograph (AASHTO 1993 Fig.3.3)", use_container_width=True)

                    # Save annotated image
                    buf1b = io.BytesIO()
                    img1_draw.save(buf1b, format='PNG')
                    ss['nomograph_img_k'] = buf1b.getvalue()
            else:
                st.markdown("""
                <div class="result-warn">
                    👆 กรุณา Upload รูป <b>Figure 3.3</b> จาก AASHTO 1993 Guide<br>
                    <small>รองรับ PNG, JPG, JPEG</small>
                </div>
                """, unsafe_allow_html=True)
                # Fallback: กรอก k∞ ตรงๆ
                st.markdown('<div class="card"><h4>📝 หรือกรอกค่า k∞ โดยตรง</h4>', unsafe_allow_html=True)
                k_inf_manual = st.number_input("k∞ (pci)", value=float(ss.k_inf) if ss.k_inf else 200.0,
                                                step=10.0, min_value=10.0, key="k_inf_manual")
                if st.button("✅ ใช้ค่านี้", key="use_kinf_manual"):
                    ss.k_inf = k_inf_manual
                    st.success(f"✅ k∞ = {k_inf_manual:.0f} pci บันทึกแล้ว")
                st.markdown('</div>', unsafe_allow_html=True)

        # ─────────────────────────────────────────
        #  Sub-tab B: Loss of Support
        # ─────────────────────────────────────────
        with sub_ls:
            from PIL import Image as PILImage, ImageDraw as PILDraw

            st.markdown(f'<div class="result-info">k∞ จาก Tab k∞ = <b>{ss.k_inf:.0f} pci</b></div>', unsafe_allow_html=True)

            uploaded_ls = st.file_uploader(
                "📂 Upload ภาพ Figure 3.4 (Loss of Support)",
                type=['png','jpg','jpeg'], key='uploader_ls'
            )
            if uploaded_ls is not None:
                raw_ls = uploaded_ls.read()
                ss['img2_original'] = raw_ls
            elif ss.get('img2_original'):
                raw_ls = ss['img2_original']
            else:
                raw_ls = None

            if raw_ls:
                img2 = PILImage.open(io.BytesIO(raw_ls)).convert("RGB")
                w2, h2 = img2.size
                img2_draw = img2.copy()
                draw2 = PILDraw.Draw(img2_draw)

                col_ctrl2, col_img2 = st.columns([1, 2])

                with col_ctrl2:
                    st.markdown('<div class="card"><h4>⚙️ กำหนดเส้นกราฟ</h4>', unsafe_allow_html=True)

                    # LS presets จาก Rigid_Pavement__3_.py
                    LS_PRESETS_LOCAL = {
                        0.0: (138, 715, 753, 84),
                        0.5: (129, 728, 908,  0),
                        1.0: (150, 718, 903, 84),
                        1.5: (153, 721, 928,138),
                        2.0: (164, 718, 929,220),
                        3.0: (212, 719, 929,328),
                    }
                    ls_opts = [0.0, 0.5, 1.0, 1.5, 2.0, 3.0]
                    cur_ls  = ss.get('ls_select_box', 1.0)
                    def_idx = ls_opts.index(cur_ls) if cur_ls in ls_opts else 2
                    ls_sel  = st.selectbox("เลือกค่า LS", ls_opts, index=def_idx, key="ls_select_box")

                    # Auto-preset เมื่อ LS เปลี่ยน
                    if ss.get('last_ls_select') != ls_sel:
                        ss['last_ls_select'] = ls_sel
                        coords = LS_PRESETS_LOCAL.get(ls_sel, (150,718,903,84))
                        ss['_ls_px1'], ss['_ls_py1'] = coords[0], coords[1]
                        ss['_ls_px2'], ss['_ls_py2'] = coords[2], coords[3]

                    with st.expander("ปรับแต่งเส้น LS ละเอียด", expanded=False):
                        ls_x1 = st.slider("จุดเริ่ม X", -100, w2+100, ss.get('_ls_px1', 150), key="_ls_x1")
                        ls_y1 = st.slider("จุดเริ่ม Y", -100, h2+100, ss.get('_ls_py1', 718), key="_ls_y1")
                        ls_x2 = st.slider("จุดจบ X",   -100, w2+100, ss.get('_ls_px2', 903), key="_ls_x2")
                        ls_y2 = st.slider("จุดจบ Y",   -100, h2+100, ss.get('_ls_py2',  84), key="_ls_y2")

                    draw2.line([(ls_x1,ls_y1),(ls_x2,ls_y2)], fill="red", width=6)
                    m_red = (ls_y2-ls_y1)/(ls_x2-ls_x1) if ls_x2-ls_x1 != 0 else None
                    c_red = ls_y1 - m_red*ls_x1 if m_red else 0

                    st.markdown("---")
                    with st.expander("📍 ตั้งค่าตำแหน่งแกนกราฟ", expanded=True):
                        axis_left_x  = st.number_input("ตำแหน่งแกน Y (ซ้ายสุด)", value=ss.get('axis_left',100),     step=5, key="axis_left")
                        axis_bottom_y= st.number_input("ตำแหน่งแกน X (ล่างสุด)", value=ss.get('axis_bottom',h2-50), step=5, key="axis_bottom")

                    st.caption(f"k∞ จาก Tab ก่อนหน้า = {ss.k_inf:.0f} pci")
                    k_pos_x   = st.slider("ตำแหน่ง k∞ บนแกน X", 0, w2, ss.get('k_pos_x', w2//2), key="k_pos_x")
                    intersect_y = int(m_red*k_pos_x + c_red) if m_red else h2//2

                    draw2.line([(k_pos_x,axis_bottom_y),(k_pos_x,intersect_y)], fill="blue", width=5)

                    def _arrow2(drw, start, end, color, lw=5, arrow=14):
                        drw.line([start, end], fill=color, width=lw)
                        dx=end[0]-start[0]; dy=end[1]-start[1]
                        L=math.sqrt(dx*dx+dy*dy)
                        if L>0:
                            dx/=L; dy/=L; px=-dy; py=dx
                            bx=end[0]-arrow*dx; by=end[1]-arrow*dy
                            drw.polygon([(end[0],end[1]),(int(bx+arrow*0.5*px),int(by+arrow*0.5*py)),
                                         (int(bx-arrow*0.5*px),int(by-arrow*0.5*py))], fill=color)

                    _arrow2(draw2, (k_pos_x, intersect_y), (int(axis_left_x), intersect_y), "blue")
                    draw2.ellipse([(k_pos_x-8,intersect_y-8),(k_pos_x+8,intersect_y+8)],
                                   fill="black", outline="white", width=2)

                    st.markdown('</div>', unsafe_allow_html=True)
                    st.markdown('<div class="card"><h4>📝 บันทึกผลลัพธ์</h4>', unsafe_allow_html=True)

                    k_corr_val = st.number_input(
                        "k_eff Corrected (pci)",
                        value=int(ss.get('k_corrected', max(10, ss.k_inf - 100)) if ss.k_inf else 200),
                        step=10, min_value=10, key="k_corr_input"
                    )
                    if k_corr_val > ss.k_inf and ss.k_inf > 0:
                        st.warning(f"⚠️ k_eff ({k_corr_val:.0f}) ต้องไม่เกิน k∞ ({ss.k_inf:.0f} pci)")
                        k_corr_val = ss.k_inf

                    if st.button("✅ บันทึก k_eff", type="primary", key="save_keff"):
                        ss.k_corrected = float(k_corr_val)
                        ss.ls_value    = ls_sel
                        buf2 = io.BytesIO()
                        img2_draw.save(buf2, format='PNG')
                        ss['nomograph_img_ls'] = buf2.getvalue()
                        ss['img2_original']    = raw_ls
                        st.markdown(f'<div class="result-pass">k_eff = <b>{k_corr_val} pci</b> (LS={ls_sel}) บันทึกแล้ว → ใช้ใน Tab Rigid Design</div>', unsafe_allow_html=True)

                    st.markdown('</div>', unsafe_allow_html=True)

                with col_img2:
                    if ss.get('k_corrected'):
                        st.markdown(f'<div class="result-pass" style="margin-bottom:0.5rem;">k_eff = <b>{ss.k_corrected:.0f} pci</b>  |  LS = {ss.ls_value}</div>', unsafe_allow_html=True)
                    st.image(img2_draw, caption=f"Loss of Support (LS={ls_sel}) — AASHTO 1993 Fig.3.4", use_container_width=True)

                    buf2b = io.BytesIO()
                    img2_draw.save(buf2b, format='PNG')
                    ss['nomograph_img_ls'] = buf2b.getvalue()
            else:
                st.markdown("""
                <div class="result-warn">
                    👆 กรุณา Upload รูป <b>Figure 3.4</b> จาก AASHTO 1993 Guide<br>
                    <small>รองรับ PNG, JPG, JPEG</small>
                </div>
                """, unsafe_allow_html=True)
                st.markdown('<div class="card"><h4>📝 หรือกรอกค่า k_eff โดยตรง</h4>', unsafe_allow_html=True)
                ls_opts2 = [0.0, 0.5, 1.0, 1.5, 2.0, 3.0]
                ls_man   = st.select_slider("Loss of Support (LS)", ls_opts2,
                                            value=ss.ls_value if ss.ls_value in ls_opts2 else 1.0,
                                            key="ls_manual")
                k_eff_man= st.number_input("k_eff (pci)", value=float(ss.k_corrected) if ss.k_corrected else 200.0,
                                            step=10.0, min_value=10.0, key="k_eff_manual")
                if st.button("✅ ใช้ค่านี้", key="use_keff_manual"):
                    ss.k_corrected = k_eff_man
                    ss.ls_value    = ls_man
                    st.success(f"✅ k_eff = {k_eff_man:.0f} pci บันทึกแล้ว")
                st.markdown('</div>', unsafe_allow_html=True)

    # ════════════════════════════════════════
    #  SECTION B: Rigid Design Parameters
    # ════════════════════════════════════════
    st.markdown("---")
    st.markdown("### ⚙️ SECTION B — Design Parameters & Slab Thickness")

    # ── Shared design parameters ──
    st.markdown('<div class="card"><h4>⚙️ พารามิเตอร์ร่วม (ใช้กับทุก Type)</h4>', unsafe_allow_html=True)
    rp1, rp2, rp3, rp4, rp5, rp6 = st.columns(6)
    with rp1:
        fc_cube = st.number_input("f'c (ksc)", value=350, step=10, min_value=200, key="fc_cube")
    with rp2:
        fc_cyl  = 0.8 * fc_cube
        fc_psi  = fc_cyl * 14.223
        ec_psi  = 57000 * math.sqrt(fc_psi)
        sc_auto = min(600, 10.0 * math.sqrt(fc_psi))
        sc_inp  = st.number_input("Sc (psi)", value=int(sc_auto), step=10,
                                  min_value=100, max_value=700, key="sc_inp")
    with rp3:
        r0_rig = st.selectbox("Reliability R0 (%)", list(ZR_MAP.keys()),
                               index=10, key="r0_rig")
        zr_rig = ZR_MAP[r0_rig]
    with rp4:
        so_rig = st.number_input("So", value=0.35, step=0.01,
                                 min_value=0.2, max_value=0.5, key="so_rig")
    with rp5:
        pi_rig = st.number_input("Pi", value=4.5, step=0.1, key="pi_rig")
    with rp6:
        use_pt_global_rig = st.checkbox("ใช้ Pt Global", value=ss.get('use_pt_global_rig', True), key="use_pt_global_rig")
        if use_pt_global_rig:
            pt_rig2 = float(ss.get('pt_global', 2.5))
            st.caption(f"Pt = {pt_rig2} (Global)")
        else:
            pt_rig2 = st.number_input("Pt (Override)", value=float(ss.get('pt_rig2_override', ss.get('pt_global', 2.5))),
                                       step=0.1, min_value=2.0, max_value=3.0, key="pt_rig2_override")

    st.markdown(
        f"Ec = **{ec_psi:,.0f} psi** &nbsp;|&nbsp; "
        f"f'c cylinder = **{fc_cyl:.0f} ksc** &nbsp;|&nbsp; "
        f"ZR = **{zr_rig}**"
    )
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Sub-tabs JPCP/JRCP | CRCP ──
    sub_jpcp_jrcp, sub_crcp = st.tabs(["🟦 JPCP / JRCP", "🟥 CRCP"])

    def rigid_design_panel(ptype, tab_key):
        """
        Panel ออกแบบ Rigid Pavement แต่ละประเภท
        - Layer editor: วัสดุ + หนา + E_MPa (auto-fill แก้ไขได้)
        - คำนวณ E_eq, DSB → ส่งไป session state
        - ดึง ESAL, k_eff อัตโนมัติ แก้ไขได้
        """
        j_key = "JPCP/JRCP" if ptype in ("JPCP", "JRCP") else "CRCP"
        j_default  = J_VALUES[j_key]
        mat_opts   = list(RIGID_LAYER_MATERIALS.keys())

        col_rd_l, col_rd_r = st.columns([1.1, 1])

        # ════════════════════════════════
        #  คอลัมน์ซ้าย — Layer Editor
        # ════════════════════════════════
        with col_rd_l:
            st.markdown(
                f'<div class="card"><h4>🔩 ชั้นโครงสร้าง – {ptype}</h4>',
                unsafe_allow_html=True
            )

            # Header row
            h0, h1, h2, h3 = st.columns([3, 1.2, 1.5, 0.5])
            h0.markdown("**วัสดุ**")
            h1.markdown("**หนา (cm)**")
            h2.markdown("**E (MPa)**")
            h3.markdown("")

            layer_r    = []
            total_h_cm = 0.0
            e_eq_psi   = 0.0

            for li in range(5):
                lc_a, lc_b, lc_c, lc_d = st.columns([3, 1.2, 1.5, 0.5])

                with lc_a:
                    mat_r = st.selectbox(
                        f"ชั้น {li+1}", mat_opts,
                        key=f"rmat_{tab_key}_{li}",
                        label_visibility="collapsed"
                    )
                with lc_b:
                    h_r = st.number_input(
                        "cm", value=0, step=1, min_value=0,
                        key=f"rh_{tab_key}_{li}",
                        label_visibility="collapsed"
                    )
                with lc_c:
                    prev_mat_key = f"_prev_mat_{tab_key}_{li}"
                    e_key        = f"re_{tab_key}_{li}"
                    e_default    = RIGID_LAYER_E_DEFAULT.get(mat_r, 100) if mat_r != "ไม่เลือก" else 0
                    if ss.get(prev_mat_key) != mat_r:
                        ss[prev_mat_key] = mat_r
                        ss[e_key]        = e_default
                    e_mpa = st.number_input(
                        "MPa", value=ss.get(e_key, e_default),
                        step=50, min_value=0,
                        key=e_key,
                        label_visibility="collapsed",
                        disabled=(mat_r == "ไม่เลือก" or h_r == 0)
                    )
                with lc_d:
                    if mat_r != "ไม่เลือก" and h_r > 0:
                        st.markdown("✅")

                if mat_r != "ไม่เลือก" and h_r > 0 and e_mpa > 0:
                    layer_r.append({
                        "name": mat_r,
                        "thickness_cm": h_r,
                        "E_MPa": e_mpa
                    })
                    total_h_cm += h_r

            # ── คำนวณ E_equivalent (Odemark) ──
            if layer_r:
                total_valid_cm = sum(l["thickness_cm"] for l in layer_r)
                sum_h_e_cbrt   = sum(
                    l["thickness_cm"] * (l["E_MPa"] ** (1/3))
                    for l in layer_r
                )
                e_eq_mpa = (sum_h_e_cbrt / total_valid_cm) ** 3 if total_valid_cm > 0 else 0
                e_eq_psi = e_eq_mpa * 145.038
                dsb_in   = total_valid_cm / 2.54

                # ── ส่งค่าไป session state → Section A K-Nomograph ดึงได้ ──
                ss['layer_esb_psi'] = int(e_eq_psi)
                ss['layer_dsb_in']  = round(dsb_in, 2)

                st.markdown(f"""
                <div class="result-info" style="margin-top:0.5rem;font-size:0.88rem;">
                    📐 รวมชั้นทาง = <b>{total_valid_cm:.0f} cm</b>
                    ({dsb_in:.1f} in) &nbsp;|&nbsp;
                    E_eq = <b>{e_eq_psi:,.0f} psi</b>
                    ({e_eq_mpa:.1f} MPa)<br>
                    <small>→ ส่งค่าไป Tab K-Value อัตโนมัติ (ESB={e_eq_psi:,.0f} psi, DSB={dsb_in:.1f} in)</small>
                </div>""", unsafe_allow_html=True)
            else:
                st.markdown(
                    '<div class="result-warn">กรุณากรอกชั้นทางอย่างน้อย 1 ชั้น</div>',
                    unsafe_allow_html=True
                )

            st.markdown('</div>', unsafe_allow_html=True)

        # ════════════════════════════════
        #  คอลัมน์ขวา — Parameters + Check
        # ════════════════════════════════
        with col_rd_r:
            st.markdown(
                f'<div class="card"><h4>⚙️ พารามิเตอร์ – {ptype}</h4>',
                unsafe_allow_html=True
            )

            pr1, pr2 = st.columns(2)
            with pr1:
                j_hint = "default 2.8, ช่วง 2.5–3.2 (กรมทางหลวง)" if ptype != "CRCP" else "default 2.6, ช่วง 2.5–3.2 (กรมทางหลวง)"
                j_val  = st.number_input(
                    f"J ({ptype})", value=j_default,
                    step=0.1, min_value=1.0, max_value=5.0,
                    key=f"j_{tab_key}",
                    help=j_hint
                )
                cd_val = st.number_input(
                    "Cd (Drainage)", value=1.0, step=0.05,
                    min_value=0.5, max_value=1.25,
                    key=f"cd_{tab_key}"
                )
            with pr2:
                d_sel = st.selectbox(
                    "Slab Thickness",
                    SLAB_THICKNESSES,
                    index=2,
                    format_func=lambda d: SLAB_LABELS[SLAB_THICKNESSES.index(d)],
                    key=f"d_{tab_key}"
                )

            # ── ESAL: ดึงอัตโนมัติจาก Tab 1 ──
            # normalize key: ลอง int และ float เพื่อ match key ใน esal_rigid
            _esal_r = ss.esal_rigid or {}
            esal_auto = int(_esal_r.get(d_sel,
                            _esal_r.get(int(d_sel),
                            _esal_r.get(float(d_sel), 0))))
            if esal_auto > 0:
                st.markdown(
                    f'<div class="badge-ready">📊 ESAL จาก Tab 1 (Slab {d_sel} cm) = {esal_auto:,}</div>',
                    unsafe_allow_html=True
                )
            else:
                st.markdown(
                    '<div class="badge-wait">⚠️ ยังไม่มี ESAL — คำนวณใน Tab 1 ก่อน หรือกรอกเอง</div>',
                    unsafe_allow_html=True
                )

            # อัปเดต session state เมื่อ ESAL ใน Tab 1 เปลี่ยน
            _w18_key     = f"w18_{tab_key}"
            _esal_key    = f"_prev_esal_{tab_key}_{d_sel}"
            if esal_auto > 0 and ss.get(_esal_key) != esal_auto:
                ss[_esal_key] = esal_auto
                ss[_w18_key]  = esal_auto

            w18_req = st.number_input(
                "W18 Design (ESAL) — แก้ไขได้",
                value=ss.get(_w18_key, esal_auto),
                step=100000, min_value=0,
                key=_w18_key
            )

            # ── k_eff: ดึงอัตโนมัติจาก Tab K-Value ──
            k_eff_auto = float(ss.k_corrected) if ss.k_corrected else 0.0
            if k_eff_auto > 0:
                st.markdown(
                    f'<div class="badge-ready">📐 k_eff จาก Tab K-Value = {k_eff_auto:.0f} pci</div>',
                    unsafe_allow_html=True
                )
            else:
                st.markdown(
                    '<div class="badge-wait">⚠️ ยังไม่มี k_eff — คำนวณใน Tab K-Value ก่อน หรือกรอกเอง</div>',
                    unsafe_allow_html=True
                )

            k_eff_inp = st.number_input(
                "k_eff (pci) — แก้ไขได้",
                value=k_eff_auto if k_eff_auto > 0 else 200.0,
                step=10.0, min_value=10.0,
                key=f"keff_{tab_key}"
            )

            st.markdown('</div>', unsafe_allow_html=True)

            # ── Design Check ──
            if st.button(f"✅ Design Check – {ptype}",
                          type="primary", key=f"dc_{tab_key}"):
                if w18_req <= 0:
                    st.warning("⚠️ W18 Design = 0 — กรุณาใส่ ESAL หรือคำนวณใน Tab 1")
                else:
                    w18_cap = aashto_rigid_w18(
                        d_sel, pi_rig, pt_rig2, zr_rig, so_rig,
                        sc_inp, cd_val, j_val, ec_psi, k_eff_inp
                    )
                    if w18_cap is None:
                        st.error("ไม่สามารถคำนวณได้ — ตรวจสอบพารามิเตอร์")
                    else:
                        passed = w18_cap >= w18_req
                        margin = (w18_cap / w18_req - 1) * 100 if w18_req > 0 else float('inf')
                        css    = "result-pass" if passed else "result-fail"
                        chk    = "✅ PASS" if passed else "❌ FAIL"

                        w18_ratio = w18_cap / w18_req if w18_req > 0 else 0.0
                        c1, c2, c3, c4 = st.columns(4)
                        with c1:
                            st.markdown(f"""<div class="metric-box">
                                <div class="val">{w18_cap:,.0f}</div>
                                <div class="lbl">W18 Capacity</div>
                            </div>""", unsafe_allow_html=True)
                        with c2:
                            st.markdown(f"""<div class="metric-box">
                                <div class="val">{w18_req:,.0f}</div>
                                <div class="lbl">W18 Required</div>
                            </div>""", unsafe_allow_html=True)
                        with c3:
                            color = '#1B5E20' if passed else '#B71C1C'
                            st.markdown(f"""<div class="metric-box">
                                <div class="val" style="color:{color}">{margin:+.1f}%</div>
                                <div class="lbl">Safety Margin</div>
                            </div>""", unsafe_allow_html=True)
                        with c4:
                            ratio_color = '#1B5E20' if w18_ratio >= 1.0 else '#B71C1C'
                            st.markdown(f"""<div class="metric-box">
                                <div class="val" style="color:{ratio_color}">{w18_ratio:.3f}</div>
                                <div class="lbl">W18 Ratio (≥1.0 = ผ่าน)</div>
                            </div>""", unsafe_allow_html=True)

                        st.markdown(f"""
                        <div class="{css}" style="margin-top:0.8rem;font-size:1rem;">
                            <b>{chk}</b><br>
                            Slab {d_sel} cm &nbsp;|&nbsp;
                            k_eff = {k_eff_inp:.0f} pci &nbsp;|&nbsp;
                            J = {j_val} &nbsp;|&nbsp;
                            Cd = {cd_val}<br>
                            Ec = {ec_psi:,.0f} psi &nbsp;|&nbsp;
                            Sc = {sc_inp} psi
                        </div>""", unsafe_allow_html=True)

                        # ── k_eff Optimum (W18 Ratio = 1.0) ──
                        try:
                            def _w18_at_k(k_try):
                                return aashto_rigid_w18(
                                    d_sel, pi_rig, pt_rig2, zr_rig, so_rig,
                                    sc_inp, cd_val, j_val, ec_psi, k_try
                                )
                            # หา k_eff ที่ทำให้ W18_capacity = W18_required
                            def _eq_k(k_try):
                                w = _w18_at_k(k_try)
                                return (w - w18_req) if w else -w18_req
                            k_opt = None
                            if _eq_k(10) * _eq_k(1000) < 0:
                                k_opt = _brentq(_eq_k, 10, 1000, xtol=0.1)

                            if k_opt is not None:
                                diff_pci = k_eff_inp - k_opt
                                diff_pct = (diff_pci / k_opt) * 100
                                opt_css  = "result-info"
                                st.markdown(f"""
                                <div class="{opt_css}" style="margin-top:0.6rem;font-size:0.92rem;">
                                    🎯 <b>k_eff Optimum (W18 Ratio = 1.0 พอดี) = {k_opt:.0f} pci</b><br>
                                    <small>ใช้อยู่ {k_eff_inp:.0f} pci
                                    — เผื่อไว้ {diff_pci:+.0f} pci ({diff_pct:+.1f}%)</small>
                                </div>""", unsafe_allow_html=True)

                                # ── ตาราง Sensitivity k_eff ──
                                k_base   = [100, 200, 300, 400, 500,
                                            600, 700, 800, 900, 1000]
                                k_steps  = sorted(set(
                                    k_base
                                    + [int(round(k_opt / 10) * 10)]
                                    + [int(round(k_eff_inp / 10) * 10)]
                                ))
                                k_steps = [k for k in k_steps if 10 <= k <= 1000]
                                rows_k = []
                                for kv in k_steps:
                                    wc = _w18_at_k(kv)
                                    if wc is None: continue
                                    r  = wc / w18_req if w18_req > 0 else 0
                                    is_opt    = abs(kv - round(k_opt))  <= 5
                                    is_cur    = abs(kv - round(k_eff_inp)) <= 5
                                    tag = "🎯 Optimum" if is_opt else ("← ใช้อยู่" if is_cur else "")
                                    rows_k.append({
                                        "k_eff (pci)": kv,
                                        "W18 Capacity": f"{wc:,.0f}",
                                        "W18 Ratio": f"{r:.3f}",
                                        "สถานะ": ("✅ PASS" if r >= 1.0 else "❌ FAIL") + (f"  {tag}" if tag else ""),
                                    })
                                st.markdown("**📊 Sensitivity: W18 Ratio ตาม k_eff**")
                                st.dataframe(
                                    pd.DataFrame(rows_k),
                                    use_container_width=True, hide_index=True
                                )
                            else:
                                st.markdown("""
                                <div class="result-warn" style="margin-top:0.6rem;font-size:0.88rem;">
                                    ⚠️ ไม่พบ k_eff Optimum ในช่วง 10–5,000 pci
                                    (W18 Required อาจสูงหรือต่ำเกินช่วง)
                                </div>""", unsafe_allow_html=True)
                        except Exception:
                            pass

                        # ── บันทึกผลลัพธ์ ──
                        if not isinstance(ss.rigid_results, dict):
                            ss.rigid_results = {}
                        ss.rigid_results[ptype] = {
                            'd_cm':    d_sel,
                            'k_eff':   k_eff_inp,
                            'fc':      fc_cube,
                            'sc':      sc_inp,
                            'j':       j_val,
                            'cd':      cd_val,
                            'w18_cap': w18_cap,
                            'w18_req': w18_req,
                            'pass':    passed,
                            'margin':  margin,
                            'layers':  layer_r,
                            'e_eq_psi': e_eq_psi,
                        }

                        # ── แสดงตารางชั้นทาง ──
                        if layer_r:
                            st.markdown("#### 📋 ชั้นโครงสร้างทาง")
                            df_lr = pd.DataFrame([{
                                "ชั้น": i+1,
                                "วัสดุ": l["name"],
                                "หนา (cm)": l["thickness_cm"],
                                "E (MPa)": l["E_MPa"],
                            } for i, l in enumerate(layer_r)])
                            df_lr["h×E^(1/3)"] = df_lr.apply(
                                lambda r: f"{r['หนา (cm)'] * r['E (MPa)']**( 1/3):.1f}", axis=1
                            )
                            st.dataframe(df_lr, use_container_width=True, hide_index=True)

                            if e_eq_psi > 0:
                                st.markdown(f"""
                                <div class="result-info" style="font-size:0.88rem;">
                                    E_equivalent = <b>{e_eq_psi:,.0f} psi</b>
                                    ({e_eq_psi/145.038:.1f} MPa) &nbsp;|&nbsp;
                                    DSB = <b>{total_h_cm/2.54:.1f} in</b>
                                </div>""", unsafe_allow_html=True)

                            # ── รูปโครงสร้างชั้นทาง ──
                            rigid_layers_for_fig = [
                                {"name": l["name"],
                                 "thickness_cm": l["thickness_cm"],
                                 "E_MPa": l["E_MPa"]}
                                for l in layer_r
                            ]
                            fig_rig = draw_pavement_structure(
                                rigid_layers_for_fig, mode="rigid",
                                cbr_subgrade=float(ss.cbr_design or 3.0),
                                d_concrete_cm=d_sel,
                                ptype=ptype,
                            )
                            if fig_rig:
                                st.markdown(f"#### 🖼️ รูปโครงสร้างชั้นทางคอนกรีต ({ptype})")
                                st.pyplot(fig_rig, use_container_width=True)
                                ss[f'rigid_structure_img_{ptype}'] = fig_to_bytes(fig_rig)
                                plt.close(fig_rig)

    # ── เรียก panel แต่ละประเภท ──
    with sub_jpcp_jrcp:
        st.markdown("""
        <div class="result-info">
            💡 <b>JPCP และ JRCP</b> ใช้สมการเดียวกัน — ต่างกันเฉพาะงานก่อสร้าง
            (JPCP = ไม่มีเหล็กตะแกรง | JRCP = มีเหล็กตะแกรง)
        </div>
        """, unsafe_allow_html=True)
        rigid_design_panel("JPCP", "jpcp_jrcp")
    with sub_crcp:
        rigid_design_panel("CRCP", "crcp")

# ══════════════════════════════════════════════
#  TAB 5: REPORT & SAVE
# ══════════════════════════════════════════════
with tab5:
    st.markdown("### 📄 Report & Save — Word / JSON")

    if not DOCX_OK:
        st.warning("⚠️ ไม่พบ python-docx — ติดตั้งด้วย `pip install python-docx`")

    st.markdown('<div class="card"><h4>📋 สถานะข้อมูลแต่ละส่วน</h4>', unsafe_allow_html=True)
    sc1,sc2,sc3,sc4,sc5 = st.columns(5)
    with sc1: st.markdown(status_badge('esal_rigid','ESAL Rigid'), unsafe_allow_html=True)
    with sc2: st.markdown(status_badge('cbr_values','CBR Analysis'), unsafe_allow_html=True)
    with sc3: st.markdown(status_badge('flex_results','Flexible Design'), unsafe_allow_html=True)
    with sc4: st.markdown(status_badge('k_corrected','K-Value'), unsafe_allow_html=True)
    with sc5: st.markdown(status_badge('rigid_results','Rigid Design'), unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("#### 📥 Download รายงาน Word")

    # Project info
    col_pi1, col_pi2 = st.columns(2)
    with col_pi1: proj_name = st.text_input("ชื่อโครงการ", value="", placeholder="กรอกชื่อโครงการ...", key="proj_name")
    with col_pi2: eng_name  = st.text_input("ผู้ออกแบบ", value="รศ.ดร.อิทธิพล มีผล", key="eng_name")

    st.markdown("**เลือกส่วนที่ต้องการ Report:**")
    r1,r2,r3,r4,r5 = st.columns(5)
    with r1: chk_esal  = st.checkbox("🚛 ESAL", value=True)
    with r2: chk_cbr   = st.checkbox("📊 CBR Analysis", value=True)
    with r3: chk_flex  = st.checkbox("🔧 Flexible Design", value=True)
    with r4: chk_kval  = st.checkbox("📐 K-Value", value=True)
    with r5: chk_rigid = st.checkbox("🏗️ Rigid Design", value=True)

    col_dl1, col_dl2 = st.columns(2)

    def get_ss_dict():
        return {k: ss.get(k) for k in [
            'esal_rigid','esal_flex','ldf','ddf','pt_rigid','pt_flex',
            'cbr_values','cbr_percentile','cbr_design','mr_subgrade_psi',
            'k_subgrade_pci','flex_results','k_inf','k_corrected','ls_value',
            'nomograph_img_k','nomograph_img_ls','rigid_results'
        ]}

    with col_dl1:
        st.markdown("**แยกส่วน:**")
        ss_d = get_ss_dict()

        if chk_esal and ss.esal_rigid:
            b = build_report_esal(ss_d)
            if b: st.download_button("📥 ESAL Report", b, "ESAL_Report.docx",
                                     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                     use_container_width=True)

        if chk_cbr and ss.cbr_values:
            b = build_report_cbr(ss_d)
            if b: st.download_button("📥 CBR Report", b, "CBR_Report.docx",
                                     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                     use_container_width=True)

        if chk_flex and ss.flex_results:
            b = build_report_flexible(ss_d)
            if b: st.download_button("📥 Flexible Report", b, "Flexible_Report.docx",
                                     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                     use_container_width=True)

        if chk_kval and ss.k_corrected:
            b = build_report_kvalue(ss_d)
            if b: st.download_button("📥 K-Value Report", b, "KValue_Report.docx",
                                     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                     use_container_width=True)

        if chk_rigid and ss.rigid_results:
            b = build_report_rigid(ss_d)
            if b: st.download_button("📥 Rigid Report (พื้นฐาน)", b, "Rigid_Report.docx",
                                     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                     use_container_width=True)
            # ── รายงานแบบที่ปรึกษา ──
            st.markdown("---")
            st.markdown("**📋 รายงาน Rigid แบบที่ปรึกษา**")
            r_col1, r_col2 = st.columns(2)
            with r_col1:
                sec_pfx = st.text_input("เลขหัวข้อ", value="4.5", key="rpt_sec_pfx")
                fig_pfx = st.text_input("Prefix รูป", value="4-",  key="rpt_fig_pfx")
            with r_col2:
                fig_start = st.number_input("เลขรูปเริ่มต้น", value=1, min_value=1, step=1, key="rpt_fig_start2")
                proj_rpt  = st.text_input("ชื่อโครงการ (รายงาน)", value=ss.get('proj_name',''), key="rpt_proj2")
            inc_j = st.checkbox("รวม JPCP/JRCP", value=True,  key="rpt_inc_jpcp")
            inc_c = st.checkbox("รวม CRCP",       value=True,  key="rpt_inc_crcp")
            inc_s = st.checkbox("รวมบทสรุป",      value=True,  key="rpt_inc_sum")
            if st.button("📄 สร้างรายงานแบบที่ปรึกษา", type="primary",
                         use_container_width=True, key="rpt_rigid_pro"):
                try:
                    from report_rigid import build_rigid_word_report
                    b_pro = build_rigid_word_report(
                        ss              = ss,
                        section_prefix  = sec_pfx,
                        fig_prefix      = fig_pfx,
                        fig_start_num   = int(fig_start),
                        project_name    = proj_rpt,
                        include_jpcp    = inc_j,
                        include_crcp    = inc_c,
                        include_summary = inc_s,
                    )
                    st.download_button(
                        "📥 Download Rigid Report (ที่ปรึกษา)",
                        b_pro,
                        f"Rigid_Pro_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key="dl_rigid_pro"
                    )
                except Exception as e:
                    st.error(f"❌ ไม่สามารถสร้างรายงานได้: {e}")

    with col_dl2:
        st.markdown("**รวมทุกส่วน:**")
        if st.button("🗂️ สร้างรายงานรวม", type="primary", use_container_width=True):
            ss_d = get_ss_dict()
            b_full = build_report_full(ss_d)
            if b_full:
                st.download_button("📥 Download Full Report", b_full,
                                   f"ITM_Pave_Full_Report_{datetime.now().strftime('%Y%m%d')}.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                   use_container_width=True)
            else:
                st.warning("ไม่มีข้อมูลสำหรับสร้างรายงาน หรือ python-docx ไม่พร้อม")

    st.markdown("---")
    st.markdown("#### 💾 JSON Save / Load (Sidebar)")
    st.info("💡 ใช้ปุ่ม Save/Load JSON ใน Sidebar ด้านซ้ายเพื่อบันทึกและโหลดโปรเจกต์ทั้งหมด")

    st.divider()
    st.markdown("""
    <div style='text-align:center;color:#558B2F;font-size:0.85rem;padding:0.5rem;'>
        🛣️ <b>ITM Pave Pro v2.1</b> — AASHTO 1993 Pavement Design System<br>
        พัฒนาโดย รศ.ดร.อิทธิพล มีผล | ภาควิชาครุศาสตร์โยธา | คณะครุศาสตร์อุตสาหกรรม | มจพ.
    </div>
    """, unsafe_allow_html=True)
