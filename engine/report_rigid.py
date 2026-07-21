# ╔══════════════════════════════════════════════════════════════════╗
# ║  engine/report_rigid.py — ITM Pave Pro                          ║
# ║  Rigid Pavement Word Report — format ตาม Rigid V7               ║
# ║  พัฒนาโดย รศ.ดร.อิทธิพล มีผล | ภาควิชาครุศาสตร์โยธา มจพ.    ║
# ╚══════════════════════════════════════════════════════════════════╝

import io
import math
from datetime import datetime

TH  = 'TH SarabunPSK'
EQ  = 'Times New Roman'
HBG = 'BDD7EE'   # header blue
SBG = 'FFF2CC'   # sum yellow
PBG = 'CCFFCC'   # pass green
FBG = 'FFCCCC'   # fail red
SEL = 'FFFFAA'   # selected yellow

DEFAULT_INTRO = (
    'การออกแบบความหนาแผ่นคอนกรีตตามแนวทางของ AASHTO 1993 จำเป็นต้องอาศัยสมเหตุสมผลที่'
    'พัฒนามาจากผลการทดสอบ AASHO Road Test ซึ่งสะท้อนพฤติกรรมการรับน้ำหนักและการเสื่อมสภาพของแผ่น'
    'คอนกรีตภายใต้สภาพการใช้งานจริง สมการดังกล่าวรวมปัจจัยสำคัญหลายด้าน ทั้งด้านปริมาณจราจร '
    'ความน่าเชื่อถือของการออกแบบ คุณสมบัติวัสดุ และสภาพชั้นรองรับ เพื่อให้สามารถประเมินความหนา'
    'ที่เหมาะสมสำหรับรองรับปริมาณจราจรตลอดอายุโครงการได้อย่างแม่นยำ '
    'สมการหลักที่ใช้ในการออกแบบความหนาถนนคอนกรีตตาม AASHTO 1993 มีดังนี้'
)
DEFAULT_PAVEMENT_DESC = (
    'โดยมาตรฐานการออกแบบตามวิธี AASHTO 1993 ได้แบ่งโครงสร้างทางคอนกรีตออกเป็นหลายรูปแบบตาม'
    'ลักษณะการควบคุมความแตกร้าวและการถ่ายแรงระหว่างแผ่นคอนกรีต ได้แก่ Jointed Plain Concrete Pavement (JPCP), '
    'Jointed Reinforced Concrete Pavement (JRCP) และ Continuously Reinforced Concrete Pavement (CRCP)'
)
DEFAULT_SUMMARY = (
    'จากการคำนวณตามวิธีของ AASHTO 1993 ผิวทางคอนกรีต (Concrete Pavement) สามารถสรุปรูปแบบของ'
    'โครงสร้างชั้นทางที่ออกแบบได้ดังแสดงในตารางและรูปด้านล่าง'
)

LAYER_NAMES_EN = {
    'หินคลุกปรับปรุงคุณภาพด้วยปูนซีเมนต์ (CTB)': 'Cement Treated Base (CTB)',
    'หินคลุก CBR 80%':                            'Crushed Rock CBR 80%',
    'รองพื้นทางวัสดุมวลรวม CBR 25%':             'Granular Subbase CBR 25%',
    'วัสดุคัดเลือก ก':                            'Select Material A',
    'รองผิวทางคอนกรีตด้วย AC':                   'AC Leveling Course',
}


def _fmt_name(name):
    for k, v in LAYER_NAMES_EN.items():
        if k in name:
            return v
    return name


def build_rigid_report(ss: dict) -> bytes | None:
    """
    สร้าง Word Rigid Pavement Design Report

    ss ต้องมี:
      rigid_results    : {'JPCP': {...}, 'CRCP': {...}}
      jpcp_layers, crcp_layers : list of {'name', 'thickness_cm', 'E_MPa'}
      jpcp_design_rows, crcp_design_rows : list of design result rows
      jpcp_fig33_bytes, jpcp_fig34_bytes : bytes (optional)
      crcp_fig33_bytes, crcp_fig34_bytes : bytes (optional)
      jpcp_design_params, crcp_design_params : dict
      cbr_design : float
      report_settings: dict (optional)
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return None

    rr = ss.get('rigid_results', {})
    has_j = bool(rr.get('JPCP/JRCP') or rr.get('JPCP') or rr.get('jpcp'))
    has_c = bool(rr.get('CRCP') or rr.get('crcp'))
    if not has_j and not has_c:
        return None

    rs = ss.get('report_settings', {})
    sec_prefix = rs.get('section_number',   '4.5')
    fig_prefix = rs.get('figure_prefix',    '4-')
    fig_start  = int(rs.get('figure_start', 5))
    intro_text = rs.get('intro_text',       DEFAULT_INTRO)
    summary_text = rs.get('summary_text',   DEFAULT_SUMMARY)
    inc_summary  = rs.get('inc_summary',    True)
    proj_name    = ss.get('project_name',   '')
    calc_date    = datetime.now().strftime('%d/%m/%Y')
    cbr_design   = float(ss.get('cbr_design', 4.0))

    fig_counter = [fig_start]
    def nfig():
        n = fig_counter[0]; fig_counter[0] += 1; return n

    # ── Document setup ──────────────────────────────────────────
    doc   = Document()
    style = doc.styles['Normal']
    style.font.name = TH; style.font.size = Pt(15)
    sec = doc.sections[0]
    sec.page_width  = Cm(21); sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin  = Cm(2.5); sec.bottom_margin = Cm(2.0)

    # ── Helper functions ─────────────────────────────────────────
    def _shd(cell, color):
        tcPr = cell._tc.get_or_add_tcPr()
        s = OxmlElement('w:shd')
        s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto'); s.set(qn('w:fill'), color)
        tcPr.append(s)

    def _set_cw(cell, w):
        tcPr = cell._tc.get_or_add_tcPr()
        tcW  = OxmlElement('w:tcW'); tcW.set(qn('w:w'), str(w)); tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

    def _set_vmerge(cell, restart=False):
        tcPr = cell._tc.get_or_add_tcPr()
        vM   = OxmlElement('w:vMerge')
        if restart: vM.set(qn('w:val'), 'restart')
        tcPr.append(vM)

    def _set_valign(cell, val='center'):
        tcPr = cell._tc.get_or_add_tcPr()
        vA   = OxmlElement('w:vAlign'); vA.set(qn('w:val'), val); tcPr.append(vA)

    def _cell_margin(cell, mar=80):
        tcPr = cell._tc.get_or_add_tcPr()
        tcM  = OxmlElement('w:tcMar')
        for side in ['top','bottom','left','right']:
            m = OxmlElement(f'w:{side}'); m.set(qn('w:w'), str(mar)); m.set(qn('w:type'), 'dxa')
            tcM.append(m)
        tcPr.append(tcM)

    def _sc(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, bg=None, sz=15):
        cell.text = ''
        p = cell.paragraphs[0]; p.alignment = align
        r = p.add_run(text); r.font.name = TH; r.font.size = Pt(sz); r.bold = bold
        if bg: _shd(cell, bg)

    def _set_col_w(row, widths):
        for i, cell in enumerate(row.cells):
            _set_cw(cell, widths[i])

    def _add_heading(text, level=1):
        p    = doc.add_paragraph()
        run  = p.add_run(text)
        run.font.name = TH; run.font.size = Pt(15)
        run.bold = True; run.underline = (level <= 2)
        return p

    def _add_para(text, bold=False, indent_cm=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        run = p.add_run(text); run.font.name = TH; run.font.size = Pt(15); run.bold = bold
        if indent_cm > 0: p.paragraph_format.left_indent = Cm(indent_cm)
        return p

    def _add_fig_caption(text):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.font.name = TH; r.font.size = Pt(15); r.bold = True; r.underline = True

    def _eq_run(p, text, sub=False, sup=False, bold=False):
        run = p.add_run(text)
        run.font.name = EQ; run.font.size = Pt(11); run.bold = bold
        if sub or sup:
            rPr = run._r.get_or_add_rPr()
            va  = OxmlElement('w:vertAlign')
            va.set(qn('w:val'), 'subscript' if sub else 'superscript')
            rPr.append(va)
        return run

    def _th_run(p, text, bold=False):
        run = p.add_run(text); run.font.name = TH; run.font.size = Pt(15); run.bold = bold
        return run

    def _eq_line(indent_cm=1.5):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent  = Cm(indent_cm)
        p.paragraph_format.space_after  = Pt(2)
        return p

    def _sec_num(base, sub=None):
        return base if sub is None else f'{base}.{sub}'

    # ── หน้าปก ──────────────────────────────────────────────────
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('รายการคำนวณออกแบบ\nผิวทางคอนกรีต')
    r.font.name = TH; r.font.size = Pt(20); r.bold = True
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('ตามวิธี AASHTO 1993')
    r.font.name = TH; r.font.size = Pt(16)
    if proj_name:
        doc.add_paragraph()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'โครงการ: {proj_name}'); r.font.name = TH; r.font.size = Pt(15)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'วันที่: {calc_date}'); r.font.name = TH; r.font.size = Pt(15)
    doc.add_page_break()

    # ── หัวข้อหลัก + บทเกริ่นนำ ────────────────────────────────
    _add_heading(f'{sec_prefix}  การออกแบบผิวทางคอนกรีต', level=1)
    _add_para(intro_text)
    doc.add_paragraph()

    # สมการ AASHTO 1993 Rigid
    p1 = _eq_line()
    _eq_run(p1,'log'); _eq_run(p1,'10',sub=True); _eq_run(p1,'(W'); _eq_run(p1,'18',sub=True)
    _eq_run(p1,') = Z'); _eq_run(p1,'R',sub=True); _eq_run(p1,' × S'); _eq_run(p1,'o',sub=True)
    _eq_run(p1,' + 7.35 × log'); _eq_run(p1,'10',sub=True); _eq_run(p1,'(D+1) − 0.06')

    p2 = _eq_line()
    _eq_run(p2,'        + log'); _eq_run(p2,'10',sub=True)
    _eq_run(p2,'(ΔPSI / (4.5 − 1.5)) / (1 + 1.624×10'); _eq_run(p2,'7',sup=True)
    _eq_run(p2,' / (D+1)'); _eq_run(p2,'8.46',sup=True); _eq_run(p2,')')

    p3 = _eq_line()
    _eq_run(p3,'        + (4.22 − 0.32×P'); _eq_run(p3,'t',sub=True)
    _eq_run(p3,') × log'); _eq_run(p3,'10',sub=True)
    _eq_run(p3,' [(S'); _eq_run(p3,'c',sub=True); _eq_run(p3,'×C'); _eq_run(p3,'d',sub=True)
    _eq_run(p3,'×(D'); _eq_run(p3,'0.75',sup=True); _eq_run(p3,'−1.132))')
    _eq_run(p3,' / (215.63×J×(D'); _eq_run(p3,'0.75',sup=True)
    _eq_run(p3,' − 18.42 / (E'); _eq_run(p3,'c',sub=True)
    _eq_run(p3,'/k)'); _eq_run(p3,'0.25',sup=True); _eq_run(p3,')]')

    doc.add_paragraph()
    _th_run(doc.add_paragraph(), 'โดยที่:')

    cw_sym = [1396, 6281, 1395]
    tsym   = doc.add_table(rows=1, cols=3)
    tsym.style = 'Table Grid'; tsym.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_col_w(tsym.rows[0], cw_sym)
    for i, h in enumerate(['สัญลักษณ์','ความหมาย','หน่วย']):
        _sc(tsym.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HBG)
    for sym, meaning, unit in [
        ('W₁₈','จำนวนแกนเดี่ยว 18 kip ที่รองรับได้','ESALs'),
        ('ZR','Standard Normal Deviate ที่ความเชื่อมั่น R','-'),
        ('So','Overall Standard Deviation','-'),
        ('D','ความหนาแผ่นคอนกรีต','นิ้ว'),
        ('ΔPSI','การสูญเสีย Serviceability (4.5 − Pt)','-'),
        ('Pt','Terminal Serviceability Index','-'),
        ('Sc','Modulus of Rupture ของคอนกรีต','psi'),
        ('Cd','Drainage Coefficient','-'),
        ('J','Load Transfer Coefficient','-'),
        ('Ec','Modulus of Elasticity ของคอนกรีต','psi'),
        ('k','Modulus of Subgrade Reaction','pci'),
    ]:
        row = tsym.add_row(); _set_col_w(row, cw_sym)
        _sc(row.cells[0], sym, align=WD_ALIGN_PARAGRAPH.CENTER)
        _sc(row.cells[1], meaning)
        _sc(row.cells[2], unit, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    _add_para(DEFAULT_PAVEMENT_DESC)
    doc.add_paragraph()

    # ── helper: layer table ──────────────────────────────────────
    def _layer_table(layers, d_cm, ptype, cbr):
        cw = [756, 4536, 1728, 2052]
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = tbl.rows[0]; _set_col_w(hdr, cw)
        for i, h in enumerate(['ลำดับ','ชนิดวัสดุ','ความหนา (ซม.)','Modulus E (MPa)']):
            _sc(hdr.cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HBG)
        # แผ่นคอนกรีต
        row = tbl.add_row(); _set_col_w(row, cw)
        _sc(row.cells[0], '1', align=WD_ALIGN_PARAGRAPH.CENTER)
        _sc(row.cells[1], f'ผิวทางคอนกรีต {ptype}')
        _sc(row.cells[2], str(int(d_cm)), align=WD_ALIGN_PARAGRAPH.CENTER)
        _sc(row.cells[3], '-', align=WD_ALIGN_PARAGRAPH.CENTER)
        rn = 2
        for layer in layers:
            t = layer.get('thickness_cm', 0)
            if t <= 0: continue
            e = layer.get('E_MPa', 0)
            row = tbl.add_row(); _set_col_w(row, cw)
            _sc(row.cells[0], str(rn), align=WD_ALIGN_PARAGRAPH.CENTER)
            _sc(row.cells[1], layer.get('name', ''))
            _sc(row.cells[2], str(int(t)), align=WD_ALIGN_PARAGRAPH.CENTER)
            _sc(row.cells[3], f'{int(e):,}' if e > 0 else '-', align=WD_ALIGN_PARAGRAPH.CENTER)
            rn += 1
        mr_psi = int(1500 * cbr)
        mr_mpa = round(mr_psi / 145.038)
        row = tbl.add_row(); _set_col_w(row, cw)
        _sc(row.cells[0], str(rn), align=WD_ALIGN_PARAGRAPH.CENTER)
        _sc(row.cells[1], 'ดินคันทาง')
        _sc(row.cells[2], f'CBR ≥ {cbr:.1f} %', align=WD_ALIGN_PARAGRAPH.CENTER)
        _sc(row.cells[3], f'{mr_mpa:,} ({mr_psi:,} psi)', align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()

    # ── helper: ESB section ──────────────────────────────────────
    def _esb_section(layers):
        valid = [l for l in layers if l.get('thickness_cm', 0) > 0 and l.get('E_MPa', 0) > 0]
        if not valid: return
        _add_heading('การคำนวณ Subbase Elastic Modulus (ESB)', level=2)
        p_d = doc.add_paragraph(); p_d.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        _th_run(p_d, 'ค่า Subbase Elastic Modulus (E')
        r = p_d.add_run('SB'); r.font.name = EQ; r.font.size = Pt(11)
        rPr = r._r.get_or_add_rPr()
        va  = OxmlElement('w:vertAlign'); va.set(qn('w:val'), 'subscript'); rPr.append(va)
        _th_run(p_d, ') คำนวณจากโมดูลัสเทียบเท่าของชั้นวัสดุรองพื้นทาง โดยใช้สมการดังนี้')
        peq = _eq_line(2.0)
        _eq_run(peq,'E'); _eq_run(peq,'SB',sub=True); _eq_run(peq,'  =  [ Σ ( h')
        _eq_run(peq,'i',sub=True); _eq_run(peq,' × E'); _eq_run(peq,'i',sub=True)
        _eq_run(peq,'1/3',sup=True); _eq_run(peq,' ) / Σ h'); _eq_run(peq,'i',sub=True)
        _eq_run(peq,'  ]'); _eq_run(peq,'3',sup=True)
        doc.add_paragraph()
        _th_run(doc.add_paragraph(), 'การคำนวณแสดงในตารางดังนี้')

        cw2 = [570, 2900, 1400, 1400, 1400, 1400]
        tbl2 = doc.add_table(rows=1, cols=6)
        tbl2.style = 'Table Grid'; tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr2 = tbl2.rows[0]; _set_col_w(hdr2, cw2)
        for i, h in enumerate(['ลำดับ','ชั้นวัสดุ','h (ซม.)','E (MPa)','E^(1/3)','h×E^(1/3)']):
            _sc(hdr2.cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HBG)
        sh = 0.0; shE = 0.0
        for idx, layer in enumerate(valid, 1):
            h = layer['thickness_cm']; E = layer['E_MPa']
            E13 = E**(1/3); hE = h * E13; sh += h; shE += hE
            r2 = tbl2.add_row(); _set_col_w(r2, cw2)
            _sc(r2.cells[0], str(idx), align=WD_ALIGN_PARAGRAPH.CENTER)
            _sc(r2.cells[1], layer.get('name',''))
            _sc(r2.cells[2], f'{h:,}', align=WD_ALIGN_PARAGRAPH.CENTER)
            _sc(r2.cells[3], f'{E:,}', align=WD_ALIGN_PARAGRAPH.CENTER)
            _sc(r2.cells[4], f'{E13:.4f}', align=WD_ALIGN_PARAGRAPH.CENTER)
            _sc(r2.cells[5], f'{hE:,.2f}', align=WD_ALIGN_PARAGRAPH.CENTER)
        rs2 = tbl2.add_row(); _set_col_w(rs2, cw2)
        _sc(rs2.cells[0], '', bg=SBG)
        _sc(rs2.cells[1], 'รวม (Σ)', bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, bg=SBG)
        _sc(rs2.cells[2], f'{sh:.0f}', bold=True, bg=SBG, align=WD_ALIGN_PARAGRAPH.CENTER)
        _sc(rs2.cells[3], '', bg=SBG); _sc(rs2.cells[4], '', bg=SBG)
        _sc(rs2.cells[5], f'{shE:,.2f}', bold=True, bg=SBG, align=WD_ALIGN_PARAGRAPH.CENTER)
        if sh > 0:
            esb_mpa = (shE/sh)**3; esb_psi = esb_mpa * 145.038
            doc.add_paragraph()
            pr1 = _eq_line(1.5); _th_run(pr1, 'แทนค่า  ')
            _eq_run(pr1,'E'); _eq_run(pr1,'SB',sub=True)
            _eq_run(pr1, f'  =  [ {shE:,.2f} / {sh:.0f} ]'); _eq_run(pr1, '3', sup=True)
            pr2 = _eq_line(1.5); _th_run(pr2, 'ดังนั้น  ')
            _eq_run(pr2,'E'); _eq_run(pr2,'SB',sub=True)
            _eq_run(pr2, f'  =  {esb_mpa:,.2f}'); _th_run(pr2, '  MPa')
            _eq_run(pr2, f'  =  {esb_psi:,.0f}'); _th_run(pr2, '  psi', bold=True)
        doc.add_paragraph()

    # ── helper: k-value section ──────────────────────────────────
    def _kvalue_section(params, fig33, fig34, fig_n_k):
        cw_k = [5772, 1924, 1376]
        _add_para('ขั้นตอนที่ 1: หาค่า Composite Modulus of Subgrade Reaction (k∞)', bold=True)
        tk1 = doc.add_table(rows=1, cols=3)
        tk1.style = 'Table Grid'; tk1.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = tk1.rows[0]; _set_col_w(hdr, cw_k)
        for i, h in enumerate(['พารามิเตอร์','ค่า','หน่วย']):
            _sc(hdr.cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HBG)
        for p_n, val, unit in [
            ('Roadbed Soil Resilient Modulus (MR)',  f"{params.get('MR_psi',0):,.0f}",  'psi'),
            ('Subbase Elastic Modulus (ESB)',         f"{params.get('ESB_psi',0):,.0f}", 'psi'),
            ('Subbase Thickness (DSB)',               f"{params.get('DSB_in',0):.1f}",   'inches'),
            ('Composite Modulus k∞',                  f"{params.get('k_inf',0):,.0f}",   'pci'),
        ]:
            row = tk1.add_row(); _set_col_w(row, cw_k)
            _sc(row.cells[0], p_n)
            _sc(row.cells[1], val, align=WD_ALIGN_PARAGRAPH.CENTER)
            _sc(row.cells[2], unit, align=WD_ALIGN_PARAGRAPH.CENTER)
        if fig33:
            doc.add_paragraph()
            p_img = doc.add_paragraph(); p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(io.BytesIO(fig33), width=Inches(5.0))
            _add_fig_caption(f'รูปที่ {fig_prefix}{fig_n_k}  ค่า Composite Modulus of Subgrade Reaction, k∞ (pci)')
        doc.add_paragraph()
        ls = params.get('ls', 0.0)
        _add_para('ขั้นตอนที่ 2: ปรับแก้ค่า Loss of Support (LS)', bold=True)
        tk2 = doc.add_table(rows=1, cols=3)
        tk2.style = 'Table Grid'; tk2.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = tk2.rows[0]; _set_col_w(hdr, cw_k)
        for i, h in enumerate(['พารามิเตอร์','ค่า','หน่วย']):
            _sc(hdr.cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HBG)
        for p_n, val, unit in [
            ('Effective Modulus k∞ (จาก Step 1)',   f"{params.get('k_inf',0):,.0f}",  'pci'),
            ('Loss of Support Factor (LS)',           f"{ls:.1f}",                      '-'),
            ('Corrected Modulus k (ที่ใช้ออกแบบ)', f"{params.get('k_eff',0):,.0f}",  'pci'),
        ]:
            row = tk2.add_row(); _set_col_w(row, cw_k)
            _sc(row.cells[0], p_n)
            _sc(row.cells[1], val, align=WD_ALIGN_PARAGRAPH.CENTER)
            _sc(row.cells[2], unit, align=WD_ALIGN_PARAGRAPH.CENTER)
        if fig34 and ls > 0:
            doc.add_paragraph()
            p_img = doc.add_paragraph(); p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(io.BytesIO(fig34), width=Inches(5.0))
            _add_fig_caption(f'รูปที่ {fig_prefix}{fig_n_k+1}  การปรับแก้ค่า Modulus of Subgrade Reaction เนื่องจาก Loss of Support')
        doc.add_paragraph()

    # ── helper: design result ────────────────────────────────────
    def _design_result(params, rows, sel_d_cm):
        p_l = doc.add_paragraph()
        r   = p_l.add_run('ข้อมูลนำเข้าการออกแบบ:')
        r.font.name = TH; r.font.size = Pt(15); r.bold = True; r.underline = True

        cw_in = [3923, 1471, 2207, 1471]
        tin = doc.add_table(rows=1, cols=4)
        tin.style = 'Table Grid'; tin.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = tin.rows[0]; _set_col_w(hdr, cw_in)
        for i, h in enumerate(['พารามิเตอร์','สัญลักษณ์','ค่า','หน่วย']):
            _sc(hdr.cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HBG)
        dpsi = params.get('dpsi', 4.5 - params.get('pt', 2.5))
        zr   = params.get('ZR',  params.get('zr_rig', -1.282))
        ls   = params.get('ls',  params.get('ls_value', 0.0))
        for rd in [
            ('ESAL ออกแบบ','W₁₈',           f"{params.get('w18',0):,.0f}",         'ESALs'),
            ('Terminal Serviceability','Pt',  f"{params.get('pt',2.5):.1f}",          '-'),
            ('การสูญเสีย Serviceability','ΔPSI', f"{dpsi:.1f}",                      '-'),
            ('Reliability','R',              f"{params.get('R',90):.0f}",             '%'),
            ('Standard Normal Deviate','ZR', f"{zr:.3f}",                             '-'),
            ('Standard Deviation','So',      f"{params.get('so',0.35):.2f}",          '-'),
            ('Modulus of Subgrade Reaction','k_eff', f"{params.get('k_eff',0):,.0f}",'pci'),
            ('Loss of Support','LS',         f"{ls:.1f}",                             '-'),
            ("กำลังคอนกรีต","f'c",           f"{params.get('fc_cube',350):.0f} Cube",'ksc'),
            ('Modulus of Elasticity','Ec',   f"{params.get('ec',0):,.0f}",            'psi'),
            ('Modulus of Rupture','Sc',      f"{params.get('sc',600):.0f}",           'psi'),
            ('Load Transfer Coefficient','J',f"{params.get('j',2.8):.1f}",            '-'),
            ('Drainage Coefficient','Cd',    f"{params.get('cd',1.0):.2f}",           '-'),
        ]:
            row = tin.add_row(); _set_col_w(row, cw_in)
            _sc(row.cells[0], rd[0])
            _sc(row.cells[1], rd[1], align=WD_ALIGN_PARAGRAPH.CENTER)
            _sc(row.cells[2], rd[2], align=WD_ALIGN_PARAGRAPH.CENTER)
            _sc(row.cells[3], rd[3], align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()

        p_l2 = doc.add_paragraph()
        r = p_l2.add_run('ผลการตรวจสอบความหนาแผ่นคอนกรีต:')
        r.font.name = TH; r.font.size = Pt(15); r.bold = True; r.underline = True
        cw_res = [1188, 1188, 1620, 2052, 1512, 1512]
        tres = doc.add_table(rows=1, cols=6)
        tres.style = 'Table Grid'; tres.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr2 = tres.rows[0]; _set_col_w(hdr2, cw_res)
        for i, h in enumerate(['D (ซม.)','D (นิ้ว)','log₁₀(W₁₈)','W₁₈ รองรับได้','อัตราส่วน','ผล']):
            _sc(hdr2.cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HBG)
        for rv in rows:
            is_sel = (rv.get('d_cm', 0) == sel_d_cm)
            bg_row = SEL if is_sel else None
            bg_res = PBG if rv.get('passed') else FBG
            row2 = tres.add_row(); _set_col_w(row2, cw_res)
            _sc(row2.cells[0], f"{rv.get('d_cm',0):.0f}",      bold=is_sel, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg_row)
            _sc(row2.cells[1], f"{rv.get('d_inch',0):.0f}",                 align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg_row)
            _sc(row2.cells[2], f"{rv.get('log_w18',0):.4f}",               align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg_row)
            _sc(row2.cells[3], f"{rv.get('w18_cap',0):,.0f}",              align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg_row)
            _sc(row2.cells[4], f"{rv.get('ratio',0):.2f}",                 align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg_row)
            _sc(row2.cells[5], '✓ ผ่าน' if rv.get('passed') else '✗ ไม่ผ่าน',
                align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg_res)
        doc.add_paragraph()

        # สรุป
        p_l3 = doc.add_paragraph()
        r = p_l3.add_run('สรุปผลการออกแบบ:')
        r.font.name = TH; r.font.size = Pt(15); r.bold = True; r.underline = True
        sel_row = next((rv for rv in rows if rv.get('d_cm') == sel_d_cm), None)
        w18_cap = sel_row['w18_cap'] if sel_row else 0
        w18_req = sel_row.get('w18_req', params.get('w18', 0)) if sel_row else 0
        passed  = sel_row['passed']  if sel_row else False
        ratio   = sel_row['ratio']   if sel_row else 0
        for item in [
            f"ความหนาที่เลือก : {int(sel_d_cm)} ซม. ({round(sel_d_cm/2.54):.0f} นิ้ว)",
            f"ESAL ที่ต้องการ  : {w18_req:,.0f} ESALs",
            f"ESAL ที่รองรับได้ : {w18_cap:,.0f} ESALs",
            f"อัตราส่วน        : {ratio:.2f}",
            f"ผลการตรวจสอบ  : {'✅ ผ่านเกณฑ์' if passed else '❌ ไม่ผ่านเกณฑ์'}",
        ]:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(36)
            r = p.add_run(item); r.font.name = TH; r.font.size = Pt(15)
        doc.add_paragraph()

    # ── helper: summary table ────────────────────────────────────
    def _summary_table(layers, d_cm, ptype, cbr, fig_caption_text):
        valid = [l for l in layers if l.get('thickness_cm', 0) > 0]
        data_rows = [{'thick': str(int(d_cm)), 'material': f'ผิวทางคอนกรีต\n{ptype}'}]
        for l in valid:
            data_rows.append({'thick': str(int(l.get('thickness_cm',0))),
                              'material': l.get('name','')})
        data_rows.append({'thick': 'Existing',
                          'material': f'Earth Embankment\nor Subgrade, CBR≥\n{cbr:.0f} %'})

        # สร้างรูปตัดขวาง
        fig_bytes = None
        try:
            from engine.rigid_nomograph import plot_structure
            import matplotlib.pyplot as plt
            fig = plot_structure(valid, d_cm, title='')
            if fig:
                buf = io.BytesIO()
                fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
                buf.seek(0); fig_bytes = buf.read()
                plt.close(fig)
        except Exception:
            pass

        col_w = [3800, 1400, 3872]
        n_data = len(data_rows)
        tbl = doc.add_table(rows=1+n_data, cols=3)
        tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = tbl.rows[0]
        for i, cell in enumerate(hdr.cells): _set_cw(cell, col_w[i])
        _sc(hdr.cells[0], 'รายละเอียด', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HBG)
        _sc(hdr.cells[1], 'หนา\n(ซม.)',  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HBG)
        _sc(hdr.cells[2], 'ชนิดวัสดุ', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HBG)

        for i, dr in enumerate(data_rows):
            row = tbl.rows[1+i]
            for j, cell in enumerate(row.cells): _set_cw(cell, col_w[j])
            lc = row.cells[0]
            if i == 0:
                _set_vmerge(lc, restart=True); _cell_margin(lc)
                lc.text = ''; p_img = lc.paragraphs[0]
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if fig_bytes:
                    p_img.add_run().add_picture(io.BytesIO(fig_bytes), width=Inches(2.4))
            else:
                _set_vmerge(lc, restart=False); lc.text = ''
            _set_valign(lc, 'center')
            _sc(row.cells[1], dr['thick'], align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_valign(row.cells[1], 'center')
            _sc(row.cells[2], dr['material'])
            _set_valign(row.cells[2], 'center')

        if fig_caption_text:
            _add_fig_caption(fig_caption_text)
        doc.add_paragraph()

    # ── ดึง params จาก ss ────────────────────────────────────────
    def _get_params(prefix):
        p = dict(ss.get(f'{prefix}_design_params', {}))
        if not p:
            rr_data = rr.get(prefix.upper()) or rr.get(prefix) or {}
            p = {
                'w18':     rr_data.get('w18_req', 0),
                'k_eff':   rr_data.get('k_eff', 0),
                'k_inf':   ss.get(f'{prefix}_k_inf', 0) or 0,
                'fc_cube': rr_data.get('fc', 350),
                'ec':      rr_data.get('ec', 0) or ss.get('fc_cube', 350) * 14.223 * 0.8,
                'sc':      rr_data.get('sc', 600),
                'j':       rr_data.get('j', 2.8),
                'cd':      rr_data.get('cd', 1.0),
                'pt':      rr_data.get('pt', ss.get('pt_global', 2.5)),
                'so':      ss.get('so_rig', 0.35),
                'R':       ss.get('r0_rig', 90),
                'dpsi':    4.5 - float(rr_data.get('pt', ss.get('pt_global', 2.5))),
                'ls':      ss.get(f'{prefix}_ls_val', 0.0) or ss.get('ls_value', 0.0) or 0.0,
                'MR_psi':  ss.get('mr_subgrade_psi', 7000) or 7000,
                'ESB_psi': ss.get(f'{prefix}_esb', 0) or 0,
                'DSB_in':  ss.get(f'{prefix}_dsb', 0) or 0,
            }
        p['sel_d']  = ss.get(f'{prefix}_rec_d_cm') or 30
        p['cbr']    = ss.get('cbr_design', cbr_design)
        p['k_inf']  = ss.get(f'{prefix}_k_inf', p.get('k_inf', 0)) or 0
        p['DSB_in'] = ss.get(f'{prefix}_dsb', p.get('DSB_in', 0)) or 0
        p['ESB_psi']= ss.get(f'{prefix}_esb', p.get('ESB_psi', 0)) or 0
        p['MR_psi'] = ss.get('mr_subgrade_psi', p.get('MR_psi', 7000)) or 7000
        p['ls']     = ss.get(f'{prefix}_ls_val', p.get('ls', 0.0)) or 0.0
        # R และ ZR ต้องดึงจาก R0 (r0_rig) ที่ผู้ใช้เลือกจริงในหน้า Rigid Design
        # แก้บั๊ก: เดิม 'zr_rig' ไม่เคยถูก set ที่ไหนเลย ทำให้ ZR ค้างที่ -1.282
        # (ค่า R0=90%) เสมอ ไม่ว่าผู้ใช้จะเลือก R0 เท่าไหร่ก็ตาม
        from engine.rigid_nomograph import get_zr
        p['R']  = ss.get('r0_rig', 90)
        p['ZR'] = get_zr(p['R'])
        return p

    # ── JPCP section ─────────────────────────────────────────────
    if has_j:
        jpcp_layers = ss.get('jpcp_layers', [])
        jpcp_params = _get_params('jpcp')
        jpcp_rows   = ss.get('jpcp_design_rows', [])
        fig33_j     = ss.get('jpcp_fig33_bytes')
        fig34_j     = ss.get('jpcp_fig34_bytes')

        _add_heading(f'{_sec_num(sec_prefix,1)}  ชั้นโครงสร้างทางคอนกรีตประเภท JPCP/JRCP', level=2)
        fig_n = nfig()
        _layer_table(jpcp_layers, jpcp_params.get('sel_d', 30), 'JPCP/JRCP', jpcp_params.get('cbr', cbr_design))
        _esb_section(jpcp_layers)
        _add_heading(f'{_sec_num(sec_prefix,2)}  การคำนวณ k-value สำหรับ JPCP/JRCP', level=2)
        fig_n_k = nfig()
        _kvalue_section(jpcp_params, fig33_j, fig34_j, fig_n_k)
        if fig34_j and jpcp_params.get('ls', 0) > 0:
            fig_counter[0] += 1
        _add_heading('ผลการออกแบบความหนาผิวทางคอนกรีต JPCP/JRCP', level=3)
        _design_result(jpcp_params, jpcp_rows, jpcp_params.get('sel_d', 30))

    # ── CRCP section ──────────────────────────────────────────────
    if has_c:
        crcp_layers = ss.get('crcp_layers', [])
        crcp_params = _get_params('crcp')
        crcp_rows   = ss.get('crcp_design_rows', [])
        fig33_c     = ss.get('crcp_fig33_bytes')
        fig34_c     = ss.get('crcp_fig34_bytes')
        sub_off     = 2 if has_j else 0

        _add_heading(f'{_sec_num(sec_prefix, sub_off+1)}  ชั้นโครงสร้างทางคอนกรีตประเภท CRCP', level=2)
        _layer_table(crcp_layers, crcp_params.get('sel_d', 30), 'CRCP', crcp_params.get('cbr', cbr_design))
        _esb_section(crcp_layers)
        _add_heading(f'{_sec_num(sec_prefix, sub_off+2)}  การคำนวณ k-value สำหรับ CRCP', level=2)
        fig_n_k2 = nfig()
        _kvalue_section(crcp_params, fig33_c, fig34_c, fig_n_k2)
        if fig34_c and crcp_params.get('ls', 0) > 0:
            fig_counter[0] += 1
        _add_heading('ผลการออกแบบความหนาผิวทางคอนกรีต CRCP', level=3)
        _design_result(crcp_params, crcp_rows, crcp_params.get('sel_d', 30))

    # ── Summary section ───────────────────────────────────────────
    if inc_summary and (has_j or has_c):
        doc.add_page_break()
        parts = sec_prefix.split('.')
        try: parts[-1] = str(int(parts[-1])+1); h_sum = '.'.join(parts)
        except: h_sum = sec_prefix + '_สรุป'
        _add_heading(f'{h_sum}  สรุปโครงสร้างชั้นทางที่ออกแบบด้วยวิธี AASHTO 1993', level=1)
        _add_para(summary_text)
        doc.add_paragraph()
        pat = 1
        if has_j:
            fn = nfig()
            _add_para(f'รูปแบบที่ {pat}: ผิวทางคอนกรีต แบบ JPCP/JRCP  (รูปที่ {fig_prefix}{fn})', bold=True)
            _summary_table(ss.get('jpcp_layers', []), jpcp_params.get('sel_d', 30),
                           'JPCP/JRCP', jpcp_params.get('cbr', cbr_design),
                           f'รูปที่ {fig_prefix}{fn}  โครงสร้างชั้นทาง JPCP/JRCP')
            pat += 1
        if has_c:
            fn = nfig()
            _add_para(f'รูปแบบที่ {pat}: ผิวทางคอนกรีต แบบ CRCP  (รูปที่ {fig_prefix}{fn})', bold=True)
            _summary_table(ss.get('crcp_layers', []), crcp_params.get('sel_d', 30),
                           'CRCP', crcp_params.get('cbr', cbr_design),
                           f'รูปที่ {fig_prefix}{fn}  โครงสร้างชั้นทาง CRCP')

    # อ้างอิง
    doc.add_paragraph()
    _add_para('เอกสารอ้างอิง', bold=True)
    _add_para('AASHTO Guide for Design of Pavement Structures 1993. '
              'American Association of State Highway and Transportation Officials, Washington, D.C.')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ╔══════════════════════════════════════════════════════════════════╗
# ║  PDF Summary Report — รายงานย่อ (เปรียบเทียบ JPCP vs CRCP)      ║
# ║  format ตาม Rigid V7 (_create_pdf_summary)                       ║
# ╚══════════════════════════════════════════════════════════════════╝

def build_rigid_pdf_summary(ss: dict) -> bytes | None:
    """
    สร้าง PDF summary ย่อ — ชื่อโครงการ + ตารางเปรียบเทียบ JPCP vs CRCP + ตารางชั้นวัสดุ
    ss ต้องมี: rigid_results, jpcp/crcp_design_params, jpcp/crcp_design_rows,
               jpcp/crcp_layers, jpcp/crcp_rec_d_cm, jpcp/crcp_k_eff
    """
    try:
        from fpdf import FPDF
    except ImportError:
        return None

    import os
    from engine.rigid_nomograph import get_zr, mr_from_cbr

    BASE_DIR  = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))  # root ของโปรเจกต์
    FONT_REG  = os.path.join(BASE_DIR, 'Sarabun-Regular.ttf')
    FONT_BOLD = os.path.join(BASE_DIR, 'Sarabun-Bold.ttf')
    if not os.path.exists(FONT_REG):
        FONT_REG  = 'Sarabun-Regular.ttf'
        FONT_BOLD = 'Sarabun-Bold.ttf'

    rr = ss.get('rigid_results', {}) or {}
    has_j = bool(rr.get('JPCP/JRCP') or rr.get('JPCP') or rr.get('jpcp'))
    has_c = bool(rr.get('CRCP') or rr.get('crcp'))
    if not has_j and not has_c:
        return None

    proj_name = ss.get('project_name', '') or '(ไม่ระบุชื่อโครงการ)'
    date_str  = datetime.now().strftime('%d/%m/%Y %H:%M')

    pj = dict(ss.get('jpcp_design_params', {})) if has_j else {}
    pc = dict(ss.get('crcp_design_params', {})) if has_c else {}
    shared = pj if pj else pc   # พารามิเตอร์ร่วม (fc/ec/sc/cd/pt/so/dpsi) เหมือนกันทั้ง 2 ฝั่ง

    kj_eff = ss.get('jpcp_k_eff')
    kc_eff = ss.get('crcp_k_eff')
    kj_opt = pj.get('k_opt')
    kc_opt = pc.get('k_opt')

    dj_cm  = ss.get('jpcp_rec_d_cm')
    dc_cm  = ss.get('crcp_rec_d_cm')
    dj_in  = round(dj_cm / 2.54) if dj_cm else None
    dc_in  = round(dc_cm / 2.54) if dc_cm else None

    rows_j = ss.get('jpcp_design_rows', [])
    rows_c = ss.get('crcp_design_rows', [])

    def _get_row(rows, d_cm):
        return next((r for r in rows if r['d_cm'] == d_cm), None) if d_cm else None
    rj  = _get_row(rows_j, dj_cm)
    rc_ = _get_row(rows_c, dc_cm)

    fc_cube = shared.get('fc_cube', ss.get('fc_cube', 350))
    ec_psi  = shared.get('ec', 0)
    sc_val  = shared.get('sc', 600)
    cd      = shared.get('cd', ss.get('cd_rig', 1.0))
    pt      = shared.get('pt', float(ss.get('pt_global', 2.5)))
    dpsi    = shared.get('dpsi', 4.5 - pt)
    so      = shared.get('so', ss.get('so_rig', 0.35))
    R       = ss.get('r0_rig', 90)
    zr      = get_zr(R)
    cbr     = float(ss.get('cbr_design', 4.0))
    MR_psi  = float(ss.get('_shared_mr_inp') or ss.get('jpcp_mr_inp')
                     or ss.get('mr_subgrade_psi') or 0) or mr_from_cbr(cbr)

    jj = pj.get('j', ss.get('jpcp_j', 2.8))
    jc = pc.get('j', ss.get('crcp_j', 2.6))

    dkj = (kj_eff - kj_opt) if (kj_eff and kj_opt) else None
    dkc = (kc_eff - kc_opt) if (kc_eff and kc_opt) else None

    passed_j = rj['passed'] if rj else None
    passed_c = rc_['passed'] if rc_ else None
    kj_ok = (kj_eff >= kj_opt) if (kj_eff and kj_opt) else None
    kc_ok = (kc_eff >= kc_opt) if (kc_eff and kc_opt) else None
    overall_j = (passed_j and kj_ok) if (passed_j is not None and kj_ok is not None) else None
    overall_c = (passed_c and kc_ok) if (passed_c is not None and kc_ok is not None) else None

    sections = [
        ('1 · พารามิเตอร์ออกแบบ', [
            {'label': "f'c (cube)",    'val_j': f'{fc_cube:.0f} ksc',  'shared': True},
            {'label': 'Ec',            'val_j': f'{ec_psi:,.0f} psi',  'shared': True},
            {'label': 'Sc (ทล. lock)', 'val_j': f'{sc_val:.0f} psi',   'shared': True},
            {'label': 'J',             'val_j': f'{jj:.1f}' if jj else '-',
                                        'val_c': f'{jc:.1f}' if jc else '-'},
            {'label': 'Cd',            'val_j': f'{cd:.1f}',           'shared': True},
            {'label': 'Pt / DPSI',     'val_j': f'{pt:.1f} / {dpsi:.1f}', 'shared': True},
            {'label': 'ZR / So',       'val_j': f'{zr:.3f} / {so:.2f}',   'shared': True},
            {'label': 'CBR',           'val_j': f'{cbr:.1f} %',        'shared': True},
            {'label': 'MR (subgrade)', 'val_j': f'{MR_psi:,.0f} psi',  'shared': True},
        ]),
        ('2 · ความหนาแผ่นคอนกรีต', [
            {'label': 'D แนะนำ',
             'val_j': f'{dj_in} in ({dj_cm} cm)' if dj_cm else '-',
             'val_c': f'{dc_in} in ({dc_cm} cm)' if dc_cm else '-', 'bold': True},
            {'label': 'W18 required',
             'val_j': f'{rj["w18_req"]:,.0f}' if rj else '-',
             'val_c': f'{rc_["w18_req"]:,.0f}' if rc_ else '-'},
            {'label': 'W18 capacity',
             'val_j': f'{rj["w18_cap"]:,.0f}' if rj else '-',
             'val_c': f'{rc_["w18_cap"]:,.0f}' if rc_ else '-'},
            {'label': 'Ratio (cap/req)',
             'val_j': f'x{rj["ratio"]:.2f}' if rj else '-',
             'val_c': f'x{rc_["ratio"]:.2f}' if rc_ else '-', 'bold': True},
        ]),
        ('3 · k_opt vs k_eff', [
            {'label': 'k_eff',
             'val_j': f'{kj_eff:.0f} pci' if kj_eff else '-',
             'val_c': f'{kc_eff:.0f} pci' if kc_eff else '-'},
            {'label': 'k_opt (min required)',
             'val_j': f'{kj_opt:.0f} pci' if kj_opt else '-',
             'val_c': f'{kc_opt:.0f} pci' if kc_opt else '-'},
            {'label': 'Dk = k_eff - k_opt',
             'val_j': f'{dkj:+.0f} pci ({dkj/kj_opt*100:+.1f}%)' if dkj is not None else '-',
             'val_c': f'{dkc:+.0f} pci ({dkc/kc_opt*100:+.1f}%)' if dkc is not None else '-'},
        ]),
        ('4 · ผลการตรวจสอบ', [
            {'label': 'W18 cap >= W18 req',
             'val_j': 'ผ่าน' if passed_j else 'ไม่ผ่าน',
             'val_c': 'ผ่าน' if passed_c else 'ไม่ผ่าน', 'bold': True},
            {'label': 'k_eff >= k_opt',
             'val_j': 'ผ่าน' if kj_ok else 'ไม่ผ่าน',
             'val_c': 'ผ่าน' if kc_ok else 'ไม่ผ่าน', 'bold': True},
            {'label': 'สรุปผล',
             'val_j': 'ผ่าน' if overall_j else 'ไม่ผ่าน',
             'val_c': 'ผ่าน' if overall_c else 'ไม่ผ่าน',
             'bold': True, 'shade': True},
        ]),
    ]

    layers_j = ss.get('jpcp_layers', []) if has_j else []
    layers_c = ss.get('crcp_layers', []) if has_c else []

    designer = '—'

    class PDF(FPDF):
        def header(self):
            pass  # ไม่ใช้ auto-header — วาด header เองด้านล่าง

        def footer(self):
            self.set_y(-10)
            self.set_font('Sarabun', '', 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 8, f'Page {self.page_no()} | KMUTNB - ภาควิชาครุศาสตร์โยธา - มจพ.',
                      align='C')
            self.set_text_color(0, 0, 0)

    pdf = PDF(orientation='P', unit='mm', format='A4')
    pdf.add_font('Sarabun', '',  FONT_REG,  uni=True)
    pdf.add_font('Sarabun', 'B', FONT_BOLD, uni=True)
    pdf.set_font('Sarabun', '', 10)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=12)

    # ══════════════════════════════════════════════════════════
    # Header — สีน้ำเงิน
    # ══════════════════════════════════════════════════════════
    BLUE      = (21, 101, 192)   # #1565C0
    PAGE_W    = 190
    H_TOP     = 10
    COL_R_W   = 55
    COL_L_W   = PAGE_W - COL_R_W

    pdf.set_xy(10, H_TOP)
    pdf.set_font('Sarabun', 'B', 14)
    pdf.set_text_color(*BLUE)
    pdf.cell(COL_L_W, 7, 'Rigid Pavement Design Report', ln=False)

    pdf.set_xy(10 + COL_L_W, H_TOP)
    pdf.set_font('Sarabun', '', 8)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(COL_R_W, 7, f'ผู้ออกแบบ: {designer}', align='R', ln=True)

    pdf.set_xy(10, H_TOP + 7)
    pdf.set_font('Sarabun', '', 8)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(COL_L_W, 5, 'AASHTO 1993 · ITM Pave Pro', ln=False)

    pdf.set_xy(10 + COL_L_W, H_TOP + 7)
    pdf.set_font('Sarabun', '', 8)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(COL_R_W, 5, f'วันที่: {date_str}', align='R', ln=True)

    pdf.set_draw_color(*BLUE)
    pdf.set_line_width(0.6)
    pdf.line(10, H_TOP + 13, 200, H_TOP + 13)
    pdf.set_line_width(0.2)
    pdf.set_draw_color(0, 0, 0)

    pdf.set_xy(10, H_TOP + 16)
    pdf.set_font('Sarabun', 'B', 10)
    pdf.set_text_color(*BLUE)
    pdf.cell(0, 6, f'Project: {proj_name}', ln=True)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(2)

    def sec_header(title):
        pdf.set_fill_color(238, 242, 247)
        pdf.set_font('Sarabun', 'B', 8)
        pdf.set_text_color(84, 110, 122)
        pdf.cell(0, 5, f'  {title}', ln=True, fill=True)
        pdf.set_text_color(0, 0, 0)

    W_LABEL = 62
    W_COL   = 64

    def tbl_header():
        pdf.set_font('Sarabun', 'B', 9)
        pdf.set_fill_color(21, 101, 192)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(W_LABEL, 6, 'รายการ', border=0, fill=True)
        pdf.set_fill_color(21, 101, 192)
        pdf.cell(W_COL, 6, '  JPCP / JRCP', border=0, fill=True)
        pdf.set_fill_color(46, 125, 50)
        pdf.cell(W_COL, 6, '  CRCP', border=0, fill=True, ln=True)
        pdf.set_text_color(0, 0, 0)

    def tbl_row(label, val_j, val_c, shade=False, bold_val=False):
        pdf.set_fill_color(250, 250, 250) if shade else pdf.set_fill_color(255, 255, 255)
        pdf.set_font('Sarabun', '', 8)
        pdf.set_text_color(84, 110, 122)
        pdf.cell(W_LABEL, 5, f'  {label}', border='B', fill=True)
        pdf.set_text_color(26, 35, 126)
        f = 'B' if bold_val else ''
        pdf.set_font('Sarabun', f, 8)
        pdf.cell(W_COL, 5, f'  {val_j}', border='B', fill=True)
        pdf.set_text_color(27, 94, 32)
        pdf.cell(W_COL, 5, f'  {val_c}', border='B', fill=True, ln=True)
        pdf.set_text_color(0, 0, 0)

    def tbl_row_shared(label, val):
        pdf.set_fill_color(250, 250, 250)
        pdf.set_font('Sarabun', '', 8)
        pdf.set_text_color(84, 110, 122)
        pdf.cell(W_LABEL, 5, f'  {label}', border='B', fill=True)
        pdf.set_text_color(120, 120, 120)
        pdf.cell(W_COL, 5, f'  {val}', border='B', fill=True)
        pdf.cell(W_COL, 5, f'  {val}', border='B', fill=True, ln=True)
        pdf.set_text_color(0, 0, 0)

    tbl_header()
    for sec_title, rows in sections:
        sec_header(sec_title)
        for row in rows:
            if row.get('shared'):
                tbl_row_shared(row['label'], row['val_j'])
            else:
                tbl_row(row['label'], str(row['val_j']), str(row.get('val_c', '-')),
                        shade=row.get('shade', False),
                        bold_val=row.get('bold', False))

    pdf.ln(4)

    # ── Layer Structure Table ─────────────────────────────────
    if layers_j or layers_c:
        pdf.set_font('Sarabun', 'B', 9)
        pdf.set_fill_color(21, 101, 192)
        pdf.set_text_color(255, 255, 255)
        W_NO  = 10
        W_MAT = 115
        W_LC  = 32
        pdf.cell(W_NO,  6, '#',            border=0, fill=True)
        pdf.cell(W_MAT, 6, '  วัสดุ',      border=0, fill=True)
        pdf.set_fill_color(21, 101, 192)
        pdf.cell(W_LC,  6, '  JPCP (ซม.)', border=0, fill=True)
        pdf.set_fill_color(46, 125, 50)
        pdf.cell(W_LC,  6, '  CRCP (ซม.)', border=0, fill=True, ln=True)
        pdf.set_text_color(0, 0, 0)

        pdf.set_font('Sarabun', 'B', 9)
        pdf.set_fill_color(238, 242, 247)
        pdf.set_text_color(21, 101, 192)
        pdf.cell(W_NO,  6, '0',                  border='B', fill=True)
        pdf.cell(W_MAT, 6, '  แผ่นคอนกรีต (D)', border='B', fill=True)
        pdf.cell(W_LC,  6, f'  {dj_cm or "-"}',  border='B', fill=True)
        pdf.set_text_color(46, 125, 50)
        pdf.cell(W_LC,  6, f'  {dc_cm or "-"}',  border='B', fill=True, ln=True)
        pdf.set_text_color(0, 0, 0)

        names_j = [l['name'] for l in layers_j]
        names_c = [l['name'] for l in layers_c]
        all_names = list(dict.fromkeys(names_j + names_c))

        def _thick(layers, name):
            for l in layers:
                if l['name'] == name:
                    return l['thickness_cm']
            return None

        tot_j = dj_cm or 0
        tot_c = dc_cm or 0
        for i, name in enumerate(all_names, 1):
            tj = _thick(layers_j, name)
            tc = _thick(layers_c, name)
            if tj: tot_j += tj
            if tc: tot_c += tc
            pdf.set_font('Sarabun', '', 8)
            pdf.set_fill_color(255, 255, 255)
            pdf.set_text_color(100, 100, 100)
            pdf.cell(W_NO, 6, str(i), border='B', fill=True)
            pdf.set_text_color(84, 110, 122)
            pdf.cell(W_MAT, 6, f'  {name[:65]}', border='B', fill=True)
            pdf.set_text_color(26, 35, 126)
            pdf.cell(W_LC, 6, f'  {tj if tj else "-"}', border='B', fill=True)
            pdf.set_text_color(27, 94, 32)
            pdf.cell(W_LC, 6, f'  {tc if tc else "-"}', border='B', fill=True, ln=True)
            pdf.set_text_color(0, 0, 0)

        pdf.set_font('Sarabun', 'B', 9)
        pdf.set_fill_color(240, 244, 255)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(W_NO,  6, '',                           border='B', fill=True)
        pdf.cell(W_MAT, 6, '  รวมทั้งหมด (รวมคอนกรีต)', border='B', fill=True)
        pdf.set_text_color(21, 101, 192)
        pdf.cell(W_LC,  6, f'  {tot_j} ซม.',             border='B', fill=True)
        pdf.set_text_color(46, 125, 50)
        pdf.cell(W_LC,  6, f'  {tot_c} ซม.',             border='B', fill=True, ln=True)

    return bytes(pdf.output())
