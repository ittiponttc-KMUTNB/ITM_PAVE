# ╔══════════════════════════════════════════════════════════════════╗
# ║  engine/report_cbr.py — ITM Pave Pro                            ║
# ║  CBR Word Report — format ตาม CBR Calculator V4                 ║
# ║  พัฒนาโดย รศ.ดร.อิทธิพล มีผล | ภาควิชาครุศาสตร์โยธา มจพ.    ║
# ╚══════════════════════════════════════════════════════════════════╝

import io
import math
import numpy as np
from datetime import datetime

from docx import Document as DocxDoc
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FN       = 'TH SarabunPSK'
FS       = 15
TFS      = 14
HDR_COLOR = 'D9E2F3'
MPA_PER_CBR = 1500 * 0.006895   # 10.3425 MPa/%CBR


# ─────────────────────────────────────────────
#  Calculation
# ─────────────────────────────────────────────

def calc_max_rank_percentile(cbr_values_raw):
    """Max Rank Percentile Method — return (cbr_sorted, n, unique_cbr, unique_pct, full_table)"""
    cbr_sorted  = np.sort(cbr_values_raw)
    n           = len(cbr_sorted)
    unique_cbr  = np.unique(cbr_sorted)
    unique_pct  = np.array([np.sum(cbr_sorted >= v) / n * 100 for v in unique_cbr])
    full_table  = []
    seen        = set()
    for i, v in enumerate(cbr_sorted):
        count_gte = int(np.sum(cbr_sorted >= v))
        pct_gte   = count_gte / n * 100
        show      = v not in seen
        if show:
            seen.add(v)
        full_table.append({'order': i+1, 'cbr': v,
                           'count_gte': count_gte, 'pct_gte': pct_gte, 'show_pct': show})
    return cbr_sorted, n, unique_cbr, unique_pct, full_table


def interp_cbr(target_pct, unique_pct, unique_cbr):
    return float(np.interp(target_pct, unique_pct[::-1], unique_cbr[::-1]))


# ─────────────────────────────────────────────
#  Document helpers
# ─────────────────────────────────────────────

def _new_doc():
    doc   = DocxDoc()
    style = doc.styles['Normal']
    style.font.name = FN
    style.font.size = Pt(FS)
    try:
        style._element.rPr.rFonts.set(qn('w:eastAsia'), FN)
    except Exception:
        pass
    sec = doc.sections[0]
    sec.page_width  = Cm(21.0); sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.5);  sec.right_margin  = Cm(2.5)
    sec.top_margin  = Cm(2.5);  sec.bottom_margin = Cm(2.5)
    return doc


def _run(para, text, bold=False, sz=FS, italic=False, color=None):
    r = para.add_run(text)
    r.font.name = FN; r.font.size = Pt(sz)
    r.bold = bold; r.italic = italic
    if color:
        r.font.color.rgb = color
    try:
        r._element.rPr.rFonts.set(qn('w:eastAsia'), FN)
    except Exception:
        pass
    return r


def _set_cell(cell, bold=False, align=None, sz=TFS, shading=None, color=None):
    for para in cell.paragraphs:
        if align:
            para.alignment = align
        for r in para.runs:
            r.font.name = FN; r.font.size = Pt(sz); r.bold = bold
            if color:
                r.font.color.rgb = color
            try:
                r._element.rPr.rFonts.set(qn('w:eastAsia'), FN)
            except Exception:
                pass
    if shading:
        _shd(cell, shading)


def _shd(cell, color_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color_hex)
    tcPr.append(shd)


def _thai_justify(para):
    pPr = para._element.get_or_add_pPr()
    jc  = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'thaiDistribute')
    pPr.append(jc)


def _caption(doc, num_str, title_str, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f'ตารางที่ {num_str}  {title_str}')
    r.font.name = FN; r.font.size = Pt(FS); r.bold = bold
    try:
        r._element.rPr.rFonts.set(qn('w:eastAsia'), FN)
    except Exception:
        pass


def _fig_caption(doc, num_str, title_str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'รูปที่ {num_str}  {title_str}')
    r.font.name = FN; r.font.size = Pt(FS); r.bold = True
    try:
        r._element.rPr.rFonts.set(qn('w:eastAsia'), FN)
    except Exception:
        pass


# ─────────────────────────────────────────────
#  Public builder
# ─────────────────────────────────────────────

def build_cbr_report(ss: dict) -> bytes | None:
    """
    สร้าง Word CBR Report

    ss ต้องมี:
      cbr_values        : list[float]
      cbr_percentile    : float
      cbr_design        : float        (CBR ที่ใช้ออกแบบจริง)
      section_number    : str  (เช่น '4.3')
      table_number      : str  (เช่น '4-7')
      figure_number     : str  (เช่น '4-7')
      section_title     : str
      table_caption     : str
      figure_caption    : str
      improve_soil_check: bool (optional)
      odemark_result    : dict (optional)
    """
    cbr_values = ss.get('cbr_values')
    if not cbr_values or len(cbr_values) == 0:
        return None

    cbr_values        = [float(v) for v in cbr_values]
    target_percentile = float(ss.get('cbr_percentile', 90))
    design_cbr        = float(ss.get('cbr_design', 4.0))
    section_number    = ss.get('section_number', '4.3')
    table_number      = ss.get('table_number', '4-7')
    figure_number     = ss.get('figure_number', '4-7')
    section_title     = ss.get('section_title',
                               'ข้อมูลความแข็งแรงของดินฐานรากบริเวณพื้นที่โครงการ')
    table_caption     = ss.get('table_caption',
                               'ค่าเปอร์เซ็นต์ไทล์ และค่า CBR ของตัวอย่างดินฐานรากตามแนวสายทาง')
    figure_caption    = ss.get('figure_caption',
                               'กราฟแสดงความสัมพันธ์ระหว่าง Percentile และ CBR ของดินฐานรากตามแนวสายทาง')
    improve_check     = bool(ss.get('improve_soil_check', False))
    odemark_result    = ss.get('odemark_result')

    # ── คำนวณ ──
    cbr_sorted, n, unique_cbr, unique_pct, full_table = calc_max_rank_percentile(cbr_values)
    cbr_at_pct = interp_cbr(target_percentile, unique_pct, unique_cbr)

    if improve_check and odemark_result:
        cbr_for_report = odemark_result.get('cbr_eq_design', math.floor(odemark_result['cbr_eq']))
    else:
        cbr_for_report = int(math.floor(design_cbr))

    doc = _new_doc()

    # ════════════════════════════════════════
    # 1. หัวข้อ
    # ════════════════════════════════════════
    h = doc.add_paragraph()
    h.paragraph_format.space_after = Pt(6)
    _run(h, f'{section_number}\t{section_title}', bold=True)

    # ════════════════════════════════════════
    # 2. บทเกริ่นนำ
    # ════════════════════════════════════════
    intro = doc.add_paragraph()
    intro.paragraph_format.first_line_indent = Cm(1.25)
    intro.paragraph_format.space_after = Pt(6)
    _thai_justify(intro)

    parts = [
        ('ความแข็งแรงของดินฐานรากบริเวณโดยรอบพื้นที่โครงการ หรือกำลังรับน้ำหนักของดินพื้นทางเดิม '
         'หรือพื้นทางเดิมสามารถประเมินจากรายงานสำรวจภูมิประเทศของดิน ซึ่งสามารถทำการทดสอบได้หลากหลายวิธี เช่น '
         'Plate Bearing Test CBR Test หรือ Modulus of Subgrade Reaction สำหรับการออกแบบถนนลาดยางและคอนกรีตนั้นใช้ค่า CBR '
         'ซึ่งนิยมใช้กันแพร่หลาย เมื่อกำหนดกำลังรับน้ำหนักของดินพื้นทางเดิม '
         'โดยการเจาะสำรวจดินในสนามตามรายงานการสอบดินของห้องปฏิบัติการ เพื่อหาค่า CBR '
         'ของดินพื้นทางเดินเพื่อเป็นข้อมูลในการออกแบบ ซึ่งผลการทดสอบค่า CBR ของดินฐานรากตามแนวสายทาง จำนวน ',
         False),
        (f'{n}', True),
        (' ตัวอย่าง พบว่าที่เปอร์เซ็นต์ไทล์ ร้อยละ ', False),
        (f'{target_percentile:.0f}', True),
        (' ของค่ากำลังที่พบเท่ากับ CBR เท่ากับ ', False),
        (f'{cbr_at_pct:.1f}', True),
        (' % อย่างไรก็ตาม ที่ปรึกษาเลือกค่า CBR เท่ากับ ', False),
        (f'{cbr_for_report}', True),
        (' % มาใช้ในการออกแบบโครงสร้างชั้นทาง ดังแสดงผลการวิเคราะห์ใน', False),
        (f'ตารางที่ {table_number}', True),
        (' และ', False),
        (f'รูปที่ {figure_number}', True),
    ]
    for text, is_bold in parts:
        r = intro.add_run(text)
        r.font.name = FN; r.font.size = Pt(FS); r.bold = is_bold
        try:
            r._element.rPr.rFonts.set(qn('w:eastAsia'), FN)
        except Exception:
            pass

    doc.add_paragraph()

    # ════════════════════════════════════════
    # 3. ตาราง CBR — 6 คอลัมน์ split 2 ฝั่ง
    # ════════════════════════════════════════
    _caption(doc, table_number, table_caption)
    half_n   = (n + 1) // 2
    cbr_tbl  = doc.add_table(rows=half_n + 1, cols=6)
    cbr_tbl.style   = 'Table Grid'
    cbr_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['CBR (%)', 'จำนวนที่≥', 'Percentile (%)',
               'CBR (%)', 'จำนวนที่≥', 'Percentile (%)']
    for j, h in enumerate(headers):
        cell = cbr_tbl.rows[0].cells[j]
        cell.paragraphs[0].add_run(h)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell(cell, bold=True, sz=TFS, shading=HDR_COLOR)

    for i in range(half_n):
        row = cbr_tbl.rows[i + 1]
        for side, idx in [('left', i), ('right', i + half_n)]:
            offset = 0 if side == 'left' else 3
            if idx < n:
                ft = full_table[idx]
                vals = [
                    f'{ft["cbr"]:.2f}',
                    f'{ft["count_gte"]}' if ft['show_pct'] else '',
                    f'{ft["pct_gte"]:.1f}' if ft['show_pct'] else '',
                ]
            else:
                vals = ['', '', '']
            for j, val in enumerate(vals):
                cell = row.cells[offset + j]
                cell.paragraphs[0].add_run(val)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_cell(cell, sz=TFS)

    # กำหนดความกว้างคอลัมน์
    col_widths = [Cm(2.0), Cm(2.5), Cm(3.5), Cm(2.0), Cm(2.5), Cm(3.5)]
    for row in cbr_tbl.rows:
        for j, w in enumerate(col_widths):
            row.cells[j].width = w

    doc.add_paragraph()

    # ════════════════════════════════════════
    # 4. ตารางสถิติ
    # ════════════════════════════════════════
    h_stat = doc.add_paragraph()
    _run(h_stat, 'ผลการวิเคราะห์', bold=True)

    stat_tbl = doc.add_table(rows=7, cols=2)
    stat_tbl.style     = 'Table Grid'
    stat_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    stat_data = [
        ('รายการ',                                'ค่า'),
        ('จำนวนตัวอย่าง',                         f'{n}'),
        ('ค่าต่ำสุด',                             f'{np.min(cbr_values):.2f} %'),
        ('ค่าสูงสุด',                             f'{np.max(cbr_values):.2f} %'),
        ('ค่าเฉลี่ย',                             f'{np.mean(cbr_values):.2f} %'),
        ('ส่วนเบี่ยงเบนมาตรฐาน',                 f'{np.std(cbr_values):.2f} %'),
        (f'CBR ที่ Percentile {target_percentile:.0f}%', f'{cbr_at_pct:.2f} %'),
    ]

    for i, (c1, c2) in enumerate(stat_data):
        r   = stat_tbl.rows[i]
        is_hdr = (i == 0)
        is_last = (i == len(stat_data) - 1)
        r.cells[0].paragraphs[0].add_run(c1)
        r.cells[1].paragraphs[0].add_run(c2)
        _set_cell(r.cells[0], bold=is_hdr, sz=TFS,
                  shading=HDR_COLOR if is_hdr else None)
        _set_cell(r.cells[1], bold=(is_hdr or is_last), sz=TFS,
                  shading=HDR_COLOR if is_hdr else None,
                  color=RGBColor(255, 0, 0) if is_last else None)

    for row in stat_tbl.rows:
        row.cells[0].width = Cm(12)
        row.cells[1].width = Cm(4)

    doc.add_paragraph()

    # ════════════════════════════════════════
    # 5. กราฟ matplotlib
    # ════════════════════════════════════════
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt

        fig_mpl, ax = plt.subplots(figsize=(6, 6))
        ax.plot(unique_cbr, unique_pct, 'b-', linewidth=2,
                marker='x', markersize=6,
                markerfacecolor='black', markeredgecolor='black',
                label='CBR Distribution')
        ax.plot([0, cbr_at_pct], [target_percentile, target_percentile],
                'r--', linewidth=2, label=f'Percentile {target_percentile:.0f}%')
        ax.plot([cbr_at_pct, cbr_at_pct], [0, target_percentile],
                'r--', linewidth=2, label=f'CBR = {cbr_at_pct:.2f}%')
        ax.annotate(f'{cbr_at_pct:.2f}',
                    xy=(cbr_at_pct, 0), xytext=(cbr_at_pct, -8),
                    fontsize=12, color='red', fontweight='bold', ha='center')
        ax.set_xlim(0, max(unique_cbr) * 1.1)
        ax.set_ylim(0, 100)
        ax.set_xlabel('CBR (%)', fontsize=12)
        ax.set_ylabel('Percentile (%)', fontsize=12)
        ax.set_title(f'CBR at Percentile {target_percentile:.0f}%', fontsize=14)
        ax.legend(loc='upper right', fontsize=10)
        ax.grid(True, alpha=0.3)
        for spine in ax.spines.values():
            spine.set_linewidth(2); spine.set_color('black')
        plt.tight_layout()

        chart_buf = io.BytesIO()
        fig_mpl.savefig(chart_buf, format='png', dpi=150,
                        bbox_inches='tight', facecolor='white')
        chart_buf.seek(0)
        plt.close(fig_mpl)

        chart_para = doc.add_paragraph()
        chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chart_para.add_run().add_picture(chart_buf, width=Cm(12))
    except Exception:
        doc.add_paragraph('[ไม่สามารถสร้างกราฟได้ — matplotlib error]')

    _fig_caption(doc, figure_number, figure_caption)

    # ════════════════════════════════════════
    # 6. Footer
    # ════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    f1 = doc.add_paragraph()
    f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(f1, 'พัฒนาโดย รศ.ดร.อิทธิพล มีผล', sz=TFS, italic=True)
    f2 = doc.add_paragraph()
    f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(f2, 'ภาควิชาครุศาสตร์โยธา คณะครุศาสตร์อุตสาหกรรม มจพ.', sz=TFS, italic=True)

    # ════════════════════════════════════════
    # 7. Odemark (ถ้ามี)
    # ════════════════════════════════════════
    if improve_check and odemark_result:
        res = odemark_result
        doc.add_paragraph()
        doc.add_paragraph()

        imp_h = doc.add_paragraph()
        _run(imp_h, 'การปรับปรุงดินคันทาง (Subgrade Improvement)', bold=True)

        imp_intro = doc.add_paragraph()
        imp_intro.paragraph_format.first_line_indent = Cm(1.25)
        _thai_justify(imp_intro)
        r = imp_intro.add_run(
            'เนื่องจากค่า CBR ของดินเดิมที่ได้จากการวิเคราะห์ทางสถิติมีค่าต่ำหรือแหล่งวัสดุดินถมคันทางมีค่า CBR ต่ำ '
            'จึงได้ทำการปรับปรุงดินคันทางโดยการใช้วัสดุคุณภาพดีปูทับ '
            'และคำนวณค่า CBR เทียบเท่า (CBR Equivalent) ด้วยวิธี Odemark (1974) '
            'โดยพิจารณาโครงสร้างดิน 2 ชั้น ดังนี้')
        r.font.name = FN; r.font.size = Pt(FS)
        try:
            r._element.rPr.rFonts.set(qn('w:eastAsia'), FN)
        except Exception:
            pass

        doc.add_paragraph()

        # ตารางชั้นดิน
        imp_tbl = doc.add_table(rows=3, cols=4)
        imp_tbl.style     = 'Table Grid'
        imp_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        imp_headers = ['ชั้นดิน', 'ชนิดวัสดุ', 'ความหนา (ซม.)', 'MR (MPa)']
        imp_data    = [
            ('ชั้นที่ 1 (วัสดุปรับปรุง)',
             res['mat1'], f"{res['h1_cm']:.1f}", f"{res['mr1_mpa']:.1f}"),
            ('ชั้นที่ 2 (ดินถมคันทางใหม่)',
             f"ดินถมคันทาง CBR = {res['cbr2']:.1f} %",
             f"{res['h2_cm']:.1f}", f"{res['mr2_mpa']:.2f}"),
        ]
        for j, h in enumerate(imp_headers):
            cell = imp_tbl.rows[0].cells[j]
            cell.paragraphs[0].add_run(h)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell(cell, bold=True, sz=TFS, shading=HDR_COLOR)
        for ri, row_vals in enumerate(imp_data):
            for ci, val in enumerate(row_vals):
                cell = imp_tbl.rows[ri + 1].cells[ci]
                cell.paragraphs[0].add_run(val)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_cell(cell, sz=TFS)
        col_w = [Cm(4.0), Cm(7.0), Cm(2.5), Cm(2.5)]
        for row in imp_tbl.rows:
            for j, w in enumerate(col_w):
                row.cells[j].width = w

        doc.add_paragraph()

        # วิธีการคำนวณ
        calc_h = doc.add_paragraph()
        _run(calc_h, 'วิธีการคำนวณ', bold=True)

        calc_lines = [
            'สูตร Odemark (1974):  MR_eq = ( Σ(h_i × MR_i^(1/3)) / Σh_i )^3',
            f"ชั้นที่ 1 : h = {res['h1_cm']:.1f} cm, MR = {res['mr1_mpa']:.1f} MPa, "
            f"MR^(1/3) = {res['mr1_mpa']**(1/3):.4f}",
            f"ชั้นที่ 2 : h = {res['h2_cm']:.1f} cm, MR = {res['mr2_mpa']:.2f} MPa, "
            f"MR^(1/3) = {res['mr2_mpa']**(1/3):.4f}",
            f"Σh = {res['sum_h']:.1f} cm",
            f"Σ(h·MR^(1/3)) = {res['sum_hE13']:.4f}",
            f"MR_eq = ({res['sum_hE13']:.4f} / {res['sum_h']:.1f})^3 = {res['mr_eq_mpa']:.2f} MPa",
            f"CBR_equivalent = MR_eq / (1500 × 0.006895) = {res['cbr_eq']:.2f} %",
        ]
        for line in calc_lines:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = FN; r.font.size = Pt(TFS)
            try:
                r._element.rPr.rFonts.set(qn('w:eastAsia'), FN)
            except Exception:
                pass

        doc.add_paragraph()

        # สรุปผล
        res_h = doc.add_paragraph()
        _run(res_h, 'ผลการคำนวณ', bold=True)

        sum_tbl = doc.add_table(rows=4, cols=2)
        sum_tbl.style     = 'Table Grid'
        sum_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        sum_data = [
            ('รายการ',                        'ค่า'),
            ('MR equivalent',                  f"{res['mr_eq_mpa']:.2f} MPa"),
            ('CBR equivalent (จากการคำนวณ)',   f"{res['cbr_eq']:.2f} %"),
            ('CBR equivalent (ใช้ออกแบบ)',     f"{res.get('cbr_eq_design', math.floor(res['cbr_eq']))} %"),
        ]
        for ri, (c1, c2) in enumerate(sum_data):
            is_hdr  = (ri == 0)
            is_last = (ri == len(sum_data) - 1)
            sum_tbl.rows[ri].cells[0].paragraphs[0].add_run(c1)
            sum_tbl.rows[ri].cells[1].paragraphs[0].add_run(c2)
            _set_cell(sum_tbl.rows[ri].cells[0], bold=is_hdr, sz=TFS,
                      shading=HDR_COLOR if is_hdr else None)
            _set_cell(sum_tbl.rows[ri].cells[1], bold=(is_hdr or is_last), sz=TFS,
                      shading=HDR_COLOR if is_hdr else None,
                      color=RGBColor(255, 0, 0) if is_last else None)
        for row in sum_tbl.rows:
            row.cells[0].width = Cm(12)
            row.cells[1].width = Cm(4)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
