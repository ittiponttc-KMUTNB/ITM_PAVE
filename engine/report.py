# ╔══════════════════════════════════════════════════════════════════╗
# ║  engine/report.py — ITM Pave Pro                                ║
# ║  Word Report Builder (python-docx) — Consultant Grade           ║
# ║  ไม่มี st. ใดๆ ทั้งสิ้น — pure python-docx functions           ║
# ║  พัฒนาโดย รศ.ดร.อิทธิพล มีผล | ภาควิชาครุศาสตร์โยธา มจพ.    ║
# ╚══════════════════════════════════════════════════════════════════╝

import io
import math
import numpy as np
from datetime import datetime

from docx import Document as DocxDoc
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from constants import SLAB_THICKNESSES, SLAB_LABELS
from engine.design import calc_percentile_cbr, aashto_rigid_w18
# figures imported lazily inside functions (avoid matplotlib load at startup)

# ─────────────────────────────────────────────
#  สี
# ─────────────────────────────────────────────
_HDR_NAVY   = '345E8B'   # header ตาราง
_HDR_GREEN  = 'C8E6C9'   # header ตารางเดิม
_ROW_ALT    = 'EEF4FB'   # แถวคู่
_GREEN      = RGBColor(0x00, 0x70, 0x00)
_RED        = RGBColor(0xC0, 0x00, 0x00)
_BLUE       = RGBColor(0x00, 0x47, 0xAB)
_DARK       = RGBColor(0x1A, 0x23, 0x32)

# ─────────────────────────────────────────────
#  Document Helpers (private)
# ─────────────────────────────────────────────

def _new_doc():
    doc   = DocxDoc()
    style = doc.styles['Normal']
    style.font.name = 'TH SarabunPSK'
    style.font.size = Pt(15)
    try:
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')
    except Exception:
        pass
    return doc


def _run(para, text, bold=False, size=15, color=None, italic=False):
    r             = para.add_run(text)
    r.font.name   = 'TH SarabunPSK'
    r.font.size   = Pt(size)
    r.font.bold   = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    try:
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')
    except Exception:
        pass
    return r


def _eq_run(para, text, size=11, bold=False, italic=True):
    """Times New Roman สำหรับสมการ"""
    r           = para.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold   = bold
    r.font.italic = italic
    try:
        r._element.rPr.rFonts.set(qn('w:cs'), 'Times New Roman')
    except Exception:
        pass
    return r


def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _cell(cell, text, align=WD_ALIGN_PARAGRAPH.CENTER,
          size=14, bold=False, fill=None, color=None):
    cell.text = ''
    p         = cell.paragraphs[0]
    p.alignment = align
    r = _run(p, text, bold=bold, size=size, color=color)
    if fill:
        _set_cell_bg(cell, fill)
    return r


def _doc_to_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _heading(doc, text: str, level: int = 1, size: int = 16):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    _run(p, text, bold=True, size=size)
    return p


def _para(doc, indent_cm=0, space_before=4, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(indent_cm)
    p.paragraph_format.space_before  = Pt(space_before)
    p.paragraph_format.space_after   = Pt(space_after)
    return p


def _make_table(doc, headers, rows,
                hdr_bg=_HDR_NAVY, alt_bg=_ROW_ALT,
                hdr_color=None, col_aligns=None):
    """
    สร้างตารางพร้อม header สี navy, แถวสลับสี
    col_aligns : list ของ WD_ALIGN_PARAGRAPH ต่อคอลัมน์
    """
    tbl           = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style     = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    for j, h in enumerate(headers):
        c = tbl.rows[0].cells[j]
        _cell(c, h, bold=True, size=13,
              fill=hdr_bg,
              color=RGBColor(0xFF, 0xFF, 0xFF) if hdr_bg == _HDR_NAVY else None)

    # data rows
    for i, row_data in enumerate(rows):
        bg = alt_bg if i % 2 == 1 else None
        for j, val in enumerate(row_data):
            c     = tbl.rows[i + 1].cells[j]
            align = (col_aligns[j] if col_aligns and j < len(col_aligns)
                     else WD_ALIGN_PARAGRAPH.CENTER)
            _cell(c, str(val), align=align, size=13, fill=bg)

    return tbl


def _add_figure(doc, img_bytes: bytes, caption: str = '', width_cm: float = 14):
    """เพิ่มรูปภาพพร้อม caption"""
    if not img_bytes:
        return
    doc.add_paragraph()
    p           = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r           = p.add_run()
    r.add_picture(io.BytesIO(img_bytes), width=Cm(width_cm))
    if caption:
        p2           = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p2, caption, italic=True, size=13, color=RGBColor(0x44, 0x44, 0x44))


def _add_footer(doc):
    doc.add_paragraph()
    p           = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "พัฒนาโดย รศ.ดร.อิทธิพล มีผล  |  ภาควิชาครุศาสตร์โยธา  |  มจพ.",
         size=12, color=RGBColor(80, 80, 80))


# ─────────────────────────────────────────────
#  Section 1: ESAL
# ─────────────────────────────────────────────

def build_report_esal(ss: dict) -> bytes | None:
    doc = _new_doc()
    if not doc:
        return None

    _heading(doc, "1. ผลการคำนวณ ESAL (Equivalent Single Axle Load)", size=16)
    p = _para(doc)
    _run(p, f"วิธีคำนวณ AASHTO 1993  |  LDF = {ss.get('ldf', 0.9)}  |  DDF = {ss.get('ddf', 0.5)}", size=14)

    if ss.get('esal_rigid'):
        _heading(doc, "1.1 ESAL – ผิวทางคอนกรีต", level=2, size=14)
        rows = [[lbl, f"{ss['esal_rigid'].get(D, 0):,.0f}"]
                for D, lbl in zip(SLAB_THICKNESSES, SLAB_LABELS)]
        _make_table(doc,
                    ["Slab Thickness", "ESAL (Design Lane)"],
                    rows)
        doc.add_paragraph()

    if ss.get('esal_flex'):
        _heading(doc, "1.2 ESAL – ผิวทางลาดยาง", level=2, size=14)
        rows = [[f"SN = {sn}", f"{v:,.0f}"] for sn, v in ss['esal_flex'].items()]
        _make_table(doc,
                    ["Structure Number (SN)", "ESAL (Design Lane)"],
                    rows)

    _add_footer(doc)
    return _doc_to_bytes(doc)


# ─────────────────────────────────────────────
#  Section 2: CBR
# ─────────────────────────────────────────────

def build_report_cbr(ss: dict) -> bytes | None:
    doc      = _new_doc()
    cbr_vals = ss.get('cbr_values', [])
    if not doc or not cbr_vals:
        return None

    _heading(doc, "2. ผลการวิเคราะห์ค่า CBR", size=16)

    arr, n, u_cbr, u_pct = calc_percentile_cbr(cbr_vals)
    pct   = ss.get('cbr_percentile', 90)
    cbr_d = float(np.interp(pct, u_pct[::-1], u_cbr[::-1]))
    mr_d  = 1500.0 * cbr_d
    k_d   = mr_d / 19.4

    # สถิติ
    p = _para(doc)
    _run(p, f"จำนวนตัวอย่าง n = {n}  |  Min = {np.min(cbr_vals):.2f}%  |  "
            f"Max = {np.max(cbr_vals):.2f}%  |  Mean = {np.mean(cbr_vals):.2f}%", size=14)

    # ผลลัพธ์
    p2 = _para(doc)
    _run(p2, f"CBR ออกแบบ @ P{pct:.0f} = ", size=15, bold=True)
    _run(p2, f"{cbr_d:.2f}%", size=15, bold=True, color=_BLUE)
    _run(p2, f"  →  Mr = {mr_d:,.0f} psi  →  k_subgrade = {k_d:.1f} pci", size=14)

    doc.add_paragraph()

    # ตาราง percentile
    _heading(doc, "ตารางค่า CBR และ Percentile", level=2, size=14)
    rows = []
    for i, (v, p2v) in enumerate(zip(u_cbr, u_pct)):
        cnt    = int(np.sum(arr >= v))
        mr_val = 1500.0 * v
        rows.append([str(i + 1), f"{v:.2f}", str(cnt), f"{p2v:.1f}", f"{mr_val:,.0f}"])

    _make_table(doc,
                ["ลำดับ", "CBR (%)", "จำนวน ≥", "Percentile (%)", "Mr (psi)"],
                rows,
                col_aligns=[WD_ALIGN_PARAGRAPH.CENTER] * 5)

    _add_footer(doc)
    return _doc_to_bytes(doc)


# ─────────────────────────────────────────────
#  Section 3: Flexible Pavement
# ─────────────────────────────────────────────

def build_report_flexible(ss: dict) -> bytes | None:
    doc = _new_doc()
    res = ss.get('flex_results', {})
    if not doc or not res:
        return None

    esal    = res.get('esal', 0)
    sn_req  = res.get('sn_req', 0)
    sn_prov = res.get('sn_prov', 0)
    passed  = res.get('pass', False)
    layers  = res.get('layers', [])
    cbr     = ss.get('cbr_design', 3.0)
    mr_psi  = ss.get('mr_subgrade_psi', 1500.0 * cbr)
    r0      = ss.get('r0_flex', 90)
    so      = ss.get('so_flex', 0.45)
    pi      = ss.get('pi_flex', 4.2)
    pt      = ss.get('pt_global', 2.5)
    margin  = sn_prov - sn_req
    total_h = sum(l.get('h_cm', 0) for l in layers)

    # ── หัวข้อ ──
    _heading(doc, "3. ผลการออกแบบโครงสร้างชั้นทางลาดยาง (Flexible Pavement)", size=16)

    # ── เกริ่นนำ ──
    p_intro = _para(doc, space_before=6)
    p_intro.paragraph_format.first_line_indent = Cm(1.25)
    _run(p_intro, "การออกแบบโครงสร้างถนนยืดหยุ่น (Flexible Pavement) ใช้วิธี AASHTO 1993 "
         "Guide for Design of Pavement Structures โดยพิจารณาปัจจัยด้านปริมาณจราจรสะสม ESALs "
         "ความน่าเชื่อถือ และคุณสมบัติของดินรองรับ ผลการออกแบบได้โครงสร้างชั้นทาง "
         f"{len(layers)} ชั้น รวมความหนา ")
    _run(p_intro, f"{total_h:.0f} ซม.", bold=True, color=_BLUE)
    _run(p_intro, f"  SN Required = ")
    _run(p_intro, f"{sn_req:.3f}", bold=True, color=_BLUE)
    _run(p_intro, f"  SN Provided = ")
    _run(p_intro, f"{sn_prov:.3f}", bold=True, color=_BLUE)
    _run(p_intro, "  การออกแบบ")
    _run(p_intro, "ผ่านเกณฑ์" if passed else "ไม่ผ่านเกณฑ์",
         bold=True, color=_GREEN if passed else _RED)

    # ── สมการ AASHTO 1993 ──
    doc.add_paragraph()
    _heading(doc, "3.1 วิธีการออกแบบ", level=2, size=14)
    p_eq = _para(doc, indent_cm=1.0)
    _eq_run(p_eq,
            "log₁₀(W₁₈) = Zᵣ·S₀ + 9.36·log₁₀(SN+1) − 0.20\n"
            "                   + log₁₀(ΔPSI/2.7) / [0.4 + 1094/(SN+1)⁵·¹⁹]\n"
            "                   + 2.32·log₁₀(Mᵣ) − 8.07",
            size=11, italic=True)

    # ── ตารางพารามิเตอร์ ──
    doc.add_paragraph()
    _heading(doc, "3.2 ค่าพารามิเตอร์การออกแบบ", level=2, size=14)
    _make_table(doc,
                ["พารามิเตอร์", "ค่า", "หน่วย"],
                [
                    ["Design ESALs (W₁₈)",     f"{esal:,.0f}",    "18-kip ESAL"],
                    ["Reliability (R₀)",         f"{r0}",           "%"],
                    ["Overall Std. Deviation (S₀)", f"{so:.2f}",   "-"],
                    ["Initial Serviceability (Pᵢ)", f"{pi:.1f}",   "-"],
                    ["Terminal Serviceability (Pₜ)", f"{pt:.1f}",  "-"],
                    ["ΔPSI",                     f"{pi - pt:.1f}", "-"],
                    ["CBR ดินเดิม",              f"{cbr:.2f}",     "%"],
                    ["Mᵣ = 1,500 × CBR",         f"{mr_psi:,.0f}", "psi"],
                ],
                col_aligns=[WD_ALIGN_PARAGRAPH.LEFT,
                            WD_ALIGN_PARAGRAPH.CENTER,
                            WD_ALIGN_PARAGRAPH.CENTER])

    # ── ตารางวัสดุ ──
    if layers:
        doc.add_paragraph()
        _heading(doc, "3.3 คุณสมบัติวัสดุชั้นทาง", level=2, size=14)
        rows_mat = []
        for l in layers:
            h_in  = l.get('h_cm', 0) / 2.54
            mr_l  = l.get('Mr_psi', None) or mr_psi
            rows_mat.append([
                str(l.get('layer', '')),
                l.get('material', ''),
                f"{l.get('ai', 0):.2f}",
                f"{l.get('mi', 1.0):.2f}",
                f"{l.get('h_cm', 0):.0f}",
                f"{h_in:.2f}",
                f"{mr_l:,.0f}",
            ])
        _make_table(doc,
                    ["ชั้น", "วัสดุ", "aᵢ", "mᵢ", "Dᵢ (cm)", "Dᵢ (in)", "Mᵣ (psi)"],
                    rows_mat,
                    col_aligns=[
                        WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.LEFT,
                        WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.CENTER,
                    ])

        # ── ตารางสรุป SN ──
        doc.add_paragraph()
        _heading(doc, "3.4 สรุปผลการคำนวณ Structural Number", level=2, size=14)
        rows_sn = []
        for l in layers:
            h_in = l.get('h_cm', 0) / 2.54
            rows_sn.append([
                str(l.get('layer', '')),
                l.get('material', ''),
                f"{l.get('ai', 0):.2f}",
                f"{l.get('mi', 1.0):.2f}",
                f"{h_in:.2f}",
                f"{l.get('h_cm', 0):.0f}",
                f"{l.get('sni', 0):.3f}",
                f"{l.get('cum_sn', 0):.3f}",
            ])
        _make_table(doc,
                    ["ชั้น", "วัสดุ", "aᵢ", "mᵢ", "Dᵢ (in)", "Dᵢ (cm)", "ΔSNᵢ", "ΣSN"],
                    rows_sn,
                    col_aligns=[
                        WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.LEFT,
                        WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.CENTER,
                    ])

    # ── ตารางผลตรวจสอบ ──
    doc.add_paragraph()
    _heading(doc, "3.5 ผลการตรวจสอบการออกแบบ", level=2, size=14)
    _make_table(doc,
                ["รายการ", "ค่า"],
                [
                    ["SN Required (จากสมการ AASHTO)", f"{sn_req:.3f}"],
                    ["SN Provided (จากชั้นทาง)",      f"{sn_prov:.3f}"],
                    ["Safety Margin (SN Provided − Required)", f"{margin:+.3f}"],
                    ["ผลการตรวจสอบ", "✅ PASS" if passed else "❌ FAIL"],
                ],
                col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER])

    # ── รูปโครงสร้าง ──
    if layers:
        fig_layers = [{"name": l.get("material", ""),
                       "thickness_cm": l.get("h_cm", 0),
                       "ai":  l.get("ai",  None),
                       "sni": l.get("sni", None)}
                      for l in layers]
        from engine.figures import draw_pavement_structure, fig_to_bytes
        import matplotlib.pyplot as plt
        fig = draw_pavement_structure(fig_layers, mode="flex",
                                      cbr_subgrade=float(cbr or 3.0))
        if fig:
            img_bytes = fig_to_bytes(fig)
            plt.close(fig)
            _add_figure(doc, img_bytes,
                        caption="รูปที่ 3-1  โครงสร้างชั้นทางลาดยางที่ออกแบบ")

    _add_footer(doc)
    return _doc_to_bytes(doc)


# ─────────────────────────────────────────────
#  Section 4: K-Value Nomograph
# ─────────────────────────────────────────────

def build_report_kvalue(ss: dict) -> bytes | None:
    doc = _new_doc()
    if not doc:
        return None

    _heading(doc, "4. ค่า k_eff (Effective Modulus of Subgrade Reaction)", size=16)
    p = _para(doc)
    _run(p, f"k∞ = {ss.get('k_inf', 0):.1f} pci  |  "
            f"LS = {ss.get('ls_value', 0)}  |  "
            f"k_eff = {ss.get('k_corrected', 0):.1f} pci", size=14)

    _make_table(doc,
                ["พารามิเตอร์", "ค่า", "หน่วย"],
                [
                    ["Composite k∞",          f"{ss.get('k_inf', 0):.1f}",       "pci"],
                    ["Loss of Support (LS)",   f"{ss.get('ls_value', 0)}",         "-"],
                    ["k_eff (Design k-value)", f"{ss.get('k_corrected', 0):.1f}", "pci"],
                ],
                col_aligns=[WD_ALIGN_PARAGRAPH.LEFT,
                            WD_ALIGN_PARAGRAPH.CENTER,
                            WD_ALIGN_PARAGRAPH.CENTER])

    _add_figure(doc, ss.get('nomograph_img_k'),
                caption="รูปที่ 4-1  Composite k∞ Nomograph (AASHTO 1993 Fig.3.3)")
    _add_figure(doc, ss.get('nomograph_img_ls'),
                caption="รูปที่ 4-2  Loss of Support Nomograph (AASHTO 1993 Fig.3.7)")

    _add_footer(doc)
    return _doc_to_bytes(doc)


# ─────────────────────────────────────────────
#  Section 5: Rigid Pavement
# ─────────────────────────────────────────────

def build_report_rigid(ss: dict) -> bytes | None:
    doc       = _new_doc()
    rigid_res = ss.get('rigid_results', {})
    if not doc or not rigid_res:
        return None

    _heading(doc, "5. ผลการออกแบบความหนาถนนคอนกรีต (AASHTO 1993)", size=16)

    # ── สมการ AASHTO 1993 Rigid ──
    p_eq = _para(doc, indent_cm=1.0)
    _eq_run(p_eq,
            "log₁₀(W₁₈) = Zᵣ·S₀ + 7.35·log₁₀(D+1) − 0.06\n"
            "                   + log₁₀(ΔPSI/3.0) / [1 + 1.624×10⁷/(D+1)⁸·⁴⁶]\n"
            "                   + (4.22 − 0.32Pₜ)·log₁₀[Sc·Cd·(D⁰·⁷⁵ − 1.132) / (215.63·J·(D⁰·⁷⁵ − 18.42/(Ec/k)⁰·²⁵))]",
            size=11, italic=True)

    for ptype, res in rigid_res.items():
        if not res:
            continue

        idx     = list(rigid_res.keys()).index(ptype) + 1
        d_sel   = res.get('d_cm', 30)
        k_eff   = res.get('k_eff', 200.0)
        fc      = res.get('fc', 350)
        sc      = res.get('sc', 600)
        j       = res.get('j', 2.8)
        cd      = res.get('cd', 1.0)
        w18_cap = res.get('w18_cap', 0)
        w18_req = res.get('w18_req', 0)
        passed  = res.get('pass', False)
        layers  = res.get('layers', [])
        e_eq    = res.get('e_eq_psi', 0)
        r0      = ss.get('r0_rig', 90)  or 90
        so      = float(ss.get('so_rig') or 0.35)
        pi      = float(ss.get('pi_rig') or 4.5)
        pt      = float(ss.get('pt_global') or ss.get('pt_rig_v7') or 2.5)
        fc_cyl  = 0.8 * fc
        fc_psi  = fc_cyl * 14.223
        ec_psi  = 57000 * math.sqrt(fc_psi)

        doc.add_paragraph()
        _heading(doc, f"5.{idx} {ptype}", level=2, size=15)

        # ── ตารางพารามิเตอร์ ──
        _heading(doc, f"5.{idx}.1 ค่าพารามิเตอร์การออกแบบ", level=3, size=14)
        _make_table(doc,
                    ["พารามิเตอร์", "ค่า", "หน่วย"],
                    [
                        ["Reliability (R₀)",              f"{r0}",            "%"],
                        ["Overall Std. Deviation (S₀)",   f"{so:.2f}",        "-"],
                        ["Initial Serviceability (Pᵢ)",   f"{pi:.1f}",        "-"],
                        ["Terminal Serviceability (Pₜ)",  f"{pt:.1f}",        "-"],
                        ["f'c (cube)",                     f"{fc}",            "ksc"],
                        ["f'c (cylinder) = 0.8×f'c",      f"{fc_cyl:.0f}",    "ksc"],
                        ["Ec = 57,000√f'c(psi)",          f"{ec_psi:,.0f}",   "psi"],
                        ["Modulus of Rupture (Sc)",        f"{sc}",            "psi"],
                        ["Load Transfer (J)",              f"{j:.1f}",         "-"],
                        ["Drainage Coefficient (Cd)",      f"{cd:.2f}",        "-"],
                        ["k_eff",                          f"{k_eff:.1f}",     "pci"],
                        ["Design ESAL (W₁₈)",             f"{w18_req:,.0f}",  "18-kip ESAL"],
                    ],
                    col_aligns=[WD_ALIGN_PARAGRAPH.LEFT,
                                WD_ALIGN_PARAGRAPH.CENTER,
                                WD_ALIGN_PARAGRAPH.CENTER])

        # ── ตารางชั้นทาง ──
        if layers:
            doc.add_paragraph()
            _heading(doc, f"5.{idx}.2 ชั้นโครงสร้างทาง", level=3, size=14)
            rows_lr = []
            for li, l in enumerate(layers):
                h_in = l.get('thickness_cm', 0) / 2.54
                rows_lr.append([
                    str(li + 1),
                    l.get('name', ''),
                    f"{l.get('thickness_cm', 0):.0f}",
                    f"{h_in:.2f}",
                    f"{l.get('E_MPa', 0):,}",
                    f"{l.get('thickness_cm', 0) * (l.get('E_MPa', 1) ** (1/3)):.1f}",
                ])
            _make_table(doc,
                        ["ชั้น", "วัสดุ", "หนา (cm)", "หนา (in)", "E (MPa)", "h×E^(1/3)"],
                        rows_lr,
                        col_aligns=[
                            WD_ALIGN_PARAGRAPH.CENTER,
                            WD_ALIGN_PARAGRAPH.LEFT,
                            WD_ALIGN_PARAGRAPH.CENTER,
                            WD_ALIGN_PARAGRAPH.CENTER,
                            WD_ALIGN_PARAGRAPH.CENTER,
                            WD_ALIGN_PARAGRAPH.CENTER,
                        ])
            if e_eq > 0:
                p_eq2 = _para(doc)
                dsb   = sum(l.get('thickness_cm', 0) for l in layers) / 2.54
                _run(p_eq2,
                     f"E_equivalent = {e_eq:,.0f} psi ({e_eq/145.038:.1f} MPa)  |  "
                     f"DSB = {dsb:.1f} in",
                     size=13, color=_BLUE)

        # ── ตาราง W18 ทุก Slab Thickness ──
        doc.add_paragraph()
        _heading(doc, f"5.{idx}.3 ตาราง W18 Capacity ทุกความหนา Slab", level=3, size=14)

        p_note = _para(doc)
        _run(p_note,
             f"คำนวณด้วย k_eff = {k_eff:.0f} pci  |  J = {j}  |  Cd = {cd}  |  "
             f"Sc = {sc} psi  |  Ec = {ec_psi:,.0f} psi",
             size=13, color=_DARK)

        rows_w18 = []
        for D, lbl in zip(SLAB_THICKNESSES, SLAB_LABELS):
            zr_use = float(ss.get('zr_rig') or ss.get('zr_rig_val') or -1.282)
            wc = aashto_rigid_w18(D, pi, pt, zr_use, so,
                                  sc, cd, j, ec_psi, k_eff)
            if wc is None:
                rows_w18.append([lbl, "-", f"{w18_req:,.0f}", "-",
                                  "⚠️", "← เลือก" if D == d_sel else ""])
                continue
            ratio  = wc / w18_req if w18_req > 0 else 0
            status = "✅ PASS" if ratio >= 1.0 else "❌ FAIL"
            tag    = "  ← เลือกใช้" if D == d_sel else ""
            rows_w18.append([
                lbl,
                f"{wc:,.0f}",
                f"{w18_req:,.0f}",
                f"{ratio:.3f}",
                status + tag,
            ])

        _make_table(doc,
                    ["Slab Thickness", "W18 Capacity", "W18 Required", "Ratio", "สถานะ"],
                    rows_w18,
                    col_aligns=[
                        WD_ALIGN_PARAGRAPH.LEFT,
                        WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.LEFT,
                    ])

        # ── ผลตรวจสอบ ──
        doc.add_paragraph()
        _heading(doc, f"5.{idx}.4 ผลการตรวจสอบ", level=3, size=14)
        margin_pct = (w18_cap / w18_req - 1) * 100 if w18_req > 0 else 0
        _make_table(doc,
                    ["รายการ", "ค่า"],
                    [
                        ["Slab ที่เลือก",       f"{d_sel} cm"],
                        ["W18 Capacity",         f"{w18_cap:,.0f}"],
                        ["W18 Required (ESAL)",  f"{w18_req:,.0f}"],
                        ["W18 Ratio",            f"{w18_cap/w18_req:.3f}" if w18_req > 0 else "-"],
                        ["Safety Margin",         f"{margin_pct:+.1f}%"],
                        ["ผลการตรวจสอบ",         "✅ PASS" if passed else "❌ FAIL"],
                    ],
                    col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER])

        # ── รูปโครงสร้าง ──
        img_key = f'rigid_structure_img_{ptype}'
        img_bytes = ss.get(img_key)
        if not img_bytes and layers:
            fig_layers = [{"name": l.get("name", ""),
                           "thickness_cm": l.get("thickness_cm", 0),
                           "E_MPa": l.get("E_MPa", 0)}
                          for l in layers]
            from engine.figures import draw_pavement_structure, fig_to_bytes
            import matplotlib.pyplot as plt
            fig = draw_pavement_structure(
                fig_layers, mode="rigid",
                cbr_subgrade=float(ss.get('cbr_design') or 3.0),
                d_concrete_cm=d_sel, ptype=ptype)
            if fig:
                img_bytes = fig_to_bytes(fig)
                import matplotlib.pyplot as plt
                plt.close(fig)

        _add_figure(doc, img_bytes,
                    caption=f"รูปที่ 5-{idx}  โครงสร้างชั้นทางคอนกรีต ({ptype})")

    _add_footer(doc)
    return _doc_to_bytes(doc)


# ─────────────────────────────────────────────
#  Full Report
# ─────────────────────────────────────────────

def build_report_full(ss: dict) -> bytes | None:
    doc = _new_doc()
    if not doc:
        return None

    # ── Cover page ──
    p           = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "รายการคำนวณออกแบบโครงสร้างชั้นทาง", bold=True, size=22)

    p2           = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p2, "ตามวิธี AASHTO 1993 Guide for Design of Pavement Structures", size=16)

    if ss.get('project_name'):
        p3           = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p3, f"โครงการ: {ss.get('project_name', '')}", size=15, bold=True, color=_BLUE)

    p4           = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p4, f"วันที่: {datetime.now().strftime('%d/%m/%Y %H:%M')}  |  "
             "พัฒนาโดย รศ.ดร.อิทธิพล มีผล  |  ภาควิชาครุศาสตร์โยธา มจพ.",
         size=13, color=RGBColor(80, 80, 80))

    doc.add_page_break()

    # ── Merge sections ──
    # ใช้ docxcompose.Composer แทนการ append XML element ดิบ
    # เพื่อให้รูปภาพ (relationships ใน document.xml.rels) ถูก copy
    # ตามไปด้วยอย่างถูกต้อง — แก้ปัญหารูปหายในรายงานรวม
    #
    # เรียก builder "ตัวเต็ม" จากไฟล์ report ย่อยโดยตรง (ตัวเดียวกับที่
    # รายงานเดี่ยวใช้) แทน build_report_* เวอร์ชันย่อใน report.py ที่
    # generate รูปผ่าน engine.figures (ซึ่ง import ไม่สำเร็จ ทำให้รูปหาย)
    # ผลลัพธ์: รายงานรวมได้เนื้อหา + รูป เหมือนรายงานเดี่ยวทุกประการ
    # หมายเหตุ: section K-value รวมอยู่ใน build_rigid_report แล้ว
    # จึงไม่ต้องแยกออกมาเป็น section ต่างหาก
    # lazy import — กันพังตอน startup (เหมือน figures/matplotlib)
    from docxcompose.composer import Composer
    from report_esal     import build_esal_report
    from report_cbr      import build_cbr_report
    from report_flexible import build_flexible_report
    from report_rigid    import build_rigid_report

    sections = [
        ('esal_rigid',    build_esal_report),
        ('cbr_values',    build_cbr_report),
        ('flex_results',  build_flexible_report),
        ('rigid_results', build_rigid_report),
    ]

    composer = Composer(doc)
    _first = True
    for key, fn in sections:
        if ss.get(key):
            sub_bytes = fn(ss)
            if sub_bytes:
                # ขึ้นหน้าใหม่ก่อนแต่ละ section (section แรกมี page break
                # จาก cover อยู่แล้ว จึงข้ามไป)
                if not _first:
                    doc.add_page_break()
                sub_doc = DocxDoc(io.BytesIO(sub_bytes))
                composer.append(sub_doc)
                _first = False

    _add_footer(doc)

    out = io.BytesIO()
    composer.save(out)
    out.seek(0)
    return out.read()
