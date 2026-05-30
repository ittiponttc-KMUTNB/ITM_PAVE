# ╔══════════════════════════════════════════════════════════════════╗
# ║  engine/report_esal.py — ITM Pave Pro                           ║
# ║  ESAL Word Report — format ตาม ESAL Calculator V3               ║
# ║  พัฒนาโดย รศ.ดร.อิทธิพล มีผล | ภาควิชาครุศาสตร์โยธา มจพ.    ║
# ╚══════════════════════════════════════════════════════════════════╝

import io
import re
import math
from datetime import datetime

from docx import Document as DocxDoc
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement

from constants import VEHICLE_AXLES, VEHICLE_COLS, VEHICLE_LABELS
from constants import SLAB_THICKNESSES, SN_DEFAULTS

# ─────────────────────────────────────────────
#  Constants
# ─────────────────────────────────────────────
FN       = 'TH SarabunPSK'
FS       = 15
TFS      = 14
TFS_SM   = 12
EQ_FONT  = 'Times New Roman'
EQ_SIZE  = 11
HDR_COLOR  = 'D9E2F3'
FOOT_COLOR = 'C6EFCE'

TRUCKS = {
    'MB':  'Medium Bus (รถโดยสารขนาดกลาง)',
    'HB':  'Heavy Bus (รถโดยสารขนาดใหญ่)',
    'MT':  'Medium Truck (รถบรรทุกขนาดกลาง)',
    'HT':  'Heavy Truck (รถบรรทุกขนาดใหญ่)',
    'TR':  'Full Trailer (รถพ่วง)',
    'STR': 'Semi-Trailer (รถกึ่งพ่วง)',
}

TON_TO_KIP = 2.2046


# ─────────────────────────────────────────────
#  Calculation helpers
# ─────────────────────────────────────────────

def _ealf_flex(L1_ton, L2, SN, pt):
    L1  = L1_ton * TON_TO_KIP
    Gt  = math.log10((4.2 - pt) / (4.2 - 1.5))
    Bx  = 0.40 + 0.081*(L1+L2)**3.23 / ((SN+1)**5.19 * L2**3.23)
    B18 = 0.40 + 0.081*(18+1)**3.23  / ((SN+1)**5.19 * 1.0**3.23)
    return 10**(4.79*math.log10(L1+L2) - 4.33*math.log10(L2)
                - 4.79*math.log10(19) + Gt*(1/B18 - 1/Bx))


def _ealf_rigid(L1_ton, L2, D_in, pt):
    L1  = L1_ton * TON_TO_KIP
    Gt  = math.log10((4.5 - pt) / (4.5 - 1.5))
    Bx  = 1.0 + 3.63*(L1+L2)**5.20 / ((D_in+1)**8.46 * L2**3.52)
    B18 = 1.0 + 3.63*(18+1)**5.20  / ((D_in+1)**8.46 * 1.0**3.52)
    return 10**(4.62*math.log10(L1+L2) - 3.28*math.log10(L2)
                - 4.62*math.log10(19) + Gt*(1/B18 - 1/Bx))


def _truck_factor(code, ptype, pt, param):
    axles = VEHICLE_AXLES[code]
    if ptype == 'flexible':
        return sum(_ealf_flex(L1, L2, param, pt)*cnt for L1, L2, cnt in axles)
    else:
        return sum(_ealf_rigid(L1, L2, param, pt)*cnt for L1, L2, cnt in axles)


def _calc_esal_table(traffic_df, ptype, pt, param, ldf, ddf):
    """คำนวณ ESAL รายปี + ACC.ESAL → return list of dict"""
    tf = {code: _truck_factor(code, ptype, pt, param) for code in TRUCKS}
    rows     = []
    acc_esal = 0.0
    for _, row in traffic_df.iterrows():
        year = int(row.get('Year', len(rows)+1))
        adtt = sum(int(row.get(c, 0)) for c in TRUCKS)
        yr_esal = sum(float(row.get(c, 0)) * tf[c] * ldf * ddf * 365
                      for c in TRUCKS)
        acc_esal += yr_esal
        rows.append({
            'Year': year, 'MB': int(row.get('MB', 0)), 'HB': int(row.get('HB', 0)),
            'MT': int(row.get('MT', 0)), 'HT': int(row.get('HT', 0)),
            'TR': int(row.get('TR', 0)), 'STR': int(row.get('STR', 0)),
            'ADTT': adtt, 'ESAL': int(round(yr_esal)),
            'ACC_ESAL': int(round(acc_esal)),
        })
    return rows, int(round(acc_esal))


def _increment_tbl(base, offset):
    s = base.strip()
    for pattern in [r'^(\d+\.\d+)-(\d+)$', r'^(\d+)-(\d+)$']:
        m = re.match(pattern, s)
        if m:
            return f"{m.group(1)}-{int(m.group(2))+offset}"
    m = re.match(r'^(\d+)\.(\d+)$', s)
    if m:
        return f"{m.group(1)}.{int(m.group(2))+offset}"
    m = re.match(r'^(\d+)$', s)
    if m:
        return str(int(m.group(1))+offset)
    return f"{base}+{offset}"


# ─────────────────────────────────────────────
#  Document helpers
# ─────────────────────────────────────────────

def _new_doc():
    doc   = DocxDoc()
    style = doc.styles['Normal']
    style.font.name = FN
    style.font.size = Pt(FS)
    try:
        style._element.rPr.rFonts.set(qn('w:eastAsia'), FN)
    except Exception:
        pass
    sec = doc.sections[0]
    sec.page_width  = Cm(21.0); sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.0);  sec.right_margin  = Cm(2.0)
    sec.top_margin  = Cm(2.0);  sec.bottom_margin = Cm(2.0)
    return doc


def _run(para, text, bold=False, sz=FS, italic=False):
    r = para.add_run(text)
    r.font.name = FN; r.font.size = Pt(sz)
    r.bold = bold; r.italic = italic
    try: r._element.rPr.rFonts.set(qn('w:eastAsia'), FN)
    except Exception: pass
    return r


def _eq_run(para, text, sz=EQ_SIZE, bold=False, italic=True):
    r = para.add_run(text)
    r.font.name = EQ_FONT; r.font.size = Pt(sz)
    r.bold = bold; r.italic = italic
    try: r._element.rPr.rFonts.set(qn('w:eastAsia'), EQ_FONT)
    except Exception: pass
    return r


def _sub_run(para, text, fname=EQ_FONT, sz=EQ_SIZE-1):
    r = para.add_run(text)
    r.font.name = fname; r.font.size = Pt(sz)
    r.font.subscript = True; r.italic = True
    try: r._element.rPr.rFonts.set(qn('w:eastAsia'), fname)
    except Exception: pass
    return r


def _sup_run(para, text, fname=EQ_FONT, sz=EQ_SIZE-1):
    r = para.add_run(text)
    r.font.name = fname; r.font.size = Pt(sz)
    r.font.superscript = True; r.italic = True
    try: r._element.rPr.rFonts.set(qn('w:eastAsia'), fname)
    except Exception: pass
    return r


def _shd(cell, color):
    xml = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(xml)


def _set_cell(cell, bold=False, align=None, sz=TFS, shading=None):
    for para in cell.paragraphs:
        if align:
            para.alignment = align
        for r in para.runs:
            r.font.name = FN; r.font.size = Pt(sz); r.bold = bold
            try: r._element.rPr.rFonts.set(qn('w:eastAsia'), FN)
            except Exception: pass
    if shading:
        _shd(cell, shading)


def _add_caption(doc, num_str, title_str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"ตารางที่ {num_str}  "); _run_fmt(r1, bold=True)
    r2 = p.add_run(title_str);               _run_fmt(r2)


def _run_fmt(r, bold=False, sz=TFS):
    r.font.name = FN; r.font.size = Pt(sz); r.bold = bold
    try: r._element.rPr.rFonts.set(qn('w:eastAsia'), FN)
    except Exception: pass


def _thai_justify(para):
    pPr = para._element.get_or_add_pPr()
    jc  = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'thaiDistribute')
    pPr.append(jc)


def _remove_tbl_border(tbl):
    tblPr  = tbl._tbl.tblPr
    border = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),'none'); b.set(qn('w:sz'),'0')
        b.set(qn('w:space'),'0'); b.set(qn('w:color'),'auto')
        border.append(b)
    tblPr.append(border)


def _remove_cell_border(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    bd   = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),'none'); b.set(qn('w:sz'),'0')
        b.set(qn('w:space'),'0'); b.set(qn('w:color'),'auto')
        bd.append(b)
    tcPr.append(bd)


# ─────────────────────────────────────────────
#  สมการ ESAL
# ─────────────────────────────────────────────

def _add_esal_formula_block(doc, num_years):
    # ย่อหน้าเกริ่นนำ
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.first_line_indent = Cm(1.25)
    p_intro.paragraph_format.space_after = Pt(6)
    _thai_justify(p_intro)
    _run(p_intro, 'โดยที่ค่าปริมาณเพลามาตรฐาน Equivalent Single Axle Load 18 kips (W')
    r_sub = p_intro.add_run('18')
    r_sub.font.name = FN; r_sub.font.size = Pt(FS); r_sub.font.subscript = True
    try: r_sub._element.rPr.rFonts.set(qn('w:eastAsia'), FN)
    except Exception: pass
    _run(p_intro, ') หมายถึง ปริมาณการจราจรของรถบรรทุกมาตรฐานที่วิ่งผ่านช่องจราจรออกแบบ (')
    _run(p_intro, 'Design Lane', bold=True)
    _run(p_intro, ') ในช่วงระยะเวลาออกแบบ (')
    _run(p_intro, 'Design Period', bold=True)
    _run(p_intro, ') โดยสามารถคำนวณจากสมการ')

    # สมการที่ 1
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after  = Pt(2)
    _eq_run(p1, 'W')
    _sub_run(p1, '18')
    _eq_run(p1, 'j')
    _eq_run(p1, '  =  ', italic=False)
    _eq_run(p1, '{', sz=EQ_SIZE+2, italic=False)
    _sup_run(p1, '6 ', sz=EQ_SIZE-2)
    _eq_run(p1, '\u03A3', sz=EQ_SIZE+2, italic=False)
    _sub_run(p1, 'i=1', sz=EQ_SIZE-2)
    _eq_run(p1, ' A', italic=True)
    _sub_run(p1, 'i')
    _eq_run(p1, '}', sz=EQ_SIZE+2, italic=False)
    _eq_run(p1, '  \u00D7  TF  \u00D7  L  \u00D7  D  \u00D7  365', italic=False)

    # สมการที่ 2
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(10)
    _eq_run(p2, 'W')
    _sub_run(p2, '18')
    _eq_run(p2, '  =  ', italic=False)
    _sup_run(p2, f'k={num_years} ', sz=EQ_SIZE-2)
    _eq_run(p2, '\u03A3', sz=EQ_SIZE+2, italic=False)
    _sub_run(p2, 'j=1', sz=EQ_SIZE-2)
    _eq_run(p2, '  W')
    _sub_run(p2, '18')
    _eq_run(p2, 'j', italic=True)

    # ตาราง "โดยที่"
    p_lbl = doc.add_paragraph()
    p_lbl.paragraph_format.space_after = Pt(2)
    _run(p_lbl, 'โดยที่')

    legend = [
        ([('W', False, False), ('18', True, False)],
         f'ผลรวมปริมาณเพลาเดี่ยวมาตรฐานออกแบบขนาด 18 kip ถึงอายุออกแบบ (k) เท่ากับ {num_years} ปี'),
        ([('W', False, False), ('18', True, False), ('j', True, False)],
         'ผลรวมปริมาณเพลาเดี่ยวมาตรฐานออกแบบขนาด 18 kip ณ ปีใดๆ (j)'),
        ([('TF', False, False)],
         'Truck Factor พิจารณาตามข้อมูลด้านชั่งน้ำหนักบริเวณพื้นที่โครงการ หรือใกล้เคียง'),
        ([('A', False, False), ('i', True, False)],
         'ปริมาณรถบรรทุกประเภทที่ i หน่วย คัน/วัน จากข้อมูลการคาดการณ์ปริมาณจราจรของโครงการ'),
        ([('L', False, False)], 'Lane Distribution Factor'),
        ([('D', False, False)], 'Directional Distribution Factor'),
    ]

    leg_tbl = doc.add_table(rows=len(legend), cols=3)
    leg_tbl.style = 'Table Grid'
    _remove_tbl_border(leg_tbl)

    tblPr  = leg_tbl._tbl.tblPr
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), '1134'); tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)

    tblGrid = OxmlElement('w:tblGrid')
    for w in ['1100', '400', '6700']:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), w); tblGrid.append(gc)
    leg_tbl._tbl.insert(0, tblGrid)

    for row_idx, (sym_runs, desc) in enumerate(legend):
        row = leg_tbl.rows[row_idx]
        c0, c1, c2 = row.cells[0], row.cells[1], row.cells[2]
        for cell in (c0, c1, c2):
            _remove_cell_border(cell)
            tcPr   = cell._tc.get_or_add_tcPr()
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'top')
            tcPr.append(vAlign)

        p_sym = c0.paragraphs[0]
        p_sym.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for txt, is_sub, is_sup in sym_runs:
            r = p_sym.add_run(txt)
            r.font.name = EQ_FONT; r.font.size = Pt(EQ_SIZE); r.italic = True
            r.font.subscript   = is_sub
            r.font.superscript = is_sup
            try: r._element.rPr.rFonts.set(qn('w:eastAsia'), EQ_FONT)
            except Exception: pass

        p_eq = c1.paragraphs[0]
        p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_eq = p_eq.add_run('=')
        r_eq.font.name = FN; r_eq.font.size = Pt(FS); r_eq.italic = False
        try: r_eq._element.rPr.rFonts.set(qn('w:eastAsia'), FN)
        except Exception: pass

        p_desc = c2.paragraphs[0]
        _thai_justify(p_desc)
        r_desc = p_desc.add_run(desc)
        r_desc.font.name = FN; r_desc.font.size = Pt(FS); r_desc.italic = False
        try: r_desc._element.rPr.rFonts.set(qn('w:eastAsia'), FN)
        except Exception: pass

    doc.add_paragraph()


# ─────────────────────────────────────────────
#  สร้าง section (Flexible หรือ Rigid)
# ─────────────────────────────────────────────

def _build_esal_section(doc, traffic_df, ptype, pt, param_list,
                         ldf, ddf, section_num, tbl_base):
    """สร้าง 1 section ESAL (Flexible หรือ Rigid) พร้อม 4 ตาราง"""
    pavement_text = 'Rigid Pavement'   if ptype == 'rigid'    else 'Flexible Pavement'
    pavement_thai = 'แบบแข็ง'          if ptype == 'rigid'    else 'ยืดหยุ่น'
    num_years     = len(traffic_df)

    tbl_param = tbl_base
    tbl_tf    = _increment_tbl(tbl_base, 1)
    tbl_traf  = _increment_tbl(tbl_base, 2)
    tbl_esal  = _increment_tbl(tbl_base, 3)

    # label ต่อ param
    _D_CM = {10:25, 11:28, 12:30, 13:32, 14:35}
    def p_lbl(p):
        if ptype == 'rigid':
            cm = _D_CM.get(p, '')
            return f'D = {p} นิ้ว ({cm} cm)' if cm else f'D = {p} นิ้ว'
        return f'SN = {p}'
    def p_col(p):
        if ptype == 'rigid':
            cm = _D_CM.get(p, '')
            return f'D={p}" ({cm} cm)' if cm else f'D={p}"'
        return f'SN={p}'

    # ── หัวข้อ ──
    h = doc.add_paragraph()
    h.paragraph_format.space_after = Pt(6)
    _run(h, f"{section_num}\t", bold=True)
    _run(h, f"ปริมาณเพลามาตรฐาน (ESALs) ระยะเวลาออกแบบ {num_years} ปี ผิวทาง {pavement_text}",
         bold=True)

    # ── บทเกริ่นนำ ──
    intro = doc.add_paragraph()
    intro.paragraph_format.first_line_indent = Cm(1.25)
    intro.paragraph_format.space_after = Pt(6)
    _thai_justify(intro)
    if ptype == 'flexible':
        txt = (f"ในการคำนวณปริมาณเพลามาตรฐาน สำหรับผิวทางยืดหยุ่น ที่ปรึกษาได้กำหนดค่าพารามิเตอร์ต่าง ๆ "
               f"และค่า Truck Factor ของรถบรรทุกหนัก ที่ใช้สำหรับการคำนวณ ดังแสดงในตารางที่ {tbl_param} และ {tbl_tf} "
               f"ดังนั้นค่าปริมาณเพลามาตรฐาน สำหรับผิวทางยืดหยุ่น ที่ระยะเวลาออกแบบ {num_years} ปี "
               f"แสดงดังตารางที่ {tbl_traf} และ {tbl_esal}")
    else:
        txt = (f"ในการคำนวณปริมาณเพลามาตรฐานสำหรับผิวทางแบบแข็งหรือผิวทางคอนกรีต "
               f"โดยที่ปรึกษาได้กำหนดค่าพารามิเตอร์ต่าง ๆ และค่า Truck Factor ของรถบรรทุกหนัก "
               f"ที่ใช้สำหรับการคำนวณ ดังแสดงในตารางที่ {tbl_param} และ {tbl_tf} "
               f"ดังนั้นค่าปริมาณเพลามาตรฐาน สำหรับผิวทางแบบแข็ง ที่ระยะเวลาออกแบบ {num_years} ปี "
               f"แสดงดังตารางที่ {tbl_traf} และ {tbl_esal}")
    intro.add_run(txt).font.name = FN
    doc.add_paragraph()

    # ══════════════════════════════════════════
    # ตารางที่ tbl_param — พารามิเตอร์
    # ══════════════════════════════════════════
    _add_caption(doc, tbl_param, 'ค่าพารามิเตอร์ต่าง ๆ ที่ใช้สำหรับการคำนวณ')
    params_str = ' | '.join(p_lbl(p) for p in param_list)
    param_data = [
        ('รายการ',            'ค่า'),
        ('ประเภทผิวทาง',      pavement_text),
        ('pt',                str(pt)),
        ('พารามิเตอร์',       params_str),
        ('Lane Factor',       str(ldf)),
        ('Direction Factor',  str(ddf)),
        ('จำนวนปี',           str(num_years)),
    ]
    t1 = doc.add_table(rows=len(param_data), cols=2)
    t1.style = 'Table Grid'
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (lbl, val) in enumerate(param_data):
        t1.rows[i].cells[0].text = lbl
        t1.rows[i].cells[1].text = val
        sh = HDR_COLOR if i == 0 else None
        _set_cell(t1.rows[i].cells[0], bold=(i==0), sz=TFS, shading=sh)
        _set_cell(t1.rows[i].cells[1], bold=(i==0), sz=TFS, shading=sh)
    doc.add_paragraph()

    # ══════════════════════════════════════════
    # ตารางที่ tbl_tf — Truck Factor
    # ══════════════════════════════════════════
    _add_caption(doc, tbl_tf, 'ค่า Truck Factor ของรถบรรทุกหนัก')
    tf_headers = ['รหัส', 'ประเภท'] + [p_col(p) for p in param_list]
    t2 = doc.add_table(rows=len(TRUCKS)+1, cols=len(tf_headers))
    t2.style = 'Table Grid'
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(tf_headers):
        t2.rows[0].cells[j].text = h
        _set_cell(t2.rows[0].cells[j], bold=True, sz=TFS,
                  align=WD_ALIGN_PARAGRAPH.CENTER, shading=HDR_COLOR)
    for i, (code, desc) in enumerate(TRUCKS.items()):
        row = t2.rows[i+1]
        row.cells[0].text = code
        row.cells[1].text = desc
        _set_cell(row.cells[0], sz=TFS, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(row.cells[1], sz=TFS)
        for k, p in enumerate(param_list):
            tf_val = _truck_factor(code, ptype, pt, p)
            row.cells[2+k].text = f'{tf_val:.4f}'
            _set_cell(row.cells[2+k], sz=TFS, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # ══════════════════════════════════════════
    # สมการ ESAL
    # ══════════════════════════════════════════
    _add_esal_formula_block(doc, num_years)

    # ══════════════════════════════════════════
    # ตารางที่ tbl_traf — ปริมาณจราจรรายปี
    # ══════════════════════════════════════════
    _add_caption(doc, tbl_traf,
                 f'ปริมาณจราจรรายปี ระยะเวลาออกแบบ {num_years} ปี')
    traf_headers = ['Year', 'MB', 'HB', 'MT', 'HT', 'TR', 'STR', 'ADTT']
    t3 = doc.add_table(rows=num_years+1, cols=len(traf_headers))
    t3.style = 'Table Grid'
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(traf_headers):
        t3.rows[0].cells[j].text = h
        _set_cell(t3.rows[0].cells[j], bold=True, sz=TFS_SM,
                  align=WD_ALIGN_PARAGRAPH.CENTER, shading=HDR_COLOR)

    # คำนวณ rows จาก param แรก
    first_rows, _ = _calc_esal_table(traffic_df, ptype, pt, param_list[0], ldf, ddf)
    for i, rd in enumerate(first_rows):
        row = t3.rows[i+1]
        vals = [rd['Year'], rd['MB'], rd['HB'], rd['MT'],
                rd['HT'], rd['TR'], rd['STR'], rd['ADTT']]
        for j, v in enumerate(vals):
            row.cells[j].text = f'{v:,}' if j > 0 else str(v)
            align = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            _set_cell(row.cells[j], sz=TFS_SM, align=align)
    doc.add_paragraph()

    # ══════════════════════════════════════════
    # ตารางที่ tbl_esal — ESAL รายปี + ACC.ESAL
    # ══════════════════════════════════════════
    _add_caption(doc, tbl_esal,
                 f'ค่าปริมาณเพลามาตรฐาน สำหรับผิวทาง{pavement_thai} '
                 f'ที่ระยะเวลาออกแบบ {num_years} ปี')
    esal_headers = ['Year', 'ADTT'] + [f'ESAL ({p_col(p)})' for p in param_list]
    t4 = doc.add_table(rows=num_years+2, cols=len(esal_headers))
    t4.style = 'Table Grid'
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(esal_headers):
        t4.rows[0].cells[j].text = h
        _set_cell(t4.rows[0].cells[j], bold=True, sz=TFS_SM,
                  align=WD_ALIGN_PARAGRAPH.CENTER, shading=HDR_COLOR)

    # คำนวณ ESAL ทุก param
    all_rows = {}
    all_totals = {}
    for p in param_list:
        rows_p, total_p = _calc_esal_table(traffic_df, ptype, pt, p, ldf, ddf)
        all_rows[p]   = rows_p
        all_totals[p] = total_p

    for i, rd in enumerate(all_rows[param_list[0]]):
        row = t4.rows[i+1]
        row.cells[0].text = str(rd['Year'])
        _set_cell(row.cells[0], sz=TFS_SM, align=WD_ALIGN_PARAGRAPH.CENTER)
        row.cells[1].text = f"{rd['ADTT']:,}"
        _set_cell(row.cells[1], sz=TFS_SM, align=WD_ALIGN_PARAGRAPH.RIGHT)
        for k, p in enumerate(param_list):
            esal_val = all_rows[p][i]['ESAL']
            row.cells[2+k].text = f'{esal_val:,}'
            _set_cell(row.cells[2+k], sz=TFS_SM, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # footer ACC.ESAL
    foot = t4.rows[num_years+1]
    foot.cells[0].text = 'ACC. ESAL'
    _set_cell(foot.cells[0], bold=True, sz=TFS_SM,
              align=WD_ALIGN_PARAGRAPH.CENTER, shading=FOOT_COLOR)
    foot.cells[1].text = ''
    _set_cell(foot.cells[1], sz=TFS_SM, shading=FOOT_COLOR)
    for k, p in enumerate(param_list):
        foot.cells[2+k].text = f'{all_totals[p]:,}'
        _set_cell(foot.cells[2+k], bold=True, sz=TFS_SM,
                  align=WD_ALIGN_PARAGRAPH.RIGHT, shading=FOOT_COLOR)
    doc.add_paragraph()


# ─────────────────────────────────────────────
#  Public functions
# ─────────────────────────────────────────────

def build_esal_report(ss: dict) -> bytes | None:
    """
    สร้าง Word ESAL Report แบบ consultant grade
    รองรับ Flexible + Rigid ในไฟล์เดียว

    ss ต้องมี:
      traffic_df     : pd.DataFrame
      ldf, ddf, pt_global
      esal_flex      : {SN: total_esal}
      esal_rigid     : {D_cm: total_esal}
      project_name   : str (optional)
      report_settings: dict (optional) — section_num, tbl_start
    """
    import pandas as pd

    traffic_df = ss.get('traffic_df')
    if traffic_df is None or len(traffic_df) == 0:
        return None

    ldf    = float(ss.get('ldf', 0.9))
    ddf    = float(ss.get('ddf', 0.5))
    pt_f   = float(ss.get('pt_flex', ss.get('pt_global', 2.5)))
    pt_r   = float(ss.get('pt_rigid', ss.get('pt_global', 2.5)))
    rs     = ss.get('report_settings', {})

    doc = _new_doc()

    # ── Cover ──
    p_cov           = doc.add_paragraph()
    p_cov.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_cov, 'รายการคำนวณปริมาณเพลามาตรฐาน (ESALs)', bold=True, sz=18)
    p_cov2           = doc.add_paragraph()
    p_cov2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_cov2, 'AASHTO 1993 Guide for Design of Pavement Structures', sz=14)
    if ss.get('project_name'):
        p_proj           = doc.add_paragraph()
        p_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p_proj, f"โครงการ: {ss.get('project_name', '')}", sz=14, bold=True)
    p_date           = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_date, f"วันที่: {datetime.now().strftime('%d/%m/%Y %H:%M')}", sz=13)
    doc.add_page_break()

    # ── Flexible section ──
    if ss.get('esal_flex'):
        sn_list = list(ss['esal_flex'].keys())
        # convert SN key → float
        sn_list = sorted([float(s) for s in sn_list])
        _build_esal_section(
            doc, traffic_df, 'flexible', pt_f, sn_list, ldf, ddf,
            section_num=rs.get('flex_section_number', '4.2.2'),
            tbl_base=rs.get('flex_table_start', '4-1'),
        )
        doc.add_page_break()

    # ── Rigid section ──
    if ss.get('esal_rigid'):
        # D_cm → D_in
        D_CM_TO_IN = {25:10, 28:11, 30:12, 32:13, 35:14}
        d_in_list  = sorted([D_CM_TO_IN.get(int(d), round(int(d)/2.54))
                             for d in ss['esal_rigid'].keys()])
        _build_esal_section(
            doc, traffic_df, 'rigid', pt_r, d_in_list, ldf, ddf,
            section_num=rs.get('rigid_section_number', '4.2.3'),
            tbl_base=rs.get('rigid_table_start', '4-4'),
        )

    # ── Footer ──
    doc.add_paragraph()
    fp           = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run(
        'พัฒนาโดย รศ.ดร.อิทธิพล มีผล  |  ภาควิชาครุศาสตร์โยธา  |  มจพ.')
    r.font.name = FN; r.font.size = Pt(13); r.italic = True
    try: r._element.rPr.rFonts.set(qn('w:eastAsia'), FN)
    except Exception: pass

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
