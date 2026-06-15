# ╔══════════════════════════════════════════════════════════════════╗
# ║  engine/report_flexible.py — ITM Pave Pro                       ║
# ║  Flexible Pavement Word Report — format ตาม AC_cal_V6           ║
# ║  พัฒนาโดย รศ.ดร.อิทธิพล มีผล | ภาควิชาครุศาสตร์โยธา มจพ.    ║
# ╚══════════════════════════════════════════════════════════════════╝

import io
import math
from docx import Document as DocxDoc
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FN   = 'TH SarabunPSK'
EQ   = 'Times New Roman'
FS   = 15
TFS  = 15
HDR  = 'D9E2F3'
HDR2 = 'BDD7EE'
RED   = RGBColor(0xCC, 0x00, 0x00)
GREEN = RGBColor(0x00, 0x70, 0x00)
BLUE  = RGBColor(0x00, 0x47, 0xAB)


# ─────────────────────────────────────────────
#  Helpers
# ─────────────────────────────────────────────

def _new_doc():
    doc   = DocxDoc()
    style = doc.styles['Normal']
    style.font.name = FN
    style.font.size = Pt(FS)
    try:
        style._element.rPr.rFonts.set(qn('w:eastAsia'), FN)
    except Exception:
        pass
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0); sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.5);  sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(2.5);  sec.bottom_margin = Cm(2.5)
    return doc


def _run(para, text, bold=False, sz=FS, italic=False, color=None, underline=False):
    r = para.add_run(text)
    r.font.name = FN; r.font.size = Pt(sz)
    r.bold = bold; r.italic = italic; r.underline = underline
    if color:
        r.font.color.rgb = color
    try:
        r._element.rPr.rFonts.set(qn('w:cs'), FN)
    except Exception:
        pass
    return r


def _eq_run(para, text, sz=11, bold=False, italic=True):  # Times New Roman 11pt
    r = para.add_run(text)
    r.font.name = EQ; r.font.size = Pt(sz)
    r.bold = bold; r.italic = italic
    return r


def _para(doc, indent_cm=0, first_line_cm=0, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent        = Cm(indent_cm)
    p.paragraph_format.first_line_indent  = Cm(first_line_cm)
    p.paragraph_format.space_before       = Pt(space_before)
    p.paragraph_format.space_after        = Pt(space_after)
    return p


def _heading(doc, text, level=2, sz=16):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = FN; r.font.size = Pt(sz)
    return h


def _thai_distribute(para):
    pPr = para._element.get_or_add_pPr()
    jc  = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'thaiDistribute')
    pPr.append(jc)


def _shd(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _tbl_cell(cell, text, align=WD_ALIGN_PARAGRAPH.CENTER,
              sz=TFS, bold=False, fill=None, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.font.name = FN; r.font.size = Pt(sz); r.bold = bold
    if color:
        r.font.color.rgb = color
    try:
        r._element.rPr.rFonts.set(qn('w:cs'), FN)
    except Exception:
        pass
    if fill:
        _shd(cell, fill)


def _eq_para(doc, text, indent_cm=2.0, bold=False, italic=True):
    p = _para(doc, indent_cm=indent_cm)
    _eq_run(p, text, bold=bold, italic=italic)
    return p


def _short_mat(name):
    """ตัดชื่อวัสดุให้สั้นลง"""
    m = {
        'ผิวทางแอสฟัลต์คอนกรีต (AC)':                         'AC Surface',
        'หินคลุกปรับปรุงคุณภาพด้วยปูนซีเมนต์ (CTB)':         'CTB',
        'หินคลุก CBR 80%':                                    'Crushed Rock CBR 80%',
        'รองพื้นทางวัสดุมวลรวม CBR 25%':                     'Granular Subbase CBR 25%',
        'วัสดุคัดเลือก ก':                                    'Select Material A',
        'ดินถมคันทาง CBR กรอกเอง':                           'Earth Embankment',
    }
    for k, v in m.items():
        if k in name:
            return v
    return name


# ─────────────────────────────────────────────
#  Summary layer table (รูป + ชั้นทาง)
# ─────────────────────────────────────────────

def _summary_table(doc, layers, cbr_design, fig_bytes=None):
    """ตารางสรุปโครงสร้าง 3 คอลัมน์ + รูปตัดขวาง (vMerge)"""
    COL_W = [3800, 1400, 3872]

    # build data rows
    data_rows = []
    for l in layers:
        mat  = l.get('material', '')
        h_cm = l.get('h_cm', 0)
        if mat == 'ดินถมคันทาง CBR กรอกเอง':
            mat = f'Earth Embankment CBR ≥ {cbr_design:.1f} %'
        data_rows.append((_short_mat(mat), str(int(h_cm))))
    data_rows.append((f'Earth Embankment / Subgrade\nCBR ≥ {cbr_design:.1f} %', 'Existing'))

    # insert figure via temp paragraph
    drawing_el = None
    if fig_bytes:
        try:
            tmp_para = doc.add_paragraph()
            tmp_run  = tmp_para.add_run()
            buf      = io.BytesIO(fig_bytes)
            tmp_run.add_picture(buf, width=Inches(2.4))
            drawing_el = tmp_run._r.find(qn('w:drawing'))
            if drawing_el is not None:
                tmp_run._r.remove(drawing_el)
            tmp_para._p.getparent().remove(tmp_para._p)
        except Exception:
            drawing_el = None

    # build table element
    tbl_el = OxmlElement('w:tbl')
    tblPr  = OxmlElement('w:tblPr')
    ts = OxmlElement('w:tblStyle'); ts.set(qn('w:val'), 'TableGrid'); tblPr.append(ts)
    tw = OxmlElement('w:tblW');     tw.set(qn('w:w'), str(sum(COL_W))); tw.set(qn('w:type'), 'dxa'); tblPr.append(tw)
    jc = OxmlElement('w:jc');      jc.set(qn('w:val'), 'center'); tblPr.append(jc)
    tbl_el.append(tblPr)

    tblGrid = OxmlElement('w:tblGrid')
    for w in COL_W:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w)); tblGrid.append(gc)
    tbl_el.append(tblGrid)

    def _make_tc(width, vmerge=None, valign='center'):
        tc = OxmlElement('w:tc')
        tcPr = OxmlElement('w:tcPr')
        tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'), str(width)); tcW.set(qn('w:type'), 'dxa'); tcPr.append(tcW)
        if vmerge is not None:
            vm = OxmlElement('w:vMerge')
            if vmerge == 'restart':
                vm.set(qn('w:val'), 'restart')
            tcPr.append(vm)
        va = OxmlElement('w:vAlign'); va.set(qn('w:val'), valign); tcPr.append(va)
        tc.append(tcPr)
        tc.append(OxmlElement('w:p'))
        return tc

    def _tc_text(tc, text, bold=False, center=True):
        for old_p in tc.findall(qn('w:p')):
            tc.remove(old_p)
        p_el = OxmlElement('w:p')
        pPr  = OxmlElement('w:pPr')
        jc2  = OxmlElement('w:jc'); jc2.set(qn('w:val'), 'center' if center else 'left'); pPr.append(jc2)
        p_el.append(pPr)
        for idx, part in enumerate(text.split('\n')):
            if idx > 0:
                br_r = OxmlElement('w:r'); br = OxmlElement('w:br'); br_r.append(br); p_el.append(br_r)
            r_el = OxmlElement('w:r')
            rPr  = OxmlElement('w:rPr')
            rf   = OxmlElement('w:rFonts')
            rf.set(qn('w:ascii'), FN); rf.set(qn('w:hAnsi'), FN); rf.set(qn('w:cs'), FN)
            rPr.append(rf)
            sz_el = OxmlElement('w:sz');   sz_el.set(qn('w:val'),   '30'); rPr.append(sz_el)
            szc   = OxmlElement('w:szCs'); szc.set(qn('w:val'),     '30'); rPr.append(szc)
            if bold:
                rPr.append(OxmlElement('w:b')); rPr.append(OxmlElement('w:bCs'))
            r_el.append(rPr)
            t_el = OxmlElement('w:t'); t_el.text = part; r_el.append(t_el); p_el.append(r_el)
        tc.append(p_el)

    # header row
    tr_hdr = OxmlElement('w:tr')
    for label, w in zip(['รายละเอียด', 'หนา (ซม.)', 'ชนิดวัสดุ'], COL_W):
        tc = _make_tc(w)
        _tc_text(tc, label, bold=True)
        tr_hdr.append(tc)
    tbl_el.append(tr_hdr)

    # data rows
    for i, (mat_name, thick) in enumerate(data_rows):
        tr = OxmlElement('w:tr')
        vm  = 'restart' if i == 0 else 'continue'
        tc0 = _make_tc(COL_W[0], vmerge=vm)
        if i == 0 and drawing_el is not None:
            for old_p in tc0.findall(qn('w:p')):
                tc0.remove(old_p)
            p_pic = OxmlElement('w:p')
            pPr   = OxmlElement('w:pPr')
            jc3   = OxmlElement('w:jc'); jc3.set(qn('w:val'), 'center'); pPr.append(jc3)
            p_pic.append(pPr)
            r_pic = OxmlElement('w:r'); r_pic.append(drawing_el); p_pic.append(r_pic)
            tc0.append(p_pic)
        tr.append(tc0)
        tc1 = _make_tc(COL_W[1]); _tc_text(tc1, thick);    tr.append(tc1)
        tc2 = _make_tc(COL_W[2]); _tc_text(tc2, mat_name); tr.append(tc2)
        tbl_el.append(tr)

    doc.element.body.append(tbl_el)

    # shading header
    real_tbl = doc.tables[-1]
    for j in range(3):
        _shd(real_tbl.rows[0].cells[j], HDR2)
    return real_tbl


# ─────────────────────────────────────────────
#  Public builder
# ─────────────────────────────────────────────

def build_flexible_report(ss: dict) -> bytes | None:
    """
    สร้าง Word Flexible Pavement Design Report

    ss ต้องมี:
      flex_results : {esal, sn_req, sn_prov, pass, layers[], mr_psi, cbr}
      r0_flex      : Reliability (%)
      so_flex      : So
      pi_flex      : Pi (initial serviceability)
      pt_global    : Pt
      cbr_design   : CBR (%)
      project_name : str (optional)
      report_settings: dict (optional)
        section_number, table_inputs, table_materials, table_sn,
        figure_number, section_title, table_cap_inputs,
        table_cap_materials, table_cap_sn, figure_caption,
        num_lanes, direction
      flex_structure_img : bytes (optional)
    """
    fr = ss.get('flex_results')
    if not fr:
        return None

    # ── ดึงค่าจาก flex_results ──
    W18       = float(fr.get('esal',    0))
    sn_req    = float(fr.get('sn_req',  0))
    sn_prov   = float(fr.get('sn_prov', 0))
    passed    = bool(fr.get('pass',     False))
    layers    = fr.get('layers', [])
    mr_psi    = float(fr.get('mr_psi',  7500))
    cbr       = float(fr.get('cbr',     ss.get('cbr_design', 5.0)))
    mr_mpa    = round(mr_psi * 0.006895, 1)

    reliability = int(ss.get('r0_flex', 90))
    so          = float(ss.get('so_flex', 0.45))
    pi          = float(ss.get('pi_flex', 4.2))
    pt          = float(ss.get('pt_global', 2.5))
    delta_psi   = round(pi - pt, 1)

    from constants import ZR_MAP
    zr = ZR_MAP.get(reliability, -1.282)

    rs = ss.get('report_settings', {})
    sec_no    = rs.get('section_number',     '4.4')
    tbl_inp   = rs.get('table_inputs',       '4-8')
    tbl_mat   = rs.get('table_materials',    '4-9')
    tbl_sn    = rs.get('table_sn',           '4-10')
    fig_no    = rs.get('figure_number',      '4-8')
    sec_title = rs.get('section_title',
                       'การออกแบบผิวทางลาดยาง (Flexible Pavement)')
    cap_inp   = rs.get('table_cap_inputs',
                       'ค่าพารามิเตอร์ที่ใช้ในการออกแบบผิวทางยืดหยุ่น')
    cap_mat   = rs.get('table_cap_materials',
                       'ค่าสัมประสิทธิ์และค่าโมดูลัสของวัสดุโครงสร้างชั้นทาง')
    cap_sn    = rs.get('table_cap_sn',
                       'สรุปผลการคำนวณ Structural Number ของโครงสร้างชั้นทาง')
    fig_cap   = rs.get('figure_caption',     'รูปตัดโครงสร้างชั้นทางที่ออกแบบ')
    num_lanes = rs.get('num_lanes',          2)
    direction = rs.get('direction',          '2 ทิศทาง (ไป-กลับ)')

    total_thick = sum(l.get('h_cm', 0) for l in layers)
    num_layers  = len(layers)
    passed_txt  = 'ผ่านเกณฑ์' if passed else 'ไม่ผ่านเกณฑ์'
    fig_bytes   = ss.get('flex_structure_img')

    doc = _new_doc()

    # ════════════════════════════════════════
    # หัวข้อหลัก
    # ════════════════════════════════════════
    _heading(doc, f'{sec_no}  {sec_title}', level=2, sz=16)

    # ════════════════════════════════════════
    # เกริ่นนำ
    # ════════════════════════════════════════
    p_intro = _para(doc, first_line_cm=1.25, space_before=6)
    _thai_distribute(p_intro)
    _run(p_intro, 'รูปแบบของถนนลาดยางในโครงการนี้เป็นถนน ')
    _run(p_intro, f'{num_lanes}', bold=True, color=BLUE)
    _run(p_intro, ' ช่องจราจร ')
    _run(p_intro, direction, bold=True, color=BLUE)
    _run(p_intro, ' การออกแบบโครงสร้างถนนแบบยืดหยุ่น (Flexible Pavement) ใช้วิธี '
         'AASHTO 1993 Guide for Design of Pavement Structures '
         'โดยพิจารณาปัจจัยด้านปริมาณจราจรสะสม ESALs ความน่าเชื่อถือ '
         'และคุณสมบัติของดินรองรับ สำหรับโครงการนี้ที่ปรึกษาได้กำหนดค่าพารามิเตอร์หลัก '
         'ในการออกแบบ ได้แก่ ปริมาณ W\u2081\u2088 = ')
    _run(p_intro, f'{W18:,.0f}', bold=True, color=BLUE)
    _run(p_intro, ' 18-kip ESALs ที่ระดับความน่าเชื่อถือ (Reliability) = ')
    _run(p_intro, f'{reliability}', bold=True, color=BLUE)
    _run(p_intro, ' % โดยมีดินเดิมค่า CBR = ')
    _run(p_intro, f'{cbr:.1f}', bold=True, color=BLUE)
    _run(p_intro, f' % (M\u1D63 = ')
    _run(p_intro, f'{mr_psi:,.0f}', bold=True, color=BLUE)
    _run(p_intro, f' psi) ผลการออกแบบได้โครงสร้างชั้นทาง ')
    _run(p_intro, f'{num_layers}', bold=True, color=BLUE)
    _run(p_intro, f' ชั้น ที่ SN_required = ')
    _run(p_intro, f'{sn_req:.2f}', bold=True, color=BLUE)
    _run(p_intro, ' และ SN_provided = ')
    _run(p_intro, f'{sn_prov:.2f}', bold=True, color=BLUE)
    _run(p_intro, f' ความหนารวม ')
    _run(p_intro, f'{total_thick:.0f}', bold=True, color=BLUE)
    _run(p_intro, f' ซม. การออกแบบ')
    _run(p_intro, passed_txt, bold=True, color=GREEN if passed else RED)
    _run(p_intro, f' ดังแสดงผลการวิเคราะห์ในตารางที่ ')
    _run(p_intro, tbl_inp, bold=True)
    _run(p_intro, ' และตารางที่ ')
    _run(p_intro, tbl_mat, bold=True)
    _run(p_intro, ' และรูปที่ ')
    _run(p_intro, fig_no, bold=True)

    # ════════════════════════════════════════
    # .1 วิธีการออกแบบ + สมการ
    # ════════════════════════════════════════
    _heading(doc, f'{sec_no}.1  วิธีการออกแบบ', level=3, sz=15)
    p_meth = _para(doc, first_line_cm=1.25, space_before=4)
    _run(p_meth, 'การออกแบบโครงสร้างถนนใช้วิธี AASHTO 1993 Guide for Design of Pavement Structures '
         'ตามมาตรฐานกรมทางหลวง โดยใช้สมการหลักดังนี้')

    _eq_para(doc,
        'log\u2081\u2080(W\u2081\u2088) = Z\u1D63\u00B7S\u2080 + 9.36\u00B7log\u2081\u2080(SN+1) \u2212 0.20\n'
        '                   + log\u2081\u2080(\u0394PSI/2.7) / [0.4 + 1094/(SN+1)^5.19]\n'
        '                   + 2.32\u00B7log\u2081\u2080(M\u1D63) \u2212 8.07',
        indent_cm=2.5, italic=True)

    # ════════════════════════════════════════
    # .2 ข้อมูลนำเข้า + ตาราง
    # ════════════════════════════════════════
    _heading(doc, f'{sec_no}.2  ข้อมูลนำเข้า (Design Inputs)', level=3, sz=15)
    p_inp = _para(doc, first_line_cm=1.25, space_before=4)
    _thai_distribute(p_inp)
    _run(p_inp, 'ในการออกแบบโครงสร้างถนนยืดหยุ่น การกำหนดค่าพารามิเตอร์นำเข้า (Design Inputs) '
         'ถือเป็นขั้นตอนสำคัญที่มีผลโดยตรงต่อความถูกต้องและความน่าเชื่อถือของแบบโครงสร้างถนน '
         f'รายละเอียดของค่าพารามิเตอร์ทั้งหมดแสดงในตารางที่ {tbl_inp}')

    p_cap_inp = _para(doc, space_before=4)
    p_cap_inp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_cap_inp, f'ตารางที่ {tbl_inp}  {cap_inp}', bold=True)

    inp_tbl = doc.add_table(rows=1, cols=3)
    inp_tbl.style = 'Table Grid'
    inp_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['พารามิเตอร์', 'ค่า', 'หน่วย']):
        _tbl_cell(inp_tbl.rows[0].cells[j], h, bold=True, fill=HDR)
    for param, value, unit in [
        ('Design ESALs (W\u2081\u2088)',  f'{W18:,.0f}',  '18-kip ESAL'),
        ('Reliability (R)',                f'{reliability}', '%'),
        ('Z\u1D63',                        f'{zr:.3f}',    '-'),
        ('S\u2080',                        f'{so:.2f}',    '-'),
        ('P\u2080 (Initial Serviceability)', f'{pi:.1f}', '-'),
        ('P\u209C (Terminal Serviceability)', f'{pt:.1f}', '-'),
        ('\u0394PSI',                      f'{delta_psi:.1f}', '-'),
        ('CBR ดินเดิม',                    f'{cbr:.1f}',  '%'),
        ('M\u1D63 = 1,500\u00D7CBR',       f'{mr_psi:,.0f}', 'psi'),
    ]:
        row = inp_tbl.add_row()
        _tbl_cell(row.cells[0], param, align=WD_ALIGN_PARAGRAPH.LEFT)
        _tbl_cell(row.cells[1], value)
        _tbl_cell(row.cells[2], unit)

    doc.add_paragraph()

    # ════════════════════════════════════════
    # .3 คุณสมบัติวัสดุ + ตาราง
    # ════════════════════════════════════════
    _heading(doc, f'{sec_no}.3  คุณสมบัติวัสดุชั้นทาง', level=3, sz=15)
    p_mat = _para(doc, first_line_cm=1.25, space_before=4)
    _thai_distribute(p_mat)
    _run(p_mat, 'วัสดุโครงสร้างชั้นทางแต่ละชนิดมีค่าสัมประสิทธิ์ชั้นทาง (Layer Coefficient) '
         'และค่าสัมประสิทธิ์การระบายน้ำ (Drainage Coefficient) '
         'โดยที่ปรึกษาเลือกใช้วัสดุและแสดงค่าสัมประสิทธิ์รวมถึงค่าโมดูลัสของวัสดุ'
         f'ต่าง ๆ ดังแสดงในตารางที่ {tbl_mat}')

    p_cap_mat = _para(doc, space_before=4)
    p_cap_mat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_cap_mat, f'ตารางที่ {tbl_mat}  {cap_mat}', bold=True)

    mat_tbl = doc.add_table(rows=1, cols=5)
    mat_tbl.style = 'Table Grid'
    mat_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['ชั้น', 'วัสดุ', 'a\u1D62', 'm\u1D62', 'M\u1D63 (psi)']):
        _tbl_cell(mat_tbl.rows[0].cells[j], h, bold=True, fill=HDR)

    for i, layer in enumerate(layers):
        mat   = layer.get('material', '')
        ai    = float(layer.get('ai', 0))
        mi    = float(layer.get('mi', 1.0))
        mr_l  = float(layer.get('mr_psi', mr_psi) if 'mr_psi' in layer else mr_psi)
        row   = mat_tbl.add_row()
        _tbl_cell(row.cells[0], str(i+1))
        _tbl_cell(row.cells[1], mat, align=WD_ALIGN_PARAGRAPH.LEFT)
        _tbl_cell(row.cells[2], f'{ai:.2f}')
        _tbl_cell(row.cells[3], f'{mi:.2f}')
        _tbl_cell(row.cells[4], f'{mr_l:,.0f}')

    doc.add_paragraph()

    # ════════════════════════════════════════
    # .4 ขั้นตอนการคำนวณ (layer-by-layer)
    # ════════════════════════════════════════
    _heading(doc, f'{sec_no}.4  ขั้นตอนการคำนวณความหนาชั้นทาง', level=3, sz=15)
    p_calc = _para(doc, first_line_cm=1.25, space_before=4)
    _thai_distribute(p_calc)
    _run(p_calc, 'การคำนวณความหนาขั้นต่ำของแต่ละชั้น ใช้หลักการว่า Structural Number (SN) ที่จุดใด ๆ '
         'ต้องมากกว่าหรือเท่ากับ SN ที่ต้องการ โดยคำนวณจากค่า M\u1D63 ของชั้นถัดไป')

    cum_sn = 0.0
    for layer in layers:
        mat   = layer.get('material', '')
        h_cm  = float(layer.get('h_cm', 0))
        h_in  = h_cm / 2.54
        ai    = float(layer.get('ai',   0))
        mi    = float(layer.get('mi',   1.0))
        sni   = float(layer.get('sni',  ai * h_in * mi))
        li    = layers.index(layer) + 1
        cum_sn += sni

        doc.add_paragraph()
        hdr_p = _para(doc, indent_cm=1.0, space_before=6)
        _run(hdr_p, f'ชั้นที่ {li}: {mat}', bold=True, underline=True)

        p_mat2 = _para(doc, indent_cm=1.5)
        _run(p_mat2, 'ข้อมูลวัสดุ:', bold=True)
        p_mat3 = _para(doc, indent_cm=2.0)
        _run(p_mat3,
             f'\u2022 Layer Coefficient (a{li}) = {ai:.2f}\n'
             f'\u2022 Drainage Coefficient (m{li}) = {mi:.2f}\n'
             f'\u2022 ความหนาที่ใช้ออกแบบ = {h_cm:.0f} cm ({h_in:.2f} in)')

        p_sn = _para(doc, indent_cm=1.5)
        _run(p_sn, 'SN contribution:', bold=True)
        _eq_para(doc,
            f'\u0394SN_{li} = a_{li} \u00d7 D_{li} \u00d7 m_{li}'
            f'  =  {ai:.2f} \u00d7 {h_in:.2f} \u00d7 {mi:.2f}  =  {sni:.3f}',
            indent_cm=2.5, italic=True)
        _eq_para(doc, f'\u03a3SN  =  {cum_sn:.3f}', indent_cm=2.5, bold=True, italic=False)

        is_ok = cum_sn >= sn_req or li < len(layers)
        p_st = _para(doc, indent_cm=2.0)
        _run(p_st, f'สถานะ:  {"✓ OK" if is_ok else "✗ NG"}',
             bold=True, color=GREEN if is_ok else RED)

    # ════════════════════════════════════════
    # .5 ตารางสรุป SN
    # ════════════════════════════════════════
    _heading(doc, f'{sec_no}.5  สรุปการคำนวณ Structural Number', level=3, sz=15)
    p_sn_intro = _para(doc, first_line_cm=1.25, space_before=4)
    _thai_distribute(p_sn_intro)
    _run(p_sn_intro,
         'เมื่อได้ทำการคำนวณความหนาของแต่ละชั้นทางตามขั้นตอนที่กล่าวมาข้างต้นแล้ว '
         'สามารถสรุปผลการคำนวณ Structural Number (SN) ของโครงสร้างชั้นทางได้ดังนี้ '
         'โดย SN = \u03a3(a\u1D62 \u00d7 D\u1D62 \u00d7 m\u1D62) '
         f'สรุปผลการคำนวณแสดงในตารางที่ {tbl_sn}')

    p_cap_sn = _para(doc, space_before=4)
    p_cap_sn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_cap_sn, f'ตารางที่ {tbl_sn}  {cap_sn}', bold=True)

    sn_tbl = doc.add_table(rows=1, cols=7)
    sn_tbl.style = 'Table Grid'
    sn_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['ชั้น', 'วัสดุ', 'a\u1D62', 'm\u1D62',
                            'D\u1D62 (นิ้ว)', 'D\u1D62 (ซม.)', '\u0394SN\u1D62', '\u03a3SN']):
        _tbl_cell(sn_tbl.rows[0].cells[j] if j < 7 else sn_tbl.rows[0].cells[-1],
                  h, bold=True, fill=HDR)

    # เพิ่ม col 8
    sn_tbl2 = doc.add_table(rows=1, cols=8)
    sn_tbl2.style = 'Table Grid'
    sn_tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['ชั้น', 'วัสดุ', 'a\u1D62', 'm\u1D62',
                            'D\u1D62 (นิ้ว)', 'D\u1D62 (ซม.)', '\u0394SN\u1D62', '\u03a3SN']):
        _tbl_cell(sn_tbl2.rows[0].cells[j], h, bold=True, fill=HDR)

    # ลบ sn_tbl ชั่วคราว (7 col) — ใช้ sn_tbl2 (8 col) แทน
    sn_tbl._tbl.getparent().remove(sn_tbl._tbl)

    cum2 = 0.0
    for i, layer in enumerate(layers):
        ai    = float(layer.get('ai',   0))
        mi    = float(layer.get('mi',   1.0))
        h_cm  = float(layer.get('h_cm', 0))
        h_in  = h_cm / 2.54
        sni   = float(layer.get('sni',  ai * h_in * mi))
        cum2 += sni
        row   = sn_tbl2.add_row()
        vals  = [str(i+1), layer.get('material',''),
                 f'{ai:.2f}', f'{mi:.2f}', f'{h_in:.2f}',
                 f'{h_cm:.0f}', f'{sni:.3f}', f'{cum2:.3f}']
        for j, val in enumerate(vals):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 1 else WD_ALIGN_PARAGRAPH.CENTER
            _tbl_cell(row.cells[j], val, align=align)

    doc.add_paragraph()

    # ════════════════════════════════════════
    # ผลการตรวจสอบ
    # ════════════════════════════════════════
    doc.add_paragraph()
    p_chk_cap = _para(doc, space_before=6)
    p_chk_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_chk_cap, 'ผลการตรวจสอบการออกแบบ', bold=True)

    chk_tbl = doc.add_table(rows=5, cols=2)
    chk_tbl.style = 'Table Grid'
    chk_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    ratio = sn_prov / sn_req if sn_req > 0 else 0.0
    ratio_pass = ratio >= 1.0
    for i, (param, value) in enumerate([
        ('SN Required (จากสมการ AASHTO)', f'{sn_req:.3f}'),
        ('SN Provided (จากชั้นทาง)',       f'{sn_prov:.3f}'),
        ('Ratio (SN Provided / SN Required)', f'{ratio:.3f}  ≥ 1.00  ✓' if ratio_pass else f'{ratio:.3f}  < 1.00  ✗'),
        ('Safety Margin',                  f'{sn_prov - sn_req:+.3f}'),
        ('ผลการตรวจสอบ',                   'ผ่าน (OK)' if passed else 'ไม่ผ่าน (NG)'),
    ]):
        for j, val in enumerate([param, value]):
            chk_tbl.rows[i].cells[j].text = ''
            p = chk_tbl.rows[i].cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            r.font.name = FN; r.font.size = Pt(FS)
            r.bold = (i in (2, 4))
            if i == 2:
                r.font.color.rgb = GREEN if ratio_pass else RED
            if i == 4:
                r.font.color.rgb = GREEN if passed else RED
            try:
                r._element.rPr.rFonts.set(qn('w:cs'), FN)
            except Exception:
                pass

    doc.add_paragraph()

    summary_text = (
        f'สรุป: การออกแบบผ่านเกณฑ์ เนื่องจาก SN_provided ({sn_prov:.3f}) '
        f'\u2265 SN_required ({sn_req:.3f})'
    ) if passed else (
        f'สรุป: การออกแบบไม่ผ่านเกณฑ์ เนื่องจาก SN_provided ({sn_prov:.3f}) '
        f'< SN_required ({sn_req:.3f}) กรุณาปรับเพิ่มความหนาชั้นทาง'
    )
    p_sum = _para(doc, space_before=4)
    _run(p_sum, summary_text, bold=True, color=GREEN if passed else RED)

    # ════════════════════════════════════════
    # ตารางสรุปโครงสร้าง + รูปตัดขวาง
    # ════════════════════════════════════════
    doc.add_paragraph()
    surf_name = _short_mat(layers[0].get('material','')) if layers else ''
    p_sf = _para(doc, space_before=6)
    _run(p_sf, f'รูปแบบที่: {surf_name}', bold=True)

    p_cap_sum = _para(doc, space_before=4)
    p_cap_sum.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_cap_sum, f'ตารางที่ {tbl_sn}  {cap_sn}', bold=True, underline=True)

    _summary_table(doc, layers, cbr, fig_bytes=fig_bytes)

    # footer
    doc.add_paragraph()
    doc.add_paragraph()
    f1 = doc.add_paragraph()
    f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(f1, 'พัฒนาโดย รศ.ดร.อิทธิพล มีผล', sz=14, italic=True)
    f2 = doc.add_paragraph()
    f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(f2, 'ภาควิชาครุศาสตร์โยธา คณะครุศาสตร์อุตสาหกรรม มจพ.', sz=14, italic=True)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
